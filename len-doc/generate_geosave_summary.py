#!/usr/bin/env python3
"""
Generate a simple, 2-page Executive Summary (.docx) for the GeoSave Platform.
The document is concise and illustrated with diagrams (bagan) rendered as images.

Output: len-doc/GeoSave_Executive_Summary.docx
"""

import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


HERE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(HERE, "_assets")
OUTPUT = os.path.join(HERE, "GeoSave_Executive_Summary.docx")

# Color palette
C_USER = "#2E75B6"     # blue
C_DEV = "#548235"      # green
C_INFRA = "#7F7F7F"    # gray
C_ACCENT = "#C55A11"   # orange
C_LIGHT_USER = "#DEEBF7"
C_LIGHT_DEV = "#E2EFDA"
C_LIGHT_INFRA = "#EDEDED"


# =============================================================================
# DIAGRAM HELPERS (matplotlib)
# =============================================================================

def _box(ax, x, y, w, h, text, face, edge, fontsize=10, textcolor="black", bold=False):
    box = FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.02,rounding_size=0.06",
        linewidth=1.4, edgecolor=edge, facecolor=face, zorder=2,
    )
    ax.add_patch(box)
    ax.text(
        x + w / 2, y + h / 2, text,
        ha="center", va="center", fontsize=fontsize, zorder=3,
        color=textcolor, fontweight="bold" if bold else "normal", wrap=True,
    )


def _arrow(ax, x1, y1, x2, y2, color="#404040"):
    ax.add_patch(FancyArrowPatch(
        (x1, y1), (x2, y2),
        arrowstyle="-|>", mutation_scale=16,
        linewidth=1.6, color=color, zorder=1,
    ))


def make_two_layer_diagram(path):
    """High-level diagram: two apps connected through Pipeline / API."""
    fig, ax = plt.subplots(figsize=(7.2, 2.7), dpi=200)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 5)
    ax.axis("off")

    # Dev Apps (left)
    _box(ax, 0.3, 2.7, 3.4, 1.9, "Dev Apps\n(Build • Train • Audit • Bundle)",
         C_LIGHT_DEV, C_DEV, fontsize=10, bold=True)
    # User Apps (right)
    _box(ax, 8.3, 2.7, 3.4, 1.9, "User Apps\n(AOI • Model • Run • Visualize)",
         C_LIGHT_USER, C_USER, fontsize=10, bold=True)
    # Middle: pipeline / API
    _box(ax, 4.1, 3.0, 3.8, 1.3, "Microservice API\n+ Model Registry",
         "#FCE4D6", C_ACCENT, fontsize=10, bold=True)

    # Bottom: data/artifacts
    _box(ax, 2.4, 0.3, 7.2, 1.4,
         "Data • Configs • Checkpoints • Metadata • Evaluation Reports",
         C_LIGHT_INFRA, C_INFRA, fontsize=9.5)

    # Arrows
    _arrow(ax, 3.7, 3.65, 4.1, 3.65, C_DEV)         # dev -> api
    _arrow(ax, 7.9, 3.65, 8.3, 3.65, C_ACCENT)      # api -> user
    _arrow(ax, 6.0, 3.0, 6.0, 1.7, C_INFRA)         # api <-> data
    _arrow(ax, 6.0, 1.7, 6.0, 3.0, C_INFRA)

    ax.text(6.0, 4.75, "GeoSave Platform — Two Connected Layers",
            ha="center", va="center", fontsize=11, fontweight="bold")

    fig.tight_layout(pad=0.2)
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def make_workflow_diagram(path):
    """End-to-end horizontal workflow: Dev pipeline -> Registry -> User flow."""
    fig, ax = plt.subplots(figsize=(7.2, 3.4), dpi=200)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis("off")

    # --- Dev Apps row (top) ---
    ax.text(0.1, 6.6, "Dev Apps", fontsize=10, fontweight="bold", color=C_DEV)
    dev_steps = ["AOI +\nDatetime", "Data\nIngestion", "Band\nSelection",
                 "Preprocess", "Training", "Eval &\nAudit"]
    n = len(dev_steps)
    w, h, gap = 1.7, 1.05, 0.18
    x0 = 0.1
    y = 5.2
    centers = []
    for i, s in enumerate(dev_steps):
        x = x0 + i * (w + gap)
        _box(ax, x, y, w, h, s, C_LIGHT_DEV, C_DEV, fontsize=8.6)
        centers.append((x + w / 2, y))
        if i > 0:
            _arrow(ax, x - gap, y + h / 2, x, y + h / 2, C_DEV)

    # --- Registry / API (middle) ---
    reg_x, reg_y, reg_w, reg_h = 3.4, 3.1, 5.0, 1.1
    _box(ax, reg_x, reg_y, reg_w, reg_h,
         "Pipeline Bundle  →  Microservice API  →  Model Registry",
         "#FCE4D6", C_ACCENT, fontsize=9.2, bold=True)
    # arrow from last dev step down to registry
    last_cx = centers[-1][0]
    _arrow(ax, last_cx, y, last_cx, reg_y + reg_h, C_ACCENT)
    # arrow from registry down to user row
    _arrow(ax, reg_x + reg_w / 2, reg_y, reg_x + reg_w / 2, 2.05, C_USER)

    # --- User Apps row (bottom) ---
    ax.text(0.1, 1.95, "User Apps", fontsize=10, fontweight="bold", color=C_USER)
    user_steps = ["Select\nAOI", "Select\nModel", "Set\nParameters",
                  "Run\nAnalysis", "Inference\nAPI", "Map + Metrics\n+ Report"]
    y2 = 0.5
    for i, s in enumerate(user_steps):
        x = x0 + i * (w + gap)
        face = C_LIGHT_USER if i < 5 else "#FFF2CC"
        edge = C_USER if i < 5 else C_ACCENT
        _box(ax, x, y2, w, h, s, face, edge, fontsize=8.6)
        if i > 0:
            _arrow(ax, x - gap, y2 + h / 2, x, y2 + h / 2, C_USER)

    ax.text(6.0, 6.75, "End-to-End Workflow",
            ha="center", va="center", fontsize=11, fontweight="bold")

    fig.tight_layout(pad=0.2)
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


# =============================================================================
# DOCX HELPERS
# =============================================================================

FONT = "Calibri"


def set_run(run, size=10.5, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)


def para(doc, text, size=10.5, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         space_after=4, space_before=0, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, italic=italic, color=color)
    return p


def heading(doc, text, color=C_ACCENT, size=12, space_before=6):
    return para(doc, text, size=size, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=space_before, space_after=3, color=color)


def bullet(doc, text, size=9.8):
    p = doc.add_paragraph(style=None)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("\u2022  " + text)
    set_run(r, size=size)
    return p


def add_image(doc, path, width_cm):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run()
    r.add_picture(path, width=Cm(width_cm))
    return p


def caption(doc, text):
    return para(doc, text, size=8.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=5, color="#595959")


def set_cell(cell, text, bold=False, size=8.8, align=WD_ALIGN_PARAGRAPH.LEFT,
             color=None, shade=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    if shade:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shade}"/>')
        cell._tc.get_or_add_tcPr().append(shading)


# =============================================================================
# BUILD DOCUMENT
# =============================================================================

def build():
    os.makedirs(IMG_DIR, exist_ok=True)
    two_layer = os.path.join(IMG_DIR, "two_layer.png")
    workflow = os.path.join(IMG_DIR, "workflow.png")
    make_two_layer_diagram(two_layer)
    make_workflow_diagram(workflow)

    doc = Document()

    # Base style
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.0

    # Page setup: A4, tighter margins to fit 2 pages
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # ---- Title ----
    para(doc, "GeoSave Platform", size=18, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1, color=C_USER)
    para(doc, "Executive Summary — User Apps & Dev Apps untuk Geospatial AI Workflow",
         size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=5,
         color="#404040")

    # ---- Background + Vision ----
    heading(doc, "Latar Belakang & Visi", space_before=0)
    para(doc,
         "GeoSave Platform adalah ekosistem aplikasi untuk mempercepat pemanfaatan model "
         "AI geospasial — dari pengembangan dan audit model, deployment pipeline, hingga "
         "penggunaan model melalui analisis berbasis Area of Interest (AOI). Platform dibagi "
         "menjadi dua aplikasi yang saling terhubung: Dev Apps untuk membangun, menguji, "
         "mengaudit, dan membundel model menjadi microservice API; serta User Apps untuk "
         "menjalankan inference dan analisis.")
    para(doc,
         "Visi: membuat proses pengembangan, audit, deployment, dan penggunaan model AI "
         "geospasial menjadi lebih cepat, terstruktur, reproducible, dan mudah dioperasikan "
         "baik oleh user non-teknis maupun tim developer.", italic=True)

    # ---- Diagram 1 ----
    add_image(doc, two_layer, width_cm=15.5)
    caption(doc, "Gambar 1. Arsitektur dua layer: Dev Apps membangun model, "
                 "User Apps menggunakannya melalui Microservice API & Model Registry.")

    # ---- Two columns of roles ----
    heading(doc, "Dua Aplikasi Utama")
    bullet(doc, "User Apps (operasional): pilih AOI, pilih model, atur parameter, jalankan "
                "analysis, dan lihat hasil sebagai map layer, statistik, atau report — tanpa "
                "perlu memahami detail teknis model.")
    bullet(doc, "Dev Apps (pengembangan): ingest data berdasarkan AOI & datetime, band "
                "selection, preprocessing, training, evaluasi & audit, lalu bundle model "
                "menjadi pipeline dan expose sebagai microservice API.")

    # ---- Diagram 2 ----
    add_image(doc, workflow, width_cm=16.5)
    caption(doc, "Gambar 2. Alur end-to-end: pipeline Dev Apps menghasilkan model "
                 "ter-registrasi yang dikonsumsi oleh alur analisis User Apps.")

    # ---- Core capabilities table ----
    heading(doc, "Kapabilitas Inti")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = ["Layer", "Capability", "Tujuan"]
    for i, h in enumerate(hdr):
        set_cell(table.rows[0].cells[i], h, bold=True, size=9,
                 align=WD_ALIGN_PARAGRAPH.CENTER, color="FFFFFF", shade=C_ACCENT.lstrip("#"))
    rows = [
        ("User", "AOI & Model selection", "Menentukan area dan model inference"),
        ("User", "Multi-parameter analysis", "Menjalankan prediksi dengan parameter berbeda"),
        ("User", "Result visualization", "Output sebagai peta, layer, statistik, report"),
        ("Dev", "Ingestion & Band selection", "Mengambil data per AOI/datetime, pilih band"),
        ("Dev", "Preprocessing & Training", "Standarisasi input dan membangun model"),
        ("Dev", "Audit & Pipeline bundling", "Validasi model, kemas jadi pipeline konsisten"),
        ("Dev", "Microservice API", "Endpoint inference untuk User Apps"),
    ]
    for layer, cap, purpose in rows:
        cells = table.add_row().cells
        shade = C_LIGHT_USER.lstrip("#") if layer == "User" else C_LIGHT_DEV.lstrip("#")
        tcolor = C_USER if layer == "User" else C_DEV
        set_cell(cells[0], layer, bold=True, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER,
                 color=tcolor, shade=shade)
        set_cell(cells[1], cap, size=8.8)
        set_cell(cells[2], purpose, size=8.8)
    for row in table.rows:
        row.cells[0].width = Cm(1.6)
        row.cells[1].width = Cm(5.2)
        row.cells[2].width = Cm(10.5)

    # ---- Governance ----
    heading(doc, "Tata Kelola & Audit")
    para(doc,
         "Model hanya tersedia di User Apps setelah melewati status: "
         "Draft \u2192 Trained \u2192 Evaluated \u2192 Audited \u2192 Approved \u2192 Published. "
         "Setiap model dibundel dengan metadata lengkap (sumber dataset, AOI & datetime, band, "
         "preprocessing, konfigurasi, metrik training/validasi, versi model & pipeline, "
         "parameter inference, dan limitation notes) agar hasil dapat ditelusuri dan dipercaya.")

    # ---- Strategic value + conclusion ----
    heading(doc, "Nilai Strategis & Kesimpulan")
    bullet(doc, "User akhir: menjalankan analisis tanpa memahami proses teknis model.")
    bullet(doc, "Developer: workflow standar dari data mentah hingga API, reproducible.")
    bullet(doc, "Organisasi: pemisahan pengembangan dan penggunaan model menurunkan risiko "
                "pemakaian model yang belum tervalidasi.")
    para(doc,
         "Dengan memisahkan pengembangan (Dev Apps) dan penggunaan model (User Apps) yang "
         "terhubung lewat microservice API, GeoSave Platform menjadi fondasi operasional "
         "geospatial AI: model dikembangkan terstruktur, divalidasi melalui audit, "
         "dipublikasikan sebagai API, lalu digunakan dengan mudah — mempercepat deployment, "
         "menjaga konsistensi inference, dan memperkuat tata kelola model.",
         space_before=2)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Size : {os.path.getsize(OUTPUT)/1024:.1f} KB")


if __name__ == "__main__":
    build()
