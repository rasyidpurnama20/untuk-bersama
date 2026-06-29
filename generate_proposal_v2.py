"""
Generate PhD Proposal Document (Version 2)
==========================================
This script generates a formal PhD dissertation proposal document in .docx format
for the ITB Doctoral Program in Informatics.

Topic: Autonomous Privacy and Security Risk Prevention Framework
       for AI-Assisted Software Engineering

Output: phd-itb-2026-v2/proposal_disertasi_v2.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# =============================================================================
# CONFIGURATION
# =============================================================================

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "phd-itb-2026-v2")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "proposal_disertasi_v2.docx")

# Page setup
PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_LEFT_CM = 4.0
MARGIN_RIGHT_CM = 3.0
MARGIN_TOP_CM = 3.0
MARGIN_BOTTOM_CM = 3.0

# Font settings
DEFAULT_FONT = "Times New Roman"
DEFAULT_FONT_SIZE = Pt(12)
HEADING1_SIZE = Pt(14)
TITLE_SIZE = Pt(16)
LINE_SPACING = 1.5


# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def create_document():
    """Create a new Document with A4 page setup and default formatting."""
    doc = Document()

    # Set page size and margins
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.left_margin = Cm(MARGIN_LEFT_CM)
    section.right_margin = Cm(MARGIN_RIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_CM)
    section.bottom_margin = Cm(MARGIN_BOTTOM_CM)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = DEFAULT_FONT
    font.size = DEFAULT_FONT_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)

    # Set default paragraph format
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = LINE_SPACING
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)

    return doc


def set_run_font(run, font_name=DEFAULT_FONT, font_size=DEFAULT_FONT_SIZE,
                 bold=False, italic=False):
    """Set font properties for a run."""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_centered_text(doc, text, font_size=DEFAULT_FONT_SIZE, bold=False,
                      space_before=Pt(0), space_after=Pt(0)):
    """Add a centered paragraph with specified formatting."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = space_before
    para.paragraph_format.space_after = space_after
    run = para.add_run(text)
    set_run_font(run, font_size=font_size, bold=bold)
    return para


def add_blank_line(doc, count=1):
    """Add one or more blank lines."""
    for _ in range(count):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run("")
        set_run_font(run)
    return para


def add_justified_paragraph(doc, text, first_line_indent=Cm(1.25),
                            space_after=Pt(6)):
    """Add a justified paragraph with first line indent."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = first_line_indent
    para.paragraph_format.space_after = space_after
    run = para.add_run(text)
    set_run_font(run)
    return para


def add_chapter_heading(doc, text):
    """Add a chapter heading (BAB level) - bold, 14pt."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run(text)
    set_run_font(run, font_size=HEADING1_SIZE, bold=True)
    return para


def add_section_heading(doc, text):
    """Add a section heading (1.1 level) - bold, 12pt."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    set_run_font(run, font_size=DEFAULT_FONT_SIZE, bold=True)
    return para


def add_subsection_heading(doc, text):
    """Add a subsection heading (2.4.1 level) - bold, 12pt."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.left_indent = Cm(0.5)
    run = para.add_run(text)
    set_run_font(run, font_size=DEFAULT_FONT_SIZE, bold=True)
    return para


def add_numbered_item(doc, number, text, bold_prefix="", indent=Cm(1.0)):
    """Add a numbered list item."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.left_indent = indent
    para.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = para.add_run(f"{number}. {bold_prefix}")
        set_run_font(run, bold=True)
        run2 = para.add_run(f" {text}")
        set_run_font(run2)
    else:
        run = para.add_run(f"{number}. {text}")
        set_run_font(run)
    return para


def add_research_question(doc, rq_label, text):
    """Add a research question with bold label."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.left_indent = Cm(1.0)
    para.paragraph_format.space_after = Pt(6)
    run_label = para.add_run(f"{rq_label}: ")
    set_run_font(run_label, bold=True)
    run_text = para.add_run(text)
    set_run_font(run_text)
    return para


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


def set_cell_text(cell, text, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  font_size=Pt(10)):
    """Set text in a table cell with formatting."""
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = alignment
    run = para.add_run(text)
    run.font.name = DEFAULT_FONT
    run.font.size = font_size
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)


def set_cell_shading(cell, color="D9E2F3"):
    """Set cell background shading."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


# =============================================================================
# DOCUMENT SECTIONS
# =============================================================================

def create_cover_page(doc):
    """Create the cover page of the proposal."""
    add_blank_line(doc, 3)
    add_centered_text(doc, "PROPOSAL DISERTASI", font_size=TITLE_SIZE, bold=True)
    add_blank_line(doc)

    add_centered_text(doc, "AUTONOMOUS PRIVACY AND SECURITY RISK PREVENTION FRAMEWORK",
                      font_size=HEADING1_SIZE, bold=True)
    add_centered_text(doc, "FOR AI-ASSISTED SOFTWARE ENGINEERING",
                      font_size=HEADING1_SIZE, bold=True)

    add_blank_line(doc, 3)
    add_centered_text(doc, "Diajukan untuk memenuhi salah satu syarat")
    add_centered_text(doc, "Program Doktor Informatika")

    add_blank_line(doc, 3)
    add_centered_text(doc, "oleh:")
    add_centered_text(doc, "[Nama Mahasiswa]")
    add_centered_text(doc, "[NIM]")

    add_blank_line(doc, 4)
    add_centered_text(doc, "PROGRAM STUDI DOKTOR INFORMATIKA", bold=True)
    add_centered_text(doc, "SEKOLAH TEKNIK ELEKTRO DAN INFORMATIKA", bold=True)
    add_centered_text(doc, "INSTITUT TEKNOLOGI BANDUNG", bold=True)
    add_centered_text(doc, "2026", bold=True)

    add_page_break(doc)


def create_bab1(doc):
    """Create BAB I - PENDAHULUAN."""
    add_chapter_heading(doc, "BAB I. PENDAHULUAN")

    # 1.1 Latar Belakang
    add_section_heading(doc, "1.1 Latar Belakang")

    add_justified_paragraph(
        doc,
        "Perkembangan kecerdasan buatan generatif telah mengubah praktik rekayasa "
        "perangkat lunak secara signifikan [1][2]. Penggunaan AI-assisted software "
        "engineering, seperti GitHub Copilot, ChatGPT, Cursor, Claude Code, Gemini "
        "Code Assist, dan berbagai agentic coding tools, memungkinkan pengembang "
        "menghasilkan kode, memperbaiki bug, membuat dokumentasi, menulis unit test, "
        "melakukan code review, dan mengotomasi sebagian proses pengembangan "
        "perangkat lunak [3][4][5]."
    )

    add_justified_paragraph(
        doc,
        "Meskipun teknologi ini memberikan manfaat besar terhadap produktivitas, "
        "efisiensi, dan percepatan pengembangan perangkat lunak [6], penggunaan AI "
        "dalam software engineering juga menimbulkan risiko baru. Kode yang dihasilkan "
        "AI dapat mengandung kerentanan keamanan, seperti SQL injection, cross-site "
        "scripting, insecure authentication, hardcoded credentials, penggunaan library "
        "yang tidak aman, dan kesalahan konfigurasi [7][8][9]. Selain itu, proses "
        "interaksi antara pengembang dan AI juga dapat menyebabkan risiko privasi, "
        "seperti kebocoran source code, API key, token, data pelanggan, data bisnis "
        "rahasia, atau informasi sensitif lain yang dimasukkan ke dalam prompt [10][11]."
    )

    add_justified_paragraph(
        doc,
        "Permasalahan tersebut menjadi semakin penting karena banyak organisasi mulai "
        "mengintegrasikan AI coding tools ke dalam software development lifecycle tanpa "
        "mekanisme pencegahan, deteksi, dan mitigasi risiko yang memadai [12][13]. "
        "Pendekatan keamanan tradisional seperti static application security testing "
        "dan manual code review belum sepenuhnya dirancang untuk konteks AI-assisted "
        "software engineering yang bersifat cepat, iteratif, dan berbasis interaksi "
        "prompt-code [14][15]."
    )

    add_justified_paragraph(
        doc,
        "Oleh karena itu, diperlukan sebuah framework yang mampu secara otonom "
        "mencegah, mendeteksi, dan memitigasi risiko privasi serta keamanan pada "
        "seluruh tahapan AI-assisted software engineering. Penelitian ini mengusulkan "
        "pengembangan Autonomous Privacy and Security Risk Prevention Framework yang "
        "mengintegrasikan analisis prompt, deteksi kerentanan kode, identifikasi "
        "kebocoran data sensitif, prioritisasi risiko, perbaikan otomatis, dan "
        "verifikasi keamanan sebelum kode masuk ke repository atau pipeline "
        "pengembangan [16][17]."
    )

    # 1.2 Rumusan Masalah
    add_section_heading(doc, "1.2 Rumusan Masalah")

    add_justified_paragraph(
        doc,
        "Berdasarkan latar belakang tersebut, permasalahan utama dalam penelitian ini "
        "adalah belum tersedianya framework teknis yang mampu secara otonom mencegah, "
        "mendeteksi, dan memitigasi risiko privasi serta keamanan pada AI-assisted "
        "software engineering."
    )

    add_justified_paragraph(
        doc,
        "Permasalahan tersebut dapat dirinci sebagai berikut:"
    )

    add_numbered_item(doc, 1,
                      "AI-generated code berpotensi mengandung kerentanan keamanan "
                      "yang tidak terdeteksi oleh pengembang.")
    add_numbered_item(doc, 2,
                      "Prompt yang digunakan dalam AI-assisted development dapat "
                      "memuat data sensitif, source code rahasia, credential, atau "
                      "informasi privat.")
    add_numbered_item(doc, 3,
                      "Proses software engineering berbantuan AI belum memiliki "
                      "mekanisme otomatis untuk melakukan risk prevention, risk "
                      "detection, dan risk mitigation.")
    add_numbered_item(doc, 4,
                      "Belum terdapat framework terintegrasi yang menghubungkan "
                      "analisis prompt, analisis kode, mitigasi otomatis, dan "
                      "verifikasi keamanan dalam satu pipeline.")

    # 1.3 Pertanyaan Penelitian
    add_section_heading(doc, "1.3 Pertanyaan Penelitian")

    add_justified_paragraph(
        doc,
        "Penelitian ini diarahkan untuk menjawab pertanyaan penelitian berikut:"
    )

    add_research_question(doc, "RQ1",
                          "Risiko privasi dan keamanan apa saja yang muncul dalam "
                          "proses AI-assisted software engineering?")
    add_research_question(doc, "RQ2",
                          "Bagaimana metode otomatis dapat digunakan untuk mendeteksi "
                          "risiko privasi pada prompt dan artefak perangkat lunak yang "
                          "digunakan dalam AI-assisted software engineering?")
    add_research_question(doc, "RQ3",
                          "Bagaimana metode otomatis dapat digunakan untuk mendeteksi "
                          "risiko keamanan pada kode yang dihasilkan oleh AI coding "
                          "assistant?")
    add_research_question(doc, "RQ4",
                          "Bagaimana mekanisme mitigasi otomatis dapat dirancang untuk "
                          "memperbaiki atau mengurangi risiko privasi dan keamanan pada "
                          "AI-generated code?")
    add_research_question(doc, "RQ5",
                          "Seberapa efektif framework yang diusulkan dalam meningkatkan "
                          "keamanan, menjaga privasi, dan mengurangi risiko pada "
                          "AI-assisted software engineering?")

    # 1.4 Tujuan Penelitian
    add_section_heading(doc, "1.4 Tujuan Penelitian")

    add_justified_paragraph(
        doc,
        "Tujuan utama penelitian ini adalah merancang, mengembangkan, dan mengevaluasi "
        "Autonomous Privacy and Security Risk Prevention Framework untuk AI-Assisted "
        "Software Engineering yang mampu secara otomatis mencegah, mendeteksi, dan "
        "memitigasi risiko privasi serta keamanan."
    )

    add_justified_paragraph(
        doc,
        "Secara spesifik, tujuan penelitian ini adalah:",
        first_line_indent=Cm(0)
    )

    add_numbered_item(doc, 1,
                      "Mengidentifikasi dan mengklasifikasikan risiko privasi dan "
                      "keamanan yang muncul dalam AI-assisted software engineering.")
    add_numbered_item(doc, 2,
                      "Mengembangkan model deteksi otomatis untuk risiko privasi pada "
                      "prompt, context, dan artefak perangkat lunak.")
    add_numbered_item(doc, 3,
                      "Mengembangkan model deteksi otomatis untuk kerentanan keamanan "
                      "pada kode yang dihasilkan AI.")
    add_numbered_item(doc, 4,
                      "Merancang mekanisme mitigasi otomatis untuk memperbaiki atau "
                      "mengurangi risiko yang terdeteksi.")
    add_numbered_item(doc, 5,
                      "Mengintegrasikan seluruh komponen ke dalam satu framework yang "
                      "dapat digunakan dalam pipeline pengembangan perangkat lunak.")
    add_numbered_item(doc, 6,
                      "Mengevaluasi efektivitas framework dalam meningkatkan keamanan "
                      "dan menjaga privasi pada AI-assisted software engineering.")

    # 1.5 Manfaat Penelitian
    add_section_heading(doc, "1.5 Manfaat Penelitian")

    add_justified_paragraph(
        doc,
        "Secara teoretis, penelitian ini berkontribusi pada pengembangan body of knowledge "
        "di bidang software security, privacy engineering, dan AI-assisted software "
        "engineering. Penelitian ini menghasilkan risk taxonomy, model deteksi, dan "
        "framework konseptual yang dapat digunakan sebagai referensi bagi peneliti lain "
        "dalam mengembangkan solusi keamanan dan privasi untuk konteks AI-assisted "
        "development."
    )

    add_justified_paragraph(
        doc,
        "Secara praktis, framework yang dihasilkan dapat diimplementasikan oleh "
        "organisasi pengembang perangkat lunak untuk melindungi proses pengembangan "
        "dari risiko privasi dan keamanan saat menggunakan AI coding tools. Framework "
        "ini juga memberikan panduan teknis bagi tim keamanan dalam merancang kebijakan "
        "dan kontrol untuk AI-assisted software engineering."
    )

    # 1.6 Ruang Lingkup Penelitian
    add_section_heading(doc, "1.6 Ruang Lingkup Penelitian")

    add_justified_paragraph(
        doc,
        "Ruang lingkup penelitian ini mencakup:"
    )

    add_numbered_item(doc, 1,
                      "Risiko privasi dan keamanan yang muncul dalam konteks "
                      "AI-assisted software engineering, khususnya penggunaan LLM-based "
                      "coding assistants.")
    add_numbered_item(doc, 2,
                      "Deteksi risiko privasi pada prompt, code context, dan artefak "
                      "yang dikirimkan ke AI coding assistant.")
    add_numbered_item(doc, 3,
                      "Deteksi kerentanan keamanan pada kode yang dihasilkan oleh AI "
                      "coding assistant.")
    add_numbered_item(doc, 4,
                      "Mitigasi otomatis berupa code repair, sanitization, dan "
                      "redaction untuk risiko yang terdeteksi.")
    add_numbered_item(doc, 5,
                      "Evaluasi framework menggunakan dataset sintetis dan benchmark "
                      "yang relevan.")

    add_justified_paragraph(
        doc,
        "Penelitian ini tidak berfokus pada: (1) pengembangan model bahasa besar "
        "baru, (2) aspek hukum dan regulasi privasi secara mendalam, (3) serangan "
        "adversarial terhadap model AI itu sendiri, dan (4) evaluasi usability dari "
        "perspektif human-computer interaction."
    )

    add_page_break(doc)


def create_bab2(doc):
    """Create BAB II - TINJAUAN PUSTAKA."""
    add_chapter_heading(doc, "BAB II. TINJAUAN PUSTAKA")

    # 2.1 AI-Assisted Software Engineering
    add_section_heading(doc, "2.1 AI-Assisted Software Engineering")

    add_justified_paragraph(
        doc,
        "AI-Assisted Software Engineering (AI-ASE) merupakan paradigma pengembangan "
        "perangkat lunak yang memanfaatkan kecerdasan buatan, khususnya large language "
        "models (LLM), untuk membantu berbagai aktivitas dalam software development "
        "lifecycle. Konsep ini mencakup penggunaan AI untuk code generation, code "
        "completion, bug detection, code review, documentation generation, test "
        "generation, dan refactoring [18][19]."
    )

    add_justified_paragraph(
        doc,
        "Perkembangan LLM-based code generation dimulai dari model seperti CodeBERT "
        "[18], GraphCodeBERT [19], Codex [1], AlphaCode [33], dan CodeGen [34], "
        "hingga model terbaru seperti GPT-4, Claude, dan Gemini yang mampu menghasilkan "
        "kode kompleks dari deskripsi natural language. Model-model ini dilatih pada "
        "miliaran baris kode dari repository publik dan mampu memahami konteks, pola "
        "pemrograman, serta konvensi pengembangan perangkat lunak [20]."
    )

    add_justified_paragraph(
        doc,
        "Paradigma AI-ASE mengubah cara pengembang bekerja dari menulis kode secara "
        "manual menjadi berinteraksi dengan AI melalui prompt, menerima saran kode, "
        "melakukan iterasi, dan memvalidasi output. Pergeseran ini memerlukan "
        "pendekatan baru dalam memastikan keamanan dan privasi, karena interaksi "
        "manusia-AI menciptakan attack surface baru yang tidak ada pada pengembangan "
        "tradisional [5][12]."
    )

    # 2.2 Risiko Keamanan
    add_section_heading(doc, "2.2 Risiko Keamanan dalam AI-Assisted Software Engineering")

    add_justified_paragraph(
        doc,
        "Penggunaan AI dalam pengembangan perangkat lunak menimbulkan berbagai risiko "
        "keamanan yang perlu dipahami dan dimitigasi. Berikut adalah kategori risiko "
        "keamanan utama yang telah diidentifikasi dalam literatur [7][8][9][21][22]:"
    )

    security_risks = [
        ("Vulnerable Code Generation",
         "AI coding assistant dapat menghasilkan kode yang mengandung kerentanan "
         "keamanan seperti SQL injection, buffer overflow, cross-site scripting (XSS), "
         "path traversal, dan command injection. Studi menunjukkan bahwa sekitar 40% "
         "kode yang dihasilkan oleh AI mengandung setidaknya satu kerentanan keamanan."),
        ("Hardcoded Secrets and Credentials",
         "Model AI cenderung menghasilkan kode contoh yang mengandung hardcoded "
         "credentials, API keys, database passwords, dan encryption keys yang "
         "seharusnya disimpan dalam environment variables atau secret management "
         "system."),
        ("Insecure Authentication and Authorization",
         "Kode yang dihasilkan AI sering mengimplementasikan mekanisme autentikasi "
         "dan otorisasi yang tidak aman, seperti penggunaan algoritma hashing yang "
         "lemah, session management yang tidak tepat, atau missing access control."),
        ("Dependency and Supply Chain Risks",
         "AI dapat merekomendasikan library atau package yang sudah deprecated, "
         "memiliki known vulnerabilities, atau bahkan merekomendasikan package yang "
         "tidak ada (hallucinated packages) yang berpotensi menjadi vektor serangan "
         "supply chain."),
        ("Prompt Injection and Agentic Risks",
         "Pada agentic coding tools, terdapat risiko prompt injection dimana instruksi "
         "berbahaya dapat disisipkan dalam code comments, documentation, atau "
         "repository files yang kemudian dieksekusi oleh AI agent."),
        ("Over-reliance and Automation Bias",
         "Pengembang yang terlalu bergantung pada AI cenderung tidak melakukan "
         "review menyeluruh terhadap kode yang dihasilkan, sehingga kerentanan "
         "keamanan dapat lolos ke production tanpa terdeteksi.")
    ]

    for i, (title, desc) in enumerate(security_risks, 1):
        add_numbered_item(doc, i, desc, bold_prefix=f"{title}.")

    # 2.3 Risiko Privasi
    add_section_heading(doc, "2.3 Risiko Privasi dalam AI-Assisted Software Engineering")

    add_justified_paragraph(
        doc,
        "Selain risiko keamanan, penggunaan AI coding tools juga menimbulkan risiko "
        "privasi yang signifikan. Risiko ini muncul dari interaksi antara pengembang "
        "dan AI yang melibatkan pengiriman data sensitif [10][11][23]:"
    )

    privacy_risks = [
        ("Prompt Privacy Leakage",
         "Pengembang secara tidak sengaja memasukkan informasi sensitif seperti "
         "business logic, trade secrets, atau strategi bisnis ke dalam prompt yang "
         "dikirim ke AI service provider."),
        ("Credential and Secret Leakage",
         "API keys, tokens, passwords, dan credentials lainnya dapat terekspos "
         "melalui prompt atau code context yang dikirim ke model AI pihak ketiga."),
        ("Personal Data Exposure",
         "Data personal pelanggan, karyawan, atau pengguna yang terdapat dalam "
         "source code, database queries, atau test data dapat terkirim ke AI "
         "service tanpa consent yang memadai."),
        ("Intellectual Property Leakage",
         "Source code proprietary, algoritma rahasia, arsitektur sistem, dan "
         "desain internal dapat terekspos melalui code completion context yang "
         "dikirim ke cloud-based AI models."),
        ("Data Retention and Training Risk",
         "Data yang dikirim ke AI service provider berpotensi disimpan dan "
         "digunakan untuk melatih model selanjutnya, sehingga informasi sensitif "
         "dapat muncul dalam output model untuk pengguna lain.")
    ]

    for i, (title, desc) in enumerate(privacy_risks, 1):
        add_numbered_item(doc, i, desc, bold_prefix=f"{title}.")

    # 2.4 Prevention, Detection, and Mitigation
    add_section_heading(doc, "2.4 Prevention, Detection, and Mitigation")

    add_subsection_heading(doc, "2.4.1 Prevention")
    add_justified_paragraph(
        doc,
        "Prevention merupakan tahap pertama dalam framework keamanan yang bertujuan "
        "mencegah risiko sebelum terjadi. Dalam konteks AI-assisted software "
        "engineering, prevention mencakup mekanisme yang diterapkan sebelum prompt "
        "atau code context dikirim ke AI model. Teknik prevention meliputi input "
        "sanitization, sensitive data redaction, prompt filtering, policy enforcement, "
        "dan pre-submission validation. Tujuannya adalah memastikan bahwa tidak ada "
        "data sensitif atau informasi rahasia yang terkirim ke model AI eksternal "
        "[24][25]."
    )

    add_subsection_heading(doc, "2.4.2 Detection")
    add_justified_paragraph(
        doc,
        "Detection merupakan tahap identifikasi risiko pada artefak yang sudah "
        "dihasilkan. Dalam konteks ini, detection mencakup analisis kode yang "
        "dihasilkan AI untuk menemukan kerentanan keamanan, analisis output untuk "
        "mendeteksi kebocoran data sensitif, dan monitoring interaksi untuk "
        "mengidentifikasi anomali. Teknik detection meliputi static analysis, "
        "pattern matching, machine learning-based vulnerability detection, secret "
        "scanning, dan taint analysis. Detection dilakukan secara real-time "
        "setelah AI menghasilkan output dan sebelum kode diintegrasikan ke "
        "codebase [26][27][28]."
    )

    add_subsection_heading(doc, "2.4.3 Mitigation")
    add_justified_paragraph(
        doc,
        "Mitigation merupakan tahap perbaikan dan pengurangan dampak risiko yang "
        "telah terdeteksi. Mekanisme mitigation mencakup automated code repair, "
        "vulnerability patching, code transformation, secure code suggestion, "
        "dan rollback mechanisms. Tujuannya adalah memperbaiki kerentanan atau "
        "mengurangi risiko tanpa memerlukan intervensi manual yang signifikan "
        "dari pengembang. Pendekatan modern menggunakan LLM-based repair yang "
        "mampu memahami konteks kerentanan dan menghasilkan perbaikan yang "
        "semantically correct [29][30][31]."
    )

    # 2.5 Penelitian Terdahulu
    add_section_heading(doc, "2.5 Penelitian Terdahulu")

    add_justified_paragraph(
        doc,
        "Berikut adalah ringkasan penelitian terdahulu yang relevan dengan topik "
        "penelitian ini:"
    )

    # Create literature review table
    table = doc.add_table(rows=21, cols=6)
    table.style = 'Table Grid'

    # Header row
    headers = ["No", "Penulis (Tahun)", "Fokus", "Metode", "Temuan", "Gap"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, font_size=Pt(9))
        set_cell_shading(table.rows[0].cells[i])

    # Table data
    literature_data = [
        ("1", "Chen et al. (2021)",
         "Evaluasi LLM untuk code generation",
         "Benchmark HumanEval",
         "Codex menyelesaikan 28.8% masalah pemrograman",
         "Tidak membahas aspek keamanan kode yang dihasilkan"),
        ("2", "Pearce et al. (2022)",
         "Keamanan kode GitHub Copilot",
         "Analisis CWE pada kode output",
         "40% skenario menghasilkan kode vulnerable",
         "Tidak ada mekanisme mitigasi otomatis"),
        ("3", "Perry et al. (2023)",
         "Dampak AI assistant pada keamanan kode user",
         "User study dengan 47 partisipan",
         "Pengguna AI menulis kode lebih insecure",
         "Tidak menyediakan framework pencegahan"),
        ("4", "Sandoval et al. (2023)",
         "Implikasi keamanan LLM code assistant",
         "User study bahasa C",
         "AI-assisted code memiliki lebih banyak bug",
         "Fokus pada satu bahasa, tidak ada solusi otomatis"),
        ("5", "Peng et al. (2023)",
         "Dampak AI pada produktivitas developer",
         "Randomized controlled trial",
         "Copilot meningkatkan produktivitas 55.8%",
         "Tidak mengukur dampak keamanan"),
        ("6", "Ziegler et al. (2022)",
         "Produktivitas neural code completion",
         "Analisis telemetri",
         "30% kode diterima dari suggestion",
         "Tidak membahas kualitas keamanan suggestion"),
        ("7", "Khoury et al. (2023)",
         "Keamanan kode ChatGPT",
         "Analisis 21 program",
         "16/21 program mengandung vulnerability",
         "Tidak ada detection dan repair otomatis"),
        ("8", "Liu et al. (2024)",
         "Kualitas code generation ChatGPT",
         "Multi-dimensional evaluation",
         "Kode ChatGPT memiliki masalah keamanan",
         "Tidak fokus pada privacy risks"),
        ("9", "He & Vechev (2023)",
         "Security hardening untuk LLM code",
         "Adversarial testing",
         "LLM rentan terhadap adversarial prompt",
         "Tidak mencakup full prevention framework"),
        ("10", "Niu et al. (2024)",
         "Privacy leaks pada Copilot",
         "Empirical analysis",
         "Copilot dapat membocorkan training data",
         "Tidak ada real-time prevention mechanism"),
        ("11", "Yang et al. (2024)",
         "Memorization dalam code models",
         "Probing dan extraction",
         "Model menghafal data sensitif dari training",
         "Tidak mengusulkan solusi mitigasi"),
        ("12", "Jiang et al. (2024)",
         "Survey LLM untuk SE",
         "Systematic literature review",
         "Mapping komprehensif penggunaan LLM di SE",
         "Survey tanpa framework implementasi"),
        ("13", "Hou et al. (2024)",
         "Survey LLM untuk SE",
         "Systematic literature review",
         "Identifikasi challenges dan opportunities",
         "Tidak fokus pada security dan privacy"),
        ("14", "Li et al. (2018)",
         "Deep learning vulnerability detection",
         "BiLSTM pada code gadgets",
         "VulDeePecker efektif untuk buffer overflow",
         "Tidak dirancang untuk AI-generated code"),
        ("15", "Fu & Tantithamthavorn (2022)",
         "Line-level vulnerability prediction",
         "Transformer-based model",
         "LineVul lebih akurat dari baseline",
         "Belum terintegrasi dengan AI coding pipeline"),
        ("16", "NIST (2023)",
         "AI Risk Management Framework",
         "Framework development",
         "AI RMF 1.0 sebagai panduan manajemen risiko",
         "Tidak spesifik untuk software engineering"),
        ("17", "ISO/IEC (2022)",
         "Information security management",
         "Standard specification",
         "ISO 27001:2022 requirements",
         "Tidak mencakup AI-specific risks"),
        ("18", "Feng et al. (2020)",
         "Pre-trained model untuk code",
         "Masked language modeling",
         "CodeBERT efektif untuk code understanding",
         "Tidak digunakan untuk security detection"),
        ("19", "Guo et al. (2021)",
         "Code representation dengan data flow",
         "Graph-based pre-training",
         "GraphCodeBERT memahami semantik kode",
         "Belum diaplikasikan untuk vulnerability detection"),
        ("20", "Austin et al. (2021)",
         "Program synthesis dengan LLM",
         "Benchmark evaluation",
         "LLM mampu program synthesis",
         "Tidak membahas keamanan output")
    ]

    for row_idx, (no, author, focus, method, finding, gap) in enumerate(literature_data, 1):
        set_cell_text(table.rows[row_idx].cells[0], no, font_size=Pt(9))
        set_cell_text(table.rows[row_idx].cells[1], author,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT, font_size=Pt(9))
        set_cell_text(table.rows[row_idx].cells[2], focus,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT, font_size=Pt(9))
        set_cell_text(table.rows[row_idx].cells[3], method,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT, font_size=Pt(9))
        set_cell_text(table.rows[row_idx].cells[4], finding,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT, font_size=Pt(9))
        set_cell_text(table.rows[row_idx].cells[5], gap,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT, font_size=Pt(9))

    # Set column widths
    col_widths = [Cm(0.8), Cm(2.5), Cm(2.8), Cm(2.5), Cm(3.5), Cm(3.5)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # 2.6 Kerangka Konseptual
    add_section_heading(doc, "2.6 Kerangka Konseptual")

    add_justified_paragraph(
        doc,
        "Kerangka konseptual penelitian ini menggambarkan alur kerja framework "
        "yang diusulkan, mulai dari input pengembang hingga menghasilkan artefak "
        "perangkat lunak yang aman. Berikut adalah alur konseptual framework:"
    )

    # Conceptual framework flow
    flow_items = [
        "Developer",
        "    → Prompt / Requirement / Code Context",
        "        → Privacy Prevention Layer",
        "            → AI Coding Assistant",
        "                → AI-Generated Code",
        "                    → Security and Privacy Detection Layer",
        "                        → Risk Classification and Prioritization",
        "                            → Automatic Mitigation and Code Repair",
        "                                → Verification Layer",
        "                                    → Secure AI-Assisted Software Artifact"
    ]

    for item in flow_items:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.left_indent = Cm(1.0)
        run = para.add_run(item)
        set_run_font(run, font_size=Pt(11))

    add_blank_line(doc)

    add_justified_paragraph(
        doc,
        "Framework ini bekerja dalam dua arah: (1) secara proaktif mencegah risiko "
        "privasi sebelum data dikirim ke AI model melalui Privacy Prevention Layer, "
        "dan (2) secara reaktif mendeteksi dan memitigasi risiko keamanan pada kode "
        "yang dihasilkan AI melalui Security Detection Layer, Risk Classification, "
        "dan Automatic Mitigation. Verification Layer memastikan bahwa seluruh "
        "perbaikan valid dan tidak memperkenalkan risiko baru."
    )

    add_page_break(doc)


def create_bab3(doc):
    """Create BAB III - BAB ISI UTAMA USULAN PENELITIAN."""
    add_chapter_heading(doc, "BAB III. BAB ISI UTAMA USULAN PENELITIAN")

    # 3.1 Desain Penelitian
    add_section_heading(doc, "3.1 Desain Penelitian")

    add_justified_paragraph(
        doc,
        "Penelitian ini menggunakan pendekatan Design Science Research (DSR) [38][39] "
        "yang dikombinasikan dengan metode eksperimental untuk pengembangan dan "
        "evaluasi framework. DSR dipilih karena penelitian ini bertujuan menghasilkan "
        "artefak berupa framework teknis yang memberikan solusi terhadap permasalahan "
        "praktis. Metode eksperimental digunakan untuk mengevaluasi efektivitas "
        "komponen-komponen framework."
    )

    add_justified_paragraph(
        doc,
        "Tahapan penelitian ini terdiri dari:",
        first_line_indent=Cm(0)
    )

    stages = [
        "Identifikasi risiko dan penyusunan risk taxonomy melalui systematic literature "
        "review dan analisis empiris.",
        "Pengembangan Privacy Risk Detection Model untuk mendeteksi risiko privasi pada "
        "prompt dan artefak.",
        "Pengembangan Security Vulnerability Detection Model untuk mendeteksi kerentanan "
        "pada AI-generated code.",
        "Pengembangan Automatic Mitigation and Repair Engine untuk perbaikan otomatis.",
        "Integrasi seluruh komponen ke dalam unified framework.",
        "Evaluasi framework menggunakan benchmark dataset dan metrik yang terukur.",
        "Validasi framework melalui studi kasus pada proyek pengembangan perangkat lunak.",
        "Dokumentasi dan diseminasi hasil penelitian."
    ]

    for i, stage in enumerate(stages, 1):
        add_numbered_item(doc, i, stage)

    # 3.2 Tahap 1
    add_section_heading(doc, "3.2 Tahap 1: Identifikasi Risiko dan Penyusunan Risk Taxonomy")

    add_justified_paragraph(
        doc,
        "Tahap pertama bertujuan mengidentifikasi dan mengklasifikasikan seluruh "
        "risiko privasi dan keamanan yang muncul dalam AI-assisted software engineering. "
        "Metode yang digunakan meliputi systematic literature review [40], analisis "
        "CVE database, review OWASP Top 10 [32], dan analisis empiris pada AI coding "
        "tools. Hasil tahap ini adalah risk taxonomy yang komprehensif."
    )

    add_justified_paragraph(
        doc,
        "Risk taxonomy mencakup kategori berikut:",
        first_line_indent=Cm(0)
    )

    taxonomy_items = [
        "Code Injection Vulnerabilities (SQL injection, XSS, command injection)",
        "Authentication and Authorization Flaws (broken auth, missing access control)",
        "Cryptographic Weaknesses (weak algorithms, hardcoded keys, improper implementation)",
        "Data Exposure Risks (sensitive data in logs, unencrypted transmission)",
        "Configuration Vulnerabilities (insecure defaults, missing security headers)",
        "Dependency Risks (vulnerable libraries, deprecated packages, hallucinated packages)",
        "Prompt Privacy Risks (sensitive data in prompts, credential exposure)",
        "Context Leakage Risks (source code leakage, IP exposure, business logic disclosure)",
        "Output Privacy Risks (training data memorization, cross-user information leakage)",
        "Agentic Risks (prompt injection via code, unauthorized actions, data exfiltration)"
    ]

    for i, item in enumerate(taxonomy_items, 1):
        add_numbered_item(doc, i, item)

    # 3.3 Tahap 2
    add_section_heading(doc, "3.3 Tahap 2: Pengembangan Privacy Risk Detection Model")

    add_justified_paragraph(
        doc,
        "Privacy Risk Detection Model menerima input berupa:",
        first_line_indent=Cm(0)
    )

    input_items = [
        "Prompt text yang akan dikirim ke AI model",
        "Code context dan file yang disertakan dalam request",
        "Repository metadata dan konfigurasi",
        "Environment variables dan configuration files",
        "Conversation history dan chat context",
        "Clipboard content yang di-paste ke prompt"
    ]

    for i, item in enumerate(input_items, 1):
        add_numbered_item(doc, i, item)

    add_justified_paragraph(
        doc,
        "Model ini mendeteksi kategori risiko privasi berikut:",
        first_line_indent=Cm(0)
    )

    detect_items = [
        "API keys, tokens, dan credentials",
        "Database connection strings dan passwords",
        "Personal Identifiable Information (PII)",
        "Internal IP addresses dan infrastructure details",
        "Proprietary business logic dan algorithms",
        "Customer data dan sensitive records",
        "Internal documentation dan trade secrets",
        "Security configurations dan access control rules"
    ]

    for i, item in enumerate(detect_items, 1):
        add_numbered_item(doc, i, item)

    add_justified_paragraph(
        doc,
        "Pendekatan teknis yang digunakan meliputi:",
        first_line_indent=Cm(0)
    )

    approaches = [
        "Regular expression dan pattern matching untuk known credential formats",
        "Named Entity Recognition (NER) untuk identifikasi PII",
        "Entropy analysis untuk mendeteksi high-entropy strings (potential secrets)",
        "Contextual classification menggunakan fine-tuned language model",
        "Rule-based policy engine untuk organizational-specific sensitive data"
    ]

    for i, item in enumerate(approaches, 1):
        add_numbered_item(doc, i, item)

    add_justified_paragraph(
        doc,
        "Output dari model ini berupa risk report yang mencakup lokasi data sensitif, "
        "kategori risiko, severity level, dan recommended action (redact, mask, block, "
        "atau warn)."
    )

    # 3.4 Tahap 3
    add_section_heading(doc, "3.4 Tahap 3: Pengembangan Security Vulnerability Detection Model")

    add_justified_paragraph(
        doc,
        "Security Vulnerability Detection Model menganalisis kode yang dihasilkan AI "
        "untuk mendeteksi kerentanan keamanan. Model ini mencakup deteksi untuk "
        "10 kategori kerentanan utama:",
        first_line_indent=Cm(0)
    )

    vuln_types = [
        "SQL Injection dan NoSQL Injection",
        "Cross-Site Scripting (XSS) - Reflected, Stored, DOM-based",
        "Command Injection dan Code Injection",
        "Path Traversal dan File Inclusion",
        "Insecure Authentication (weak hashing, missing MFA, session fixation)",
        "Broken Access Control (IDOR, privilege escalation, missing authorization)",
        "Cryptographic Failures (weak algorithms, ECB mode, hardcoded keys)",
        "Insecure Deserialization dan Type Confusion",
        "Server-Side Request Forgery (SSRF)",
        "Insecure Dependencies dan Known Vulnerable Components"
    ]

    for i, item in enumerate(vuln_types, 1):
        add_numbered_item(doc, i, item)

    add_justified_paragraph(
        doc,
        "Pendekatan deteksi yang digunakan meliputi:",
        first_line_indent=Cm(0)
    )

    detection_approaches = [
        "Abstract Syntax Tree (AST) analysis untuk structural vulnerability patterns",
        "Data flow analysis dan taint tracking untuk injection vulnerabilities",
        "Transformer-based model (fine-tuned CodeBERT/GraphCodeBERT) untuk semantic "
        "vulnerability detection",
        "Pattern matching dan rule engine untuk known vulnerability signatures",
        "Control flow analysis untuk authentication dan authorization flaws",
        "Dependency analysis dan CVE database lookup untuk third-party risks"
    ]

    for i, item in enumerate(detection_approaches, 1):
        add_numbered_item(doc, i, item)

    add_justified_paragraph(
        doc,
        "Output dari model ini berupa vulnerability report yang mencakup jenis "
        "kerentanan, lokasi (file, line, function), severity (critical, high, medium, "
        "low), confidence score, CWE mapping, dan suggested fix."
    )

    # 3.5 Tahap 4
    add_section_heading(doc, "3.5 Tahap 4: Pengembangan Automatic Mitigation and Repair Engine")

    add_justified_paragraph(
        doc,
        "Automatic Mitigation and Repair Engine bertanggung jawab untuk memperbaiki "
        "risiko yang terdeteksi secara otomatis. Engine ini mencakup mekanisme berikut:",
        first_line_indent=Cm(0)
    )

    mitigation_mechanisms = [
        "Sensitive data redaction - mengganti data sensitif dalam prompt dengan placeholder",
        "Credential masking - menyembunyikan API keys, tokens, dan passwords sebelum pengiriman",
        "Code sanitization - membersihkan input validation pada AI-generated code",
        "Parameterized query transformation - mengubah string concatenation menjadi parameterized queries",
        "Secure coding pattern replacement - mengganti pola kode tidak aman dengan secure alternatives",
        "Authentication hardening - menambahkan proper hashing, salting, dan session management",
        "Access control injection - menambahkan authorization checks yang hilang",
        "Cryptographic upgrade - mengganti algoritma lemah dengan yang direkomendasikan",
        "Dependency update - memperbarui vulnerable dependencies ke versi aman",
        "Configuration hardening - memperbaiki security configurations dan adding security headers"
    ]

    for i, item in enumerate(mitigation_mechanisms, 1):
        add_numbered_item(doc, i, item)

    add_justified_paragraph(
        doc,
        "Contoh alur mitigasi: AI menghasilkan kode dengan SQL injection vulnerability → "
        "Detection model mengidentifikasi string concatenation pada SQL query → "
        "Repair engine melakukan transformasi menjadi parameterized query → "
        "Verification layer memvalidasi bahwa perbaikan correct dan tidak breaking."
    )

    add_justified_paragraph(
        doc,
        "Output dari engine ini berupa repaired code, diff report (before/after), "
        "repair explanation, dan confidence score untuk setiap perbaikan yang dilakukan."
    )

    # 3.6 Tahap 5
    add_section_heading(doc, "3.6 Tahap 5: Integrasi Framework")

    add_justified_paragraph(
        doc,
        "Seluruh komponen diintegrasikan ke dalam arsitektur unified framework "
        "dengan alur sebagai berikut:",
        first_line_indent=Cm(0)
    )

    # Architecture flow
    arch_flow = [
        "Input Layer (Prompt/Code Context)",
        "    → Privacy Prevention Module (scan, detect, redact/mask)",
        "        → Sanitized Input → AI Coding Assistant",
        "            → AI Output (Generated Code)",
        "                → Security Detection Module (AST, dataflow, ML-based)",
        "                    → Privacy Detection Module (output scanning)",
        "                        → Risk Classifier (severity, priority, category)",
        "                            → Mitigation Engine (auto-repair, transform)",
        "                                → Verification Module (validate repairs)",
        "                                    → Secure Output (to repository/pipeline)"
    ]

    for item in arch_flow:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.left_indent = Cm(1.0)
        run = para.add_run(item)
        set_run_font(run, font_size=Pt(11))

    add_blank_line(doc)

    add_justified_paragraph(
        doc,
        "Framework diimplementasikan dengan opsi deployment berikut:",
        first_line_indent=Cm(0)
    )

    impl_options = [
        "IDE Extension (VS Code, JetBrains) - integrasi langsung di editor pengembang",
        "CLI Tool - command-line interface untuk integrasi dengan CI/CD pipeline",
        "API Service - microservice yang dapat diintegrasikan dengan berbagai tools",
        "Git Hook - pre-commit dan pre-push hook untuk validasi otomatis",
        "Proxy Layer - intercepting proxy antara developer tools dan AI API endpoint"
    ]

    for i, item in enumerate(impl_options, 1):
        add_numbered_item(doc, i, item)

    # 3.7 Tahap 6
    add_section_heading(doc, "3.7 Tahap 6: Evaluasi Framework")

    add_justified_paragraph(
        doc,
        "Evaluasi framework dilakukan menggunakan dataset dan benchmark berikut:",
        first_line_indent=Cm(0)
    )

    datasets = [
        "CWE Top 25 - dataset kerentanan keamanan paling berbahaya",
        "OWASP Benchmark - test suite untuk vulnerability detection tools",
        "CodeQL Dataset - real-world vulnerabilities dari GitHub repositories",
        "SecEval - benchmark untuk evaluasi keamanan AI-generated code",
        "Custom Prompt Dataset - dataset prompt dengan data sensitif (synthetic)",
        "CVE Database - known vulnerabilities untuk dependency analysis evaluation",
        "HumanEval-Security - extended HumanEval dengan security annotations"
    ]

    for i, item in enumerate(datasets, 1):
        add_numbered_item(doc, i, item)

    add_justified_paragraph(
        doc,
        "Metrik evaluasi untuk Privacy Detection:",
        first_line_indent=Cm(0)
    )

    privacy_metrics = [
        "Precision - proporsi true positive dari seluruh positive predictions",
        "Recall - proporsi true positive dari seluruh actual sensitive data",
        "F1-Score - harmonic mean dari precision dan recall",
        "False Positive Rate - proporsi false alarms",
        "Detection Latency - waktu yang diperlukan untuk mendeteksi risiko",
        "Coverage - proporsi kategori risiko yang dapat dideteksi"
    ]

    for i, item in enumerate(privacy_metrics, 1):
        add_numbered_item(doc, i, item)

    add_justified_paragraph(
        doc,
        "Metrik evaluasi untuk Security Detection:",
        first_line_indent=Cm(0)
    )

    security_metrics = [
        "Vulnerability Detection Rate (True Positive Rate)",
        "False Positive Rate",
        "CWE Coverage - jumlah CWE types yang dapat dideteksi",
        "Severity Accuracy - ketepatan klasifikasi severity",
        "Detection Time - waktu analisis per file/function",
        "Comparison with existing tools (CodeQL, Semgrep, Snyk)"
    ]

    for i, item in enumerate(security_metrics, 1):
        add_numbered_item(doc, i, item)

    add_justified_paragraph(
        doc,
        "Metrik evaluasi untuk Mitigation:",
        first_line_indent=Cm(0)
    )

    mitigation_metrics = [
        "Repair Success Rate - proporsi vulnerability yang berhasil diperbaiki",
        "Functional Correctness - kode hasil repair tetap berfungsi dengan benar",
        "Regression Rate - proporsi repair yang memperkenalkan bug baru",
        "Semantic Preservation - kode repair mempertahankan intended behavior",
        "Repair Time - waktu yang diperlukan untuk menghasilkan perbaikan",
        "Developer Acceptance Rate - proporsi repair yang diterima developer"
    ]

    for i, item in enumerate(mitigation_metrics, 1):
        add_numbered_item(doc, i, item)

    # 3.8 Jadwal Penelitian
    add_section_heading(doc, "3.8 Jadwal Penelitian")

    add_justified_paragraph(
        doc,
        "Jadwal penelitian disusun untuk durasi 3 tahun (6 semester) sebagai berikut:",
        first_line_indent=Cm(0)
    )

    # Create schedule table
    schedule_table = doc.add_table(rows=11, cols=4)
    schedule_table.style = 'Table Grid'

    # Header
    schedule_headers = ["Kegiatan", "Tahun 1", "Tahun 2", "Tahun 3"]
    for i, header in enumerate(schedule_headers):
        set_cell_text(schedule_table.rows[0].cells[i], header, bold=True, font_size=Pt(10))
        set_cell_shading(schedule_table.rows[0].cells[i])

    schedule_data = [
        ("Literature Review dan Risk Identification", "✓", "", ""),
        ("Risk Taxonomy Development", "✓", "", ""),
        ("Privacy Detection Model Development", "✓", "✓", ""),
        ("Security Detection Model Development", "", "✓", ""),
        ("Mitigation Engine Development", "", "✓", ""),
        ("Framework Integration", "", "✓", "✓"),
        ("Evaluation dan Benchmarking", "", "", "✓"),
        ("Case Study Validation", "", "", "✓"),
        ("Paper Writing dan Publication", "✓", "✓", "✓"),
        ("Dissertation Writing", "", "✓", "✓")
    ]

    for row_idx, (activity, y1, y2, y3) in enumerate(schedule_data, 1):
        set_cell_text(schedule_table.rows[row_idx].cells[0], activity,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT, font_size=Pt(10))
        set_cell_text(schedule_table.rows[row_idx].cells[1], y1, font_size=Pt(10))
        set_cell_text(schedule_table.rows[row_idx].cells[2], y2, font_size=Pt(10))
        set_cell_text(schedule_table.rows[row_idx].cells[3], y3, font_size=Pt(10))

    # Set column widths for schedule table
    schedule_col_widths = [Cm(7.0), Cm(3.0), Cm(3.0), Cm(3.0)]
    for row in schedule_table.rows:
        for idx, width in enumerate(schedule_col_widths):
            row.cells[idx].width = width

    add_page_break(doc)


def create_bab4(doc):
    """Create BAB IV - KEBARUAN DAN ORISINALITAS."""
    add_chapter_heading(doc, "BAB IV. KEBARUAN DAN ORISINALITAS")

    # 4.1 Kebaruan Penelitian
    add_section_heading(doc, "4.1 Kebaruan Penelitian")

    add_justified_paragraph(
        doc,
        "Penelitian ini memiliki beberapa kebaruan (novelty) dibandingkan dengan "
        "penelitian-penelitian sebelumnya:"
    )

    novelties = [
        ("Integrated Privacy-Security Framework.",
         "Berbeda dengan penelitian sebelumnya yang membahas keamanan atau privasi "
         "secara terpisah, penelitian ini mengintegrasikan prevention, detection, dan "
         "mitigation untuk kedua aspek dalam satu framework unified yang bekerja "
         "end-to-end."),
        ("Autonomous Risk Prevention.",
         "Framework ini mampu bekerja secara otonom tanpa memerlukan intervensi manual "
         "yang signifikan, mulai dari deteksi hingga perbaikan. Pendekatan ini berbeda "
         "dari tools yang hanya memberikan warning tanpa automated repair."),
        ("AI-ASE Specific Design.",
         "Framework ini dirancang khusus untuk konteks AI-assisted software engineering "
         "dengan mempertimbangkan karakteristik unik seperti prompt-based interaction, "
         "non-deterministic output, dan context window limitations yang tidak ada pada "
         "traditional development."),
        ("Bidirectional Protection.",
         "Framework ini melindungi dua arah: (1) mencegah kebocoran data sensitif dari "
         "developer ke AI (privacy prevention), dan (2) mencegah kerentanan keamanan "
         "dari AI ke codebase (security detection dan mitigation)."),
        ("Verification-Integrated Repair.",
         "Berbeda dari automated program repair tradisional, mekanisme repair dalam "
         "framework ini terintegrasi dengan verification layer yang memastikan "
         "perbaikan tidak memperkenalkan risiko baru atau merusak fungsionalitas "
         "yang ada.")
    ]

    for i, (title, desc) in enumerate(novelties, 1):
        add_numbered_item(doc, i, desc, bold_prefix=title)

    # 4.2 Orisinalitas Penelitian
    add_section_heading(doc, "4.2 Orisinalitas Penelitian")

    add_justified_paragraph(
        doc,
        "Orisinalitas penelitian ini terletak pada pengembangan framework yang "
        "sepenuhnya baru dan belum pernah ada sebelumnya. Meskipun terdapat "
        "penelitian tentang keamanan AI-generated code (Pearce et al., 2022; "
        "Perry et al., 2023) dan privasi pada code models (Niu et al., 2024), "
        "belum ada penelitian yang mengintegrasikan keduanya dalam satu framework "
        "otonom yang mencakup prevention, detection, dan mitigation secara "
        "komprehensif."
    )

    add_justified_paragraph(
        doc,
        "Penelitian ini juga orisinal dalam mengusulkan pendekatan bidirectional "
        "protection yang mempertimbangkan risiko pada kedua arah interaksi "
        "(developer → AI dan AI → codebase). Kontribusi orisinal lainnya adalah "
        "pengembangan risk taxonomy yang spesifik untuk AI-assisted software "
        "engineering, yang belum tersedia dalam literatur dan standar keamanan "
        "yang ada saat ini."
    )

    add_page_break(doc)


def create_bab5_daftar_pustaka(doc):
    """Create BAB V - DAFTAR PUSTAKA."""
    add_chapter_heading(doc, "BAB V. DAFTAR PUSTAKA")

    references = [
        '[1] M. Chen et al., "Evaluating Large Language Models Trained on Code," '
        'arXiv:2107.03374, 2021.',

        '[2] H. Pearce et al., "Asleep at the Keyboard? Assessing the Security of '
        'GitHub Copilot Code Contributions," in Proc. IEEE S&P, 2022, pp. 754-768.',

        '[3] N. Perry et al., "Do Users Write More Insecure Code with AI Assistants?," '
        'in Proc. ACM CCS, 2023, pp. 2785-2799.',

        '[4] G. Sandoval et al., "Lost at C: A User Study on the Security Implications '
        'of Large Language Model Code Assistants," in Proc. USENIX Security, 2023.',

        '[5] S. Peng et al., "The Impact of AI on Developer Productivity: Evidence from '
        'GitHub Copilot," arXiv:2302.06590, 2023.',

        '[6] A. Ziegler et al., "Productivity Assessment of Neural Code Completion," '
        'in Proc. MAPS, 2022.',

        '[7] R. Khoury et al., "How Secure is Code Generated by ChatGPT?," Systems, '
        'vol. 11, no. 7, p. 364, 2023.',

        '[8] Y. Liu et al., "No Need to Lift a Finger Anymore? Assessing the Quality '
        'of Code Generation by ChatGPT," IEEE Trans. Softw. Eng., vol. 50, no. 9, 2024.',

        '[9] J. He and M. Vechev, "Large Language Models for Code: Security Hardening '
        'and Adversarial Testing," in Proc. ACM CCS, 2023.',

        '[10] B. Niu et al., "CodexLeaks: Privacy Leaks from Code Generation Language '
        'Models in GitHub Copilot," in Proc. USENIX Security, 2024.',

        '[11] R. Yang et al., "Unveiling Memorization in Code Models," in Proc. ICSE, 2024.',

        '[12] S. Jiang et al., "A Survey on Large Language Models for Software Engineering," '
        'ACM Trans. Softw. Eng. Methodol., 2024.',

        '[13] X. Hou et al., "Large Language Models for Software Engineering: A Systematic '
        'Literature Review," ACM Trans. Softw. Eng. Methodol., 2024.',

        '[14] Z. Li et al., "VulDeePecker: A Deep Learning-Based System for Vulnerability '
        'Detection," in Proc. NDSS, 2018.',

        '[15] M. Fu and C. Tantithamthavorn, "LineVul: A Transformer-based Line-Level '
        'Vulnerability Prediction," in Proc. MSR, 2022.',

        '[16] NIST, "Artificial Intelligence Risk Management Framework (AI RMF 1.0)," '
        'NIST, 2023.',

        '[17] ISO/IEC, "ISO/IEC 27001:2022 Information Security Management Systems," 2022.',

        '[18] Z. Feng et al., "CodeBERT: A Pre-Trained Model for Programming and Natural '
        'Languages," in Proc. EMNLP, 2020.',

        '[19] D. Guo et al., "GraphCodeBERT: Pre-training Code Representations with Data '
        'Flow," in Proc. ICLR, 2021.',

        '[20] J. Austin et al., "Program Synthesis with Large Language Models," '
        'arXiv:2108.07732, 2021.',

        '[21] Y. Zhu et al., "PromptBench: Towards Evaluating the Robustness of Large '
        'Language Models on Adversarial Prompts," arXiv:2306.04528, 2023.',

        '[22] R. Jesse et al., "Large Language Models and Simple, Stupid Bugs," '
        'in Proc. MSR, 2023.',

        '[23] Z. Fan et al., "Large Language Models for Software Engineering: Survey '
        'and Open Problems," arXiv:2310.03533, 2023.',

        '[24] C. Le Goues et al., "Automated Program Repair," Commun. ACM, vol. 62, '
        'no. 12, pp. 56-65, 2019.',

        '[25] B. Steenhoek et al., "An Empirical Study of Deep Learning Models for '
        'Vulnerability Detection," in Proc. ICSE, 2024.',

        '[26] M. Croft et al., "Data Quality for Software Vulnerability Datasets," '
        'in Proc. ICSE, 2023.',

        '[27] Y. Zhou et al., "Large Language Model for Vulnerability Detection: '
        'Emerging Results and Future Directions," arXiv:2403.18624, 2024.',

        '[28] D. Ding et al., "Vulnerability Detection with Fine-Grained '
        'Interpretations," in Proc. ESEC/FSE, 2021.',

        '[29] J. Zhang et al., "A Survey on Automated Program Repair Techniques," '
        'arXiv:2303.18184, 2023.',

        '[30] S. Muthukrishnan et al., "Fixing Security Vulnerabilities with AI in '
        'OSS-Fuzz," arXiv:2411.03346, 2024.',

        '[31] G. Ryan et al., "Code Red: AI-Assisted Secure Coding," in Proc. IEEE '
        'S&P Workshop, 2024.',

        '[32] OWASP, "OWASP Top 10:2021," OWASP Foundation, 2021.',

        '[33] Y. Li et al., "Competition-Level Code Generation with AlphaCode," '
        'Science, vol. 378, pp. 1092-1097, 2022.',

        '[34] E. Nijkamp et al., "CodeGen: An Open Large Language Model for Code with '
        'Multi-Turn Program Synthesis," in Proc. ICLR, 2023.',

        '[35] W. U. Ahmad et al., "Unified Pre-training for Program Understanding and '
        'Generation," in Proc. NAACL, 2021.',

        '[36] Z. Yang et al., "A Survey on Deep Learning for Software Engineering," '
        'ACM Comput. Surv., vol. 54, no. 10, 2022.',

        '[37] T. Ahmed and P. Devanbu, "Few-shot Training LLMs for Project-specific '
        'Code-summarization," in Proc. ASE, 2022.',

        '[38] K. Hevner et al., "Design Science in Information Systems Research," '
        'MIS Quarterly, vol. 28, no. 1, pp. 75-105, 2004.',

        '[39] K. Peffers et al., "A Design Science Research Methodology for Information '
        'Systems Research," J. Manage. Inf. Syst., vol. 24, no. 3, pp. 45-77, 2007.',

        '[40] V. Braun and V. Clarke, "Using Thematic Analysis in Psychology," '
        'Qualitative Research in Psychology, vol. 3, no. 2, pp. 77-101, 2006.',
    ]

    for ref in references:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.left_indent = Cm(1.27)
        para.paragraph_format.first_line_indent = Cm(-1.27)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(ref)
        set_run_font(run, font_size=Pt(11))


# =============================================================================
# MAIN EXECUTION
# =============================================================================

def main():
    """Main function to generate the PhD proposal document."""
    # Create output directory if not exists
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Create document
    print("Creating document...")
    doc = create_document()

    # Build document sections
    print("Creating cover page...")
    create_cover_page(doc)

    print("Creating BAB I - Pendahuluan...")
    create_bab1(doc)

    print("Creating BAB II - Tinjauan Pustaka...")
    create_bab2(doc)

    print("Creating BAB III - Isi Utama Usulan Penelitian...")
    create_bab3(doc)

    print("Creating BAB IV - Kebaruan dan Orisinalitas...")
    create_bab4(doc)

    print("Creating BAB V - Daftar Pustaka...")
    create_bab5_daftar_pustaka(doc)

    # Save document
    doc.save(OUTPUT_FILE)
    print(f"\nDocument saved successfully: {OUTPUT_FILE}")
    print(f"File size: {os.path.getsize(OUTPUT_FILE) / 1024:.1f} KB")


if __name__ == "__main__":
    main()
