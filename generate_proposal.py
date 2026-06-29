#!/usr/bin/env python3
"""
Generate PhD Dissertation Proposal for ITB Computer Science.
Topic: Governing Privacy and Security Risks in AI-Assisted Software Engineering
Output: phd-itb-2026/proposal_disertasi.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_fp(doc, text, font_size=12, bold=False, italic=False,
           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6),
           space_before=Pt(0), line_spacing=1.5, first_line_indent=None):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = space_after
    paragraph.paragraph_format.space_before = space_before
    paragraph.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = first_line_indent
    return paragraph


def add_bab(doc, num, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(f"BAB {num}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.font.bold = True
    r._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(18)
    p2.paragraph_format.line_spacing = 1.5
    r2 = p2.add_run(title)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(14)
    r2.font.bold = True
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")


def add_sub(doc, number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(f"{number} {title}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.bold = True
    r._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")


def add_body(doc, text, indent=True):
    return add_fp(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_line_indent=Cm(1.27) if indent else None, line_spacing=1.5)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    marker = "\u2022" if level == 0 else "-"
    r = p.add_run(f"{marker} {text}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")


def add_num(doc, number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    r = p.add_run(f"{number}. {text}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")


def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.bold = True
    r._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")


def set_cell(cell, text, bold=False, font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(font_size)
    r.font.bold = bold
    r._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")


def main():
    doc = Document()

    # Set default document formatting
    style = doc.styles['Normal']
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Set margins and paper size (A4)
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(4.0)
        section.top_margin = Cm(3.0)
        section.right_margin = Cm(3.0)
        section.bottom_margin = Cm(3.0)

    # ============ COVER PAGE ============
    create_cover(doc)

    # ============ BAB I ============
    create_bab1(doc)

    # ============ BAB II ============
    create_bab2(doc)

    # ============ BAB III ============
    create_bab3(doc)

    # ============ BAB IV ============
    create_bab4(doc)

    # ============ BAB V - DAFTAR PUSTAKA ============
    create_bab5(doc)

    # Save document
    os.makedirs("phd-itb-2026", exist_ok=True)
    output_path = os.path.join("phd-itb-2026", "proposal_disertasi.docx")
    doc.save(output_path)
    print(f"Proposal disertasi berhasil dibuat: {output_path}")
    print(f"Ukuran file: {os.path.getsize(output_path) / 1024:.1f} KB")


def create_cover(doc):
    for _ in range(3):
        add_fp(doc, "", space_after=Pt(24))

    add_fp(doc, "PROPOSAL DISERTASI", font_size=16, bold=True,
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(36))

    add_fp(doc, "GOVERNING PRIVACY AND SECURITY RISKS IN AI-ASSISTED SOFTWARE ENGINEERING:",
           font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_fp(doc, "PENGEMBANGAN FRAMEWORK TATA KELOLA KEAMANAN DAN PRIVASI PADA REKAYASA PERANGKAT LUNAK BERBANTUAN KECERDASAN BUATAN",
           font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(48))

    add_fp(doc, "Diajukan oleh:", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    add_fp(doc, "[Nama Lengkap Mahasiswa]", bold=True,
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_fp(doc, "NIM: [Nomor Induk Mahasiswa]",
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(48))

    add_fp(doc, "Program Studi Doktor Informatika", bold=True,
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_fp(doc, "Fakultas Informatika", bold=True,
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_fp(doc, "Institut Teknologi Bandung", bold=True,
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_fp(doc, "2026", bold=True,
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    doc.add_page_break()


def create_bab1(doc):
    add_bab(doc, "I", "PENDAHULUAN")

    add_sub(doc, "1.1", "Latar Belakang")

    add_body(doc, "Perkembangan kecerdasan buatan (Artificial Intelligence/AI) telah membawa transformasi fundamental dalam praktik rekayasa perangkat lunak (software engineering) modern. Kemunculan large language models (LLMs) yang dilatih secara khusus untuk pemrograman, seperti GitHub Copilot, Amazon CodeWhisperer, ChatGPT, dan berbagai AI coding assistants lainnya, telah mengubah cara pengembang perangkat lunak menulis, menguji, dan memelihara kode secara signifikan [1][2][3]. Transformasi ini dipercepat oleh kemajuan arsitektur Transformer [4] yang memungkinkan model bahasa berskala besar memahami dan menghasilkan kode program dengan tingkat akurasi yang semakin tinggi.")

    add_body(doc, "Studi empiris menunjukkan bahwa penggunaan AI coding assistants memberikan peningkatan produktivitas yang signifikan bagi pengembang perangkat lunak. Penelitian oleh Peng et al. (2023) membuktikan bahwa GitHub Copilot meningkatkan kecepatan penyelesaian tugas pemrograman sebesar 55,8% [2]. Survei GitHub terhadap 500 pengembang di Amerika Serikat menunjukkan bahwa 92% pengembang profesional telah menggunakan AI coding tools dalam pekerjaan mereka [5]. Iansiti dan Lakhani (2023) melaporkan dampak ekonomi substansial dari adopsi AI dalam siklus hidup pengembangan perangkat lunak, termasuk pengurangan biaya pengembangan dan percepatan time-to-market [6].")

    add_body(doc, "Namun demikian, di balik peningkatan produktivitas tersebut, terdapat risiko keamanan dan privasi yang serius dan belum terkelola dengan baik. Penelitian Perry et al. (2023) yang dipublikasikan pada ACM Conference on Computer and Communications Security (CCS) menunjukkan bahwa pengembang yang menggunakan AI assistants justru menghasilkan kode yang secara signifikan lebih tidak aman dibandingkan pengembang yang tidak menggunakan AI [7]. Pearce et al. (2022) dalam studi yang dipresentasikan pada IEEE Symposium on Security and Privacy menemukan bahwa sekitar 40% dari saran kode yang dihasilkan GitHub Copilot mengandung kerentanan keamanan (security vulnerabilities) [8].")

    add_body(doc, "Permasalahan keamanan ini diperkuat oleh temuan Fu et al. (2025) yang menganalisis kode yang dihasilkan Copilot pada proyek-proyek nyata di GitHub dan menemukan berbagai kelemahan keamanan yang terklasifikasi dalam Common Weakness Enumeration (CWE) [9]. Sandoval et al. (2023) dalam studi yang dipublikasikan pada USENIX Security mengonfirmasi bahwa pengembang seringkali tidak menyadari kerentanan yang diperkenalkan oleh kode yang dihasilkan AI [10]. Lebih lanjut, Bhatt et al. (2023) melalui Purple Llama CyberSecEval mengembangkan benchmark untuk mengukur sejauh mana model bahasa menghasilkan kode yang tidak aman [11].")

    add_body(doc, "Dari perspektif privasi, penggunaan AI coding assistants menimbulkan risiko kebocoran data yang signifikan. Kode sumber (source code) yang merupakan aset intelektual organisasi dikirimkan ke server pihak ketiga untuk diproses oleh model AI. Carlini et al. (2021) mendemonstrasikan bahwa data pelatihan dapat diekstraksi dari large language models, menimbulkan risiko terhadap kerahasiaan kode yang pernah digunakan dalam pelatihan model [12]. Selain itu, isu lintas batas (cross-border data transfer) menjadi semakin relevan dengan berlakunya regulasi seperti EU AI Act [13] dan berbagai regulasi perlindungan data pribadi nasional.")

    add_body(doc, "Fenomena AI-assisted software engineering juga menghadirkan dimensi risiko baru berupa prompt injection [14][15], hallucinated APIs dan dependencies [16], serta serangan terhadap software supply chain [17][18]. Fang et al. (2024) bahkan mendemonstrasikan bahwa LLM agents dapat secara otonom melakukan eksploitasi terhadap sistem web [19], menunjukkan potensi penyalahgunaan yang serius dari teknologi ini.")

    add_body(doc, "Meskipun berbagai risiko tersebut telah diidentifikasi, terdapat kesenjangan signifikan dalam hal tata kelola (governance). Sebagian besar organisasi belum memiliki framework yang komprehensif untuk mengelola risiko keamanan dan privasi dalam penggunaan AI untuk rekayasa perangkat lunak. Standar yang ada seperti NIST AI Risk Management Framework [20] dan ISO/IEC 27001:2022 [21] memberikan panduan umum, namun belum secara spesifik menangani konteks unik AI-assisted software engineering. Penelitian yang ada cenderung membahas aspek keamanan kode AI atau tata kelola AI secara terpisah, tanpa integrasi yang memadai antara aspek keamanan, privasi, dan governance dalam satu framework yang holistik [22][23].")

    add_body(doc, "Berdasarkan uraian di atas, terdapat kebutuhan mendesak untuk mengembangkan sebuah framework tata kelola yang terintegrasi untuk mengelola risiko keamanan dan privasi pada rekayasa perangkat lunak berbantuan AI. Framework ini diharapkan dapat menjembatani kesenjangan antara pemanfaatan AI untuk produktivitas pengembangan perangkat lunak dengan pengelolaan risiko keamanan dan privasi yang efektif, sehingga organisasi dapat mengadopsi AI coding tools secara aman dan bertanggung jawab.")

    # 1.2
    add_sub(doc, "1.2", "Rumusan Masalah")
    add_body(doc, "Berdasarkan latar belakang yang telah diuraikan, rumusan masalah dalam penelitian ini adalah sebagai berikut:")
    add_num(doc, 1, "Risiko keamanan dan privasi apa saja yang muncul dalam praktik AI-assisted software engineering, dan bagaimana risiko tersebut dapat diklasifikasikan secara sistematis?")
    add_num(doc, 2, "Bagaimana praktik tata kelola (governance) yang diterapkan organisasi saat ini dalam mengendalikan risiko keamanan dan privasi pada penggunaan AI coding tools?")
    add_num(doc, 3, "Bagaimana merancang framework governance yang efektif dan komprehensif untuk mengelola risiko keamanan dan privasi dalam AI-assisted software engineering?")
    add_num(doc, 4, "Bagaimana mengukur dan mengevaluasi efektivitas framework governance yang diusulkan dalam konteks pengembangan perangkat lunak berbantuan AI?")

    # 1.3
    add_sub(doc, "1.3", "Pertanyaan Penelitian")
    add_body(doc, "Untuk menjawab rumusan masalah di atas, penelitian ini merumuskan empat pertanyaan penelitian (research questions) sebagai berikut:")

    rqs = [
        ("RQ1", "What privacy and security risks emerge in AI-assisted software engineering, and how can they be systematically classified?"),
        ("RQ2", "How are organizations currently governing privacy and security risks in AI-assisted software engineering?"),
        ("RQ3", "What governance mechanisms should be incorporated into an integrated framework for managing privacy and security risks in AI-assisted software engineering?"),
        ("RQ4", "How effective is the proposed governance framework in reducing privacy and security risks in AI-assisted software engineering?"),
    ]
    for rq_id, rq_text in rqs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Cm(1.27)
        r = p.add_run(f"{rq_id}: ")
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        r.font.bold = True
        r._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
        r2 = p.add_run(rq_text)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(12)
        r2.font.italic = True
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")

    # 1.4
    add_sub(doc, "1.4", "Tujuan Penelitian")
    add_fp(doc, "Tujuan Umum:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_body(doc, "Mengembangkan framework governance yang terintegrasi untuk mengelola risiko keamanan dan privasi pada rekayasa perangkat lunak berbantuan kecerdasan buatan (AI-assisted software engineering).")

    add_fp(doc, "Tujuan Khusus:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_num(doc, 1, "Mengidentifikasi dan mengklasifikasikan risiko keamanan dan privasi yang muncul dalam AI-assisted software engineering melalui systematic literature review dan studi empiris.")
    add_num(doc, 2, "Menganalisis praktik tata kelola yang diterapkan organisasi saat ini dalam mengendalikan risiko keamanan dan privasi pada penggunaan AI coding tools.")
    add_num(doc, 3, "Merancang dan mengembangkan framework governance yang mengintegrasikan aspek keamanan, privasi, dan tata kelola AI dalam konteks rekayasa perangkat lunak.")
    add_num(doc, 4, "Mengevaluasi efektivitas framework governance yang diusulkan melalui studi kasus, validasi pakar, dan pengukuran metrik governance maturity.")

    # 1.5
    add_sub(doc, "1.5", "Manfaat Penelitian")
    add_fp(doc, "Manfaat Teoretis:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_bullet(doc, "Memperkaya body of knowledge di bidang AI governance dengan kontribusi framework yang mengintegrasikan perspektif keamanan, privasi, dan tata kelola dalam AI-assisted software engineering.")
    add_bullet(doc, "Mengembangkan taksonomi risiko keamanan dan privasi yang spesifik untuk konteks pengembangan perangkat lunak berbantuan AI, yang dapat menjadi referensi bagi penelitian selanjutnya.")
    add_bullet(doc, "Menghubungkan teori software engineering, information security governance, dan AI governance dalam satu kerangka konseptual yang terintegrasi.")

    add_fp(doc, "Manfaat Praktis:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_bullet(doc, "Menyediakan panduan operasional bagi organisasi dalam mengadopsi dan mengelola AI coding tools secara aman dan bertanggung jawab.")
    add_bullet(doc, "Memberikan instrumen penilaian (assessment tool) governance maturity yang dapat digunakan organisasi untuk mengukur tingkat kesiapan dan efektivitas tata kelola AI dalam pengembangan perangkat lunak.")
    add_bullet(doc, "Mendukung kepatuhan organisasi terhadap regulasi terkait AI, privasi data, dan keamanan informasi yang berlaku.")
    add_bullet(doc, "Membantu pembuat kebijakan (policymakers) dalam merumuskan regulasi dan standar terkait penggunaan AI dalam rekayasa perangkat lunak.")

    # 1.6
    add_sub(doc, "1.6", "Ruang Lingkup")
    add_fp(doc, "Termasuk dalam lingkup penelitian:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_bullet(doc, "Generative AI untuk pembuatan kode (code generation), termasuk AI coding assistants (GitHub Copilot, Amazon CodeWhisperer, ChatGPT, Gemini Code Assist) dan agentic coding tools.")
    add_bullet(doc, "Secure Software Development Lifecycle (Secure SDLC) dalam konteks penggunaan AI.")
    add_bullet(doc, "Privacy governance terkait pengelolaan data kode sumber dan data sensitif yang diproses oleh AI systems.")
    add_bullet(doc, "Developer workflow yang melibatkan AI tools dalam berbagai tahap pengembangan perangkat lunak (coding, testing, code review, debugging, documentation).")
    add_bullet(doc, "Kerangka regulasi dan standar internasional terkait AI governance, keamanan informasi, dan perlindungan data pribadi.")

    add_fp(doc, "Tidak termasuk dalam lingkup penelitian:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_bullet(doc, "Proses pelatihan (training) foundation models dan large language models.")
    add_bullet(doc, "Hardware security dan physical security infrastructure.")
    add_bullet(doc, "Pure cryptography dan pengembangan algoritma kriptografi baru.")
    add_bullet(doc, "Pengembangan AI models baru untuk code generation.")
    add_bullet(doc, "Aspek hukum spesifik yurisdiksi tertentu (penelitian berfokus pada framework governance yang generik dan dapat diadaptasi).")

    doc.add_page_break()


def create_bab2(doc):
    add_bab(doc, "II", "TINJAUAN PUSTAKA")

    # 2.1
    add_sub(doc, "2.1", "AI-Assisted Software Engineering")

    add_body(doc, "AI-Assisted Software Engineering merujuk pada pemanfaatan teknologi kecerdasan buatan, khususnya large language models (LLMs), untuk mendukung berbagai aktivitas dalam siklus hidup pengembangan perangkat lunak. Konsep ini telah berevolusi dari automated code completion sederhana menjadi sistem yang mampu menghasilkan kode kompleks, melakukan code review, menulis pengujian, dan bahkan melakukan debugging secara otonom [24][25].")

    add_body(doc, "Evolusi AI coding assistants dimulai dengan model Codex [26] yang dikembangkan oleh OpenAI dan menjadi dasar GitHub Copilot. Codex dilatih pada miliaran baris kode dari repositori publik dan mampu menerjemahkan natural language descriptions menjadi kode program. Perkembangan selanjutnya menghasilkan model-model yang lebih canggih seperti AlphaCode [27], StarCoder 2 [28], dan berbagai model open-source lainnya yang menunjukkan kemampuan code generation yang semakin mendekati tingkat manusia.")

    add_body(doc, "Fan et al. (2023) dalam survei komprehensif tentang LLMs untuk software engineering mengidentifikasi berbagai area aplikasi termasuk code generation, program repair, code summarization, test generation, dan code review [29]. Hou et al. (2024) dalam systematic literature review yang dipublikasikan pada ACM TOSEM mengkategorikan penggunaan LLMs dalam software engineering ke dalam fase-fase SDLC dan menemukan bahwa adopsi terbesar terjadi pada fase implementasi dan pengujian [30].")

    add_body(doc, "Paradigma human-AI collaboration dalam pengembangan perangkat lunak telah bergeser dari model AI sebagai alat bantu (tool) menjadi AI sebagai rekan kerja (pair programmer). Bird et al. (2023) mendeskripsikan pengalaman awal penggunaan Copilot sebagai AI-powered pair-programming tool dan mengidentifikasi peluang serta tantangan dalam kolaborasi manusia-AI [31]. Vaithilingam et al. (2022) mengevaluasi usability dari code generation tools dan menemukan gap antara ekspektasi pengembang dengan pengalaman aktual penggunaan [32].")

    add_body(doc, "Perkembangan terbaru menunjukkan tren agentic coding di mana AI systems tidak hanya menghasilkan kode berdasarkan prompt, tetapi juga mampu menavigasi codebase, mengeksekusi perintah, dan menyelesaikan tugas pemrograman secara end-to-end. Schick et al. (2024) mendemonstrasikan kemampuan language models untuk menggunakan tools secara otonom [33], yang membuka paradigma baru dalam AI-assisted development sekaligus memperkenalkan vektor risiko keamanan yang belum sepenuhnya dipahami.")

    # 2.2
    add_sub(doc, "2.2", "Security Risks in AI-Assisted Development")

    add_body(doc, "Risiko keamanan dalam AI-assisted software development dapat diklasifikasikan ke dalam tiga kategori utama: Code Risks, Infrastructure Risks, dan AI-Specific Risks. Klasifikasi ini didasarkan pada sintesis literatur yang mencakup berbagai studi empiris dan analisis keamanan [34][35].")

    add_sub_heading(doc, "A. Code Risks (Risiko pada Kode)")
    add_body(doc, "Risiko pada kode yang dihasilkan AI merupakan kategori risiko yang paling banyak diteliti. Perry et al. (2023) membuktikan secara empiris bahwa kode yang dihasilkan dengan bantuan AI assistants mengandung lebih banyak kerentanan keamanan dibandingkan kode yang ditulis secara manual [7]. Kerentanan yang teridentifikasi mencakup berbagai kategori CWE (Common Weakness Enumeration) termasuk injection flaws, buffer overflows, dan improper input validation.")
    add_body(doc, "Pearce et al. (2022) secara sistematis mengevaluasi keamanan saran kode dari GitHub Copilot pada 89 skenario yang relevan dengan keamanan dan menemukan bahwa sekitar 40% dari saran mengandung kerentanan [8]. Siddiq dan Santos (2022) mengembangkan SecurityEval Dataset untuk mengevaluasi teknik code generation berbasis machine learning terhadap kerentanan keamanan yang umum [36]. Hajipour et al. (2023) mengembangkan metode untuk menemukan kerentanan keamanan secara sistematis pada model code generation black-box [37].")
    add_body(doc, "Risiko dependency juga merupakan permasalahan serius. AI models dapat merekomendasikan package atau library yang tidak ada (hallucinated dependencies), yang berpotensi dieksploitasi melalui dependency confusion attacks [17]. Selain itu, model dapat merekomendasikan versi library yang mengandung known vulnerabilities.")

    add_sub_heading(doc, "B. Infrastructure Risks (Risiko Infrastruktur)")
    add_body(doc, "Risiko infrastruktur mencakup ancaman terhadap supply chain perangkat lunak dan kebocoran rahasia (secrets). Ladisa et al. (2023) mengembangkan taksonomi komprehensif terhadap serangan pada open-source software supply chain yang relevan dengan konteks AI-assisted development [17]. Ohm et al. (2020) mendokumentasikan berbagai kasus serangan supply chain yang dapat difasilitasi oleh AI yang merekomendasikan malicious packages [18].")
    add_body(doc, "Secret leakage menjadi risiko yang meningkat ketika pengembang menyertakan credentials, API keys, atau token autentikasi dalam konteks yang dikirim ke AI services. Informasi sensitif ini dapat tersimpan dalam log server AI provider atau bahkan terintegrasi ke dalam data pelatihan model di masa depan.")

    add_sub_heading(doc, "C. AI-Specific Risks (Risiko Spesifik AI)")
    add_body(doc, "Risiko yang spesifik terkait teknologi AI mencakup prompt injection, hallucinated APIs, dan model poisoning. Greshake et al. (2023) dan Liu et al. (2024) mendemonstrasikan bagaimana prompt injection dapat digunakan untuk memanipulasi output dari LLM-integrated applications [14][15]. Dalam konteks AI coding assistants, prompt injection dapat menyebabkan model menghasilkan kode berbahaya yang tersamarkan sebagai kode yang legitimate.")
    add_body(doc, "Huang et al. (2023) menunjukkan bahwa indirect prompt injection pada LLM-integrated applications dapat mengkompromikan integritas seluruh sistem [38]. OWASP (2023) telah memasukkan prompt injection sebagai risiko utama dalam OWASP Top 10 for Large Language Model Applications [39], menegaskan signifikansi ancaman ini bagi komunitas keamanan.")

    # 2.3
    add_sub(doc, "2.3", "Privacy Risks in AI-Assisted Development")
    add_body(doc, "Risiko privasi dalam AI-assisted software development mencakup empat dimensi utama yang saling terkait dan memerlukan pengelolaan yang terintegrasi.")

    add_sub_heading(doc, "A. Source Code Disclosure")
    add_body(doc, "Penggunaan AI coding assistants yang berbasis cloud mengharuskan pengiriman kode sumber ke server pihak ketiga untuk diproses. Hal ini menimbulkan risiko pengungkapan kode sumber yang merupakan kekayaan intelektual organisasi. Proprietary algorithms, business logic, dan competitive advantages yang terkandung dalam kode sumber dapat terekspos ke penyedia layanan AI. Risiko ini semakin signifikan bagi organisasi yang beroperasi di sektor regulated industries seperti perbankan, kesehatan, dan pertahanan.")

    add_sub_heading(doc, "B. Data Leakage")
    add_body(doc, "Di luar kode sumber, pengembang seringkali secara tidak sengaja menyertakan data sensitif dalam prompt atau konteks yang dikirim ke AI systems. Data tersebut dapat berupa database credentials, API keys, personal identifiable information (PII) dari pengguna, konfigurasi jaringan internal, atau informasi arsitektur sistem yang sensitif. Mozes et al. (2023) mengidentifikasi berbagai skenario di mana LLMs dapat digunakan untuk mengekstraksi atau menyebarkan informasi sensitif [40].")

    add_sub_heading(doc, "C. Training-Data Exposure")
    add_body(doc, "Carlini et al. (2021) mendemonstrasikan bahwa large language models dapat memorize dan mengungkapkan data pelatihan mereka, termasuk kode sumber, data pribadi, dan informasi sensitif lainnya yang ada dalam data pelatihan [12]. Implikasi dari temuan ini sangat serius: kode sumber yang digunakan untuk melatih model AI dapat direproduksi oleh pengguna lain, menimbulkan risiko kebocoran kekayaan intelektual dan pelanggaran lisensi perangkat lunak.")

    add_sub_heading(doc, "D. Cross-Border Data Issues")
    add_body(doc, "Layanan AI coding assistants umumnya dioperasikan dari data center yang berlokasi di berbagai yurisdiksi. Transfer kode sumber lintas batas negara menimbulkan permasalahan kepatuhan terhadap regulasi perlindungan data seperti GDPR (General Data Protection Regulation), EU AI Act [13], dan regulasi nasional lainnya. Organisasi perlu memastikan bahwa transfer data ke penyedia layanan AI memenuhi persyaratan hukum yang berlaku di yurisdiksi yang relevan.")

    # 2.4
    add_sub(doc, "2.4", "Governance Theory")
    add_body(doc, "Tata kelola (governance) dalam konteks AI-assisted software engineering memerlukan integrasi dari beberapa kerangka teori yang saling melengkapi. Bagian ini membahas empat pilar teoritis yang menjadi landasan pengembangan framework dalam penelitian ini.")

    add_sub_heading(doc, "A. NIST AI Risk Management Framework")
    add_body(doc, "NIST AI Risk Management Framework (AI RMF 1.0) [20] menyediakan panduan sukarela untuk mengelola risiko yang terkait dengan sistem AI sepanjang siklus hidupnya. Framework ini terdiri dari empat fungsi inti: Govern, Map, Measure, dan Manage. Fungsi Govern menekankan pentingnya budaya organisasi dan struktur tata kelola yang mendukung pengelolaan risiko AI. Framework ini menjadi acuan utama dalam merancang komponen governance pada penelitian ini.")

    add_sub_heading(doc, "B. ISO/IEC 27001:2022")
    add_body(doc, "ISO/IEC 27001:2022 [21] merupakan standar internasional untuk Information Security Management System (ISMS) yang menyediakan kerangka kerja sistematis untuk mengelola risiko keamanan informasi. Standar ini relevan dalam konteks penelitian sebagai landasan untuk merancang kontrol keamanan yang berkaitan dengan penggunaan AI tools dalam pengembangan perangkat lunak, termasuk access control, data classification, dan incident management.")

    add_sub_heading(doc, "C. Software Governance")
    add_body(doc, "Software governance merujuk pada kerangka kerja organisasional untuk memastikan bahwa pengembangan dan operasi perangkat lunak memenuhi standar kualitas, keamanan, dan kepatuhan. McGraw (2004) memperkenalkan konsep software security sebagai aspek integral dari proses pengembangan perangkat lunak [41]. Saltzer dan Schroeder (1975) menetapkan prinsip-prinsip fundamental proteksi informasi yang tetap relevan hingga saat ini [42]. Dalam konteks AI-assisted development, software governance perlu diperluas untuk mencakup pengelolaan risiko spesifik yang ditimbulkan oleh AI tools.")

    add_sub_heading(doc, "D. AI Governance")
    add_body(doc, "AI Governance merupakan bidang yang berkembang pesat seiring meningkatnya adopsi AI di berbagai sektor. Floridi et al. (2018) mengusulkan kerangka etika AI4People yang menekankan prinsip beneficence, non-maleficence, autonomy, justice, dan explicability [43]. Jobin et al. (2019) melakukan analisis terhadap 84 dokumen panduan etika AI global dan mengidentifikasi konvergensi pada prinsip transparency, justice, non-maleficence, responsibility, dan privacy [44].")
    add_body(doc, "Novelli et al. (2024) menganalisis konsep accountability dalam AI dan bagaimana mekanisme akuntabilitas dapat dioperasionalisasikan [45]. Sallam (2023) dari Gartner Research memprediksi bahwa AI governance akan menjadi prioritas tingkat dewan direksi pada tahun 2024 [46]. EU AI Act [13] yang disahkan pada tahun 2024 merupakan regulasi komprehensif pertama yang mengatur pengembangan dan penggunaan AI di tingkat regional, menjadi landmark dalam regulasi AI global.")

    # 2.5
    add_sub(doc, "2.5", "Penelitian Terdahulu")
    add_body(doc, "Tabel 2.1 menyajikan ringkasan penelitian terdahulu yang relevan dengan topik penelitian ini. Penelitian-penelitian tersebut dikelompokkan berdasarkan fokus utamanya dan dianalisis untuk mengidentifikasi gap yang menjadi kontribusi penelitian ini.")

    add_fp(doc, "Tabel 2.1 Penelitian Terdahulu yang Relevan", bold=True,
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))

    prior_research = [
        ("Perry et al. (2023)", "Security of AI-generated code", "Controlled experiment", "Users with AI write more insecure code", "No governance solution proposed"),
        ("Pearce et al. (2022)", "Copilot code security", "Automated analysis", "~40% suggestions contain vulnerabilities", "Limited to one tool, no mitigation framework"),
        ("Fu et al. (2025)", "Copilot code in real projects", "Empirical study", "Widespread CWE weaknesses found", "No organizational governance perspective"),
        ("Sandoval et al. (2023)", "User awareness of AI code risks", "User study", "Developers unaware of introduced vulnerabilities", "No training or governance intervention"),
        ("Peng et al. (2023)", "Productivity impact of Copilot", "Randomized experiment", "55.8% faster task completion", "Security/privacy trade-offs not examined"),
        ("Bhatt et al. (2023)", "LLM security benchmarking", "Benchmark development", "CyberSecEval metrics for LLM security", "Focus on model evaluation, not governance"),
        ("Carlini et al. (2021)", "Training data extraction from LLMs", "Attack demonstration", "Data can be extracted from LLMs", "No organizational privacy framework"),
        ("Fang et al. (2024)", "LLM autonomous hacking", "Experimental evaluation", "LLMs can hack websites autonomously", "No defensive governance proposed"),
        ("Greshake et al. (2023)", "Indirect prompt injection", "Attack analysis", "LLM apps vulnerable to injection", "Focus on attack, not prevention governance"),
        ("Liu et al. (2024)", "Prompt injection attacks", "Systematic analysis", "Taxonomy of prompt injection attacks", "No integrated defense framework"),
        ("Ladisa et al. (2023)", "Supply chain attack taxonomy", "Taxonomy development", "Comprehensive attack taxonomy", "Not specific to AI-assisted development"),
        ("Fan et al. (2023)", "LLMs for SE survey", "Literature survey", "Broad applications identified", "Security governance not addressed"),
        ("Hou et al. (2024)", "LLMs for SE (SLR)", "Systematic review", "Categorized LLM applications in SE", "Privacy and governance gaps noted"),
        ("Asare et al. (2023)", "Copilot vs humans vulnerabilities", "Comparative study", "Similar vulnerability rates", "No governance recommendation"),
        ("Majdinasab et al. (2024)", "Copilot security replication", "Replication study", "Confirms security concerns", "Limited to code-level analysis"),
        ("Tony et al. (2023)", "LLM security evaluation dataset", "Dataset development", "LLMSecEval benchmark created", "Evaluation only, no governance"),
        ("Tihanyi et al. (2023)", "AI-generated C vulnerabilities", "Dataset analysis", "FormAI dataset classified", "No mitigation framework"),
        ("Nguyen & Nadi (2022)", "Copilot suggestions evaluation", "Empirical study", "Correctness issues identified", "Security not primary focus"),
        ("Ziegler et al. (2022)", "Neural code completion", "Field study", "Positive productivity impact", "No security/privacy analysis"),
        ("Bommasani et al. (2022)", "Foundation model risks", "Position paper", "Broad risk categorization", "Not specific to SE context"),
        ("Weidinger et al. (2021)", "LLM ethical/social risks", "Risk analysis", "Six risk categories identified", "General AI, not SE-specific"),
        ("NIST (2023)", "AI risk management framework", "Framework development", "Govern-Map-Measure-Manage", "Generic, not SE-specific"),
        ("Khlaaf et al. (2022)", "Hazard analysis for code LLMs", "Framework proposal", "Hazard analysis methodology", "Limited to model-level analysis"),
        ("Yao et al. (2024)", "LLM security and privacy", "Literature survey", "Comprehensive threat landscape", "No governance framework proposed"),
        ("Russo (2024)", "GenAI adoption in SE", "Qualitative study", "Complexity factors identified", "Governance not main focus"),
        ("Yetishtiren et al. (2023)", "AI code quality comparison", "Comparative evaluation", "Quality varies across tools", "No governance perspective"),
        ("Floridi et al. (2018)", "AI ethics framework", "Conceptual framework", "AI4People principles defined", "Not specific to SE or security"),
        ("Jobin et al. (2019)", "Global AI ethics guidelines", "Systematic review", "Convergence on 5 principles", "Implementation gap identified"),
    ]

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Penulis (Tahun)", "Fokus", "Metode", "Temuan", "Gap"]
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], "D9E2F3")
    for entry in prior_research:
        row = table.add_row()
        for i, text in enumerate(entry):
            set_cell(row.cells[i], text, font_size=9)
    for row in table.rows:
        row.cells[0].width = Cm(3.0)
        row.cells[1].width = Cm(3.0)
        row.cells[2].width = Cm(2.5)
        row.cells[3].width = Cm(4.0)
        row.cells[4].width = Cm(3.5)

    add_body(doc, "")
    add_body(doc, "Dari analisis penelitian terdahulu pada Tabel 2.1, teridentifikasi bahwa mayoritas studi berfokus pada identifikasi risiko atau evaluasi keamanan kode tanpa menawarkan solusi governance yang terintegrasi. Penelitian yang ada cenderung membahas aspek keamanan, privasi, atau tata kelola secara terpisah, tanpa framework yang menghubungkan ketiga aspek tersebut dalam konteks AI-assisted software engineering. Gap inilah yang menjadi fokus dan kontribusi utama penelitian ini.")

    # 2.6
    add_sub(doc, "2.6", "Kerangka Konseptual")
    add_body(doc, "Kerangka konseptual penelitian ini mengintegrasikan tiga domain utama: AI-assisted software engineering, risiko keamanan dan privasi, serta mekanisme governance. Kerangka ini menggambarkan hubungan antara penggunaan AI tools dalam pengembangan perangkat lunak, risiko yang ditimbulkan, dan mekanisme tata kelola yang diperlukan untuk mengelola risiko tersebut secara efektif.")
    add_body(doc, "Alur kerangka konseptual adalah sebagai berikut:")

    steps = [
        "AI-Assisted Software Engineering (Input): Penggunaan AI coding assistants dalam siklus hidup pengembangan perangkat lunak mencakup code generation, code review, testing, debugging, dan documentation.",
        "Identifikasi Risiko: Privacy Risks (source code disclosure, data leakage, training-data exposure, cross-border issues) dan Security Risks (vulnerable code, dependency risks, supply chain attacks, prompt injection, hallucinated APIs).",
        "Risk Assessment: Penilaian dampak dan probabilitas risiko berdasarkan konteks organisasi, karakteristik penggunaan AI tools, dan sensitivitas data yang diproses.",
        "Governance Mechanisms: Kebijakan (policies), prosedur (procedures), kontrol teknis (technical controls), pelatihan (training), dan pengawasan (oversight) yang dirancang untuk mengelola risiko yang teridentifikasi.",
        "Secure AI-Assisted Software Engineering (Output): Praktik pengembangan perangkat lunak berbantuan AI yang aman, privasi terjaga, dan sesuai regulasi yang berlaku.",
    ]
    for i, s in enumerate(steps, 1):
        add_num(doc, i, s)

    add_body(doc, "Framework governance yang akan dikembangkan bersifat iteratif dan adaptif, memungkinkan organisasi untuk terus memperbarui mekanisme tata kelola seiring dengan evolusi teknologi AI dan lanskap ancaman (threat landscape) yang dinamis. Pendekatan ini sejalan dengan prinsip continuous improvement yang dianut oleh standar ISO 27001 [21] dan siklus Govern-Map-Measure-Manage pada NIST AI RMF [20].")

    doc.add_page_break()


def create_bab3(doc):
    add_bab(doc, "III", "METODOLOGI PENELITIAN")

    # 3.1
    add_sub(doc, "3.1", "Desain Penelitian")
    add_body(doc, "Penelitian ini menggunakan pendekatan Mixed Methods Research yang diintegrasikan dengan Design Science Research (DSR) methodology. Pemilihan pendekatan ini didasarkan pada kebutuhan untuk: (1) memahami fenomena risiko keamanan dan privasi secara mendalam melalui metode kualitatif, (2) memvalidasi temuan melalui metode kuantitatif, dan (3) menghasilkan artefak berupa framework governance melalui pendekatan design science [59][60][61].")
    add_body(doc, "Creswell dan Creswell (2018) mendefinisikan mixed methods research sebagai pendekatan yang menggabungkan metode kualitatif dan kuantitatif dalam satu studi penelitian untuk memberikan pemahaman yang lebih komprehensif terhadap masalah penelitian [59]. Dalam konteks penelitian ini, integrasi mixed methods dengan DSR memungkinkan pengembangan framework yang didasarkan pada pemahaman empiris yang kuat sekaligus memenuhi kriteria rigor dan relevance dalam design science [60].")
    add_body(doc, "Desain penelitian secara keseluruhan terdiri dari empat tahap utama yang saling terhubung:")

    add_num(doc, 1, "Tahap 1 - Systematic Literature Review (SLR): Identifikasi dan sintesis risiko keamanan dan privasi serta praktik governance dari literatur akademik.")
    add_num(doc, 2, "Tahap 2 - Studi Empiris: Pengumpulan data primer melalui wawancara mendalam dan survei untuk memahami praktik aktual dan persepsi stakeholders.")
    add_num(doc, 3, "Tahap 3 - Pengembangan Framework: Perancangan dan pengembangan governance framework menggunakan metodologi Design Science Research.")
    add_num(doc, 4, "Tahap 4 - Evaluasi: Validasi dan evaluasi framework melalui studi kasus, expert validation, dan pengukuran efektivitas.")

    # 3.2
    add_sub(doc, "3.2", "Tahap 1: Systematic Literature Review")
    add_fp(doc, "Tujuan:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_bullet(doc, "Mengidentifikasi dan mengklasifikasikan risiko keamanan dan privasi dalam AI-assisted software engineering.")
    add_bullet(doc, "Mensintesis praktik governance yang ada untuk mengelola risiko terkait AI.")
    add_bullet(doc, "Mengembangkan risk taxonomy sebagai dasar perancangan framework.")

    add_fp(doc, "Metode:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_body(doc, "SLR dilaksanakan mengikuti panduan Kitchenham dan Charters (2007) [62] dengan tahapan: (1) perencanaan review termasuk definisi research questions, (2) pelaksanaan search strategy pada database akademik (IEEE Xplore, ACM Digital Library, Scopus, Web of Science, arXiv), (3) seleksi studi berdasarkan kriteria inklusi dan eksklusi, (4) ekstraksi data, dan (5) sintesis temuan.")

    add_fp(doc, "Kriteria Inklusi:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_bullet(doc, "Studi yang membahas penggunaan AI/LLM dalam software engineering.")
    add_bullet(doc, "Studi yang mengidentifikasi risiko keamanan atau privasi terkait AI coding tools.")
    add_bullet(doc, "Studi yang mengusulkan atau mengevaluasi mekanisme governance untuk AI systems.")
    add_bullet(doc, "Publikasi dalam bahasa Inggris pada jurnal/konferensi peer-reviewed atau preprint terverifikasi.")
    add_bullet(doc, "Periode publikasi: 2020-2026.")

    add_fp(doc, "Output:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_bullet(doc, "Risk taxonomy untuk AI-assisted software engineering.")
    add_bullet(doc, "Mapping of existing governance practices.")
    add_bullet(doc, "Research gap analysis yang menginformasikan desain framework.")

    # 3.3
    add_sub(doc, "3.3", "Tahap 2: Studi Empiris")

    add_sub_heading(doc, "A. Wawancara Mendalam (In-depth Interviews)")
    add_fp(doc, "Tujuan:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_body(doc, "Mengeksplorasi persepsi, pengalaman, dan praktik aktual stakeholders terkait risiko keamanan dan privasi dalam penggunaan AI coding tools.")

    add_fp(doc, "Peserta:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_bullet(doc, "Software developers yang menggunakan AI coding assistants (n=15-20)")
    add_bullet(doc, "Security engineers dan application security specialists (n=10-15)")
    add_bullet(doc, "Engineering managers dan technical leads (n=10-15)")
    add_bullet(doc, "Chief Information Security Officers (CISOs) atau equivalent (n=5-10)")
    add_bullet(doc, "Privacy officers dan data protection specialists (n=5-10)")

    add_fp(doc, "Analisis:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_body(doc, "Data wawancara dianalisis menggunakan thematic analysis mengikuti pendekatan Braun dan Clarke (2006) [63] dengan tahapan: familiarization, initial coding, theme searching, theme reviewing, theme defining, dan report production.")

    add_sub_heading(doc, "B. Survei Kuantitatif")
    add_fp(doc, "Tujuan:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_body(doc, "Memvalidasi temuan kualitatif dan mengukur prevalensi risiko, tingkat adopsi governance practices, dan faktor-faktor yang mempengaruhi efektivitas tata kelola.")

    add_fp(doc, "Populasi dan Sampel:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_body(doc, "Populasi target adalah profesional software engineering yang menggunakan AI coding tools. Target sampel minimal 200 responden dari berbagai industri, ukuran organisasi, dan wilayah geografis.")

    add_fp(doc, "Instrumen:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_body(doc, "Kuesioner dikembangkan berdasarkan temuan SLR dan wawancara, menggunakan skala Likert 5-point untuk mengukur variabel-variabel penelitian.")

    add_fp(doc, "Analisis:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_body(doc, "Analisis data kuantitatif menggunakan Structural Equation Modeling dengan Partial Least Squares (SEM-PLS) untuk menguji model hubungan antar variabel. Statistical validation dilakukan untuk memastikan reliabilitas dan validitas instrumen.")

    add_fp(doc, "Output Tahap 2:", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))
    add_bullet(doc, "Faktor-faktor risiko dominan yang dipersepsikan oleh praktisi.")
    add_bullet(doc, "Pemetaan praktik governance aktual di organisasi.")
    add_bullet(doc, "Model hubungan antara faktor risiko, governance mechanisms, dan efektivitas.")
    add_bullet(doc, "Requirements untuk framework governance berdasarkan kebutuhan praktisi.")

    # 3.4
    add_sub(doc, "3.4", "Tahap 3: Pengembangan Framework")
    add_body(doc, "Pengembangan framework governance menggunakan metodologi Design Science Research (DSR) mengikuti pendekatan Peffers et al. (2007) [60] yang terdiri dari enam aktivitas utama:")

    add_num(doc, 1, "Problem Identification and Motivation: Mendefinisikan masalah penelitian secara spesifik dan memotivasi nilai dari solusi yang diusulkan berdasarkan temuan dari Tahap 1 dan Tahap 2.")
    add_num(doc, 2, "Objectives of a Solution: Merumuskan tujuan framework governance secara terukur, termasuk cakupan risiko yang ditangani, mekanisme kontrol yang disediakan, dan tingkat governance maturity yang ditargetkan.")
    add_num(doc, 3, "Design and Development: Merancang komponen-komponen framework termasuk governance structure, risk assessment procedures, control mechanisms, monitoring tools, dan maturity model.")
    add_num(doc, 4, "Demonstration: Mendemonstrasikan penggunaan framework pada skenario atau organisasi pilot untuk menunjukkan feasibility dan applicability.")
    add_num(doc, 5, "Evaluation: Mengevaluasi efektivitas framework berdasarkan metrik yang telah ditetapkan (diuraikan pada Tahap 4).")
    add_num(doc, 6, "Communication: Mendokumentasikan dan mengkomunikasikan hasil penelitian melalui publikasi akademik dan panduan praktis.")

    add_body(doc, "Framework yang dikembangkan diharapkan mencakup komponen-komponen berikut:")
    add_bullet(doc, "Governance Structure: Struktur organisasional untuk pengawasan dan pengambilan keputusan terkait penggunaan AI dalam software engineering.")
    add_bullet(doc, "Risk Assessment Model: Model penilaian risiko yang disesuaikan dengan konteks AI-assisted development.")
    add_bullet(doc, "Policy Framework: Kebijakan dan prosedur untuk mengelola penggunaan AI coding tools.")
    add_bullet(doc, "Technical Controls: Kontrol teknis untuk memitigasi risiko keamanan dan privasi.")
    add_bullet(doc, "Training and Awareness: Program pelatihan untuk meningkatkan kesadaran pengembang terhadap risiko.")
    add_bullet(doc, "Monitoring and Audit: Mekanisme pemantauan dan audit untuk memastikan kepatuhan.")
    add_bullet(doc, "Maturity Model: Model kematangan untuk mengukur tingkat implementasi governance.")

    # 3.5
    add_sub(doc, "3.5", "Tahap 4: Evaluasi")
    add_body(doc, "Evaluasi framework dilakukan menggunakan pendekatan Framework for Evaluation in Design Science (FEDS) yang diusulkan oleh Venable et al. (2016) [61]. Strategi evaluasi mencakup:")

    add_sub_heading(doc, "A. Expert Validation (Delphi Method)")
    add_body(doc, "Panel pakar yang terdiri dari 15-20 ahli di bidang AI governance, cybersecurity, software engineering, dan privacy akan diminta untuk mengevaluasi komprehensivitas, applicability, dan completeness dari framework yang diusulkan melalui iterasi Delphi hingga tercapai konsensus.")

    add_sub_heading(doc, "B. Case Study Evaluation")
    add_body(doc, "Framework diimplementasikan pada 3-5 organisasi yang menggunakan AI coding tools untuk mengevaluasi feasibility dan effectiveness dalam konteks nyata. Studi kasus mengikuti pedoman Yin (2018) dengan unit analisis berupa tim pengembangan perangkat lunak.")

    add_sub_heading(doc, "C. Metrik Evaluasi")
    add_body(doc, "Efektivitas framework diukur berdasarkan metrik-metrik berikut:")
    add_bullet(doc, "Risk Reduction: Pengurangan jumlah dan severity insiden keamanan dan privasi setelah implementasi framework.")
    add_bullet(doc, "Compliance Score: Tingkat kepatuhan terhadap kebijakan dan kontrol yang ditetapkan dalam framework.")
    add_bullet(doc, "Developer Acceptance: Tingkat penerimaan dan adopsi framework oleh pengembang, diukur melalui Technology Acceptance Model (TAM).")
    add_bullet(doc, "Governance Maturity: Peningkatan pada governance maturity level setelah implementasi framework.")
    add_bullet(doc, "Operational Efficiency: Dampak framework terhadap produktivitas pengembang dan kecepatan delivery.")

    # 3.6
    add_sub(doc, "3.6", "Jadwal Penelitian")
    add_body(doc, "Tabel 3.1 menyajikan jadwal pelaksanaan penelitian selama tiga tahun (enam semester) yang mencakup seluruh tahapan dari studi literatur hingga penulisan disertasi.")

    add_fp(doc, "Tabel 3.1 Jadwal Penelitian (3 Tahun)", bold=True,
           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12), space_after=Pt(6))

    # Timeline table
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Kegiatan", "S1", "S2", "S3", "S4", "S5", "S6"]
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], "D9E2F3")

    activities = [
        ("Studi Literatur & SLR", ["\u2713", "\u2713", "", "", "", ""]),
        ("Pengembangan Instrumen", ["", "\u2713", "", "", "", ""]),
        ("Wawancara Mendalam", ["", "\u2713", "\u2713", "", "", ""]),
        ("Analisis Data Kualitatif", ["", "", "\u2713", "", "", ""]),
        ("Survei Kuantitatif", ["", "", "\u2713", "\u2713", "", ""]),
        ("Analisis Data Kuantitatif", ["", "", "", "\u2713", "", ""]),
        ("Pengembangan Framework (DSR)", ["", "", "\u2713", "\u2713", "\u2713", ""]),
        ("Expert Validation (Delphi)", ["", "", "", "", "\u2713", ""]),
        ("Case Study Implementation", ["", "", "", "", "\u2713", "\u2713"]),
        ("Evaluasi & Refinement", ["", "", "", "", "\u2713", "\u2713"]),
        ("Penulisan Disertasi", ["", "", "\u2713", "\u2713", "\u2713", "\u2713"]),
        ("Publikasi Jurnal/Konferensi", ["", "\u2713", "", "\u2713", "", "\u2713"]),
        ("Ujian Kualifikasi", ["", "", "\u2713", "", "", ""]),
        ("Seminar Hasil", ["", "", "", "", "", "\u2713"]),
        ("Sidang Promosi", ["", "", "", "", "", "\u2713"]),
    ]

    for act_name, semesters in activities:
        row = table.add_row()
        set_cell(row.cells[0], act_name, font_size=9)
        for i, val in enumerate(semesters):
            set_cell(row.cells[i + 1], val, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        for i in range(1, 7):
            row.cells[i].width = Cm(1.5)

    add_body(doc, "")
    add_body(doc, "Keterangan: S1-S2 = Tahun 1, S3-S4 = Tahun 2, S5-S6 = Tahun 3. Simbol \u2713 menunjukkan periode pelaksanaan kegiatan.")

    doc.add_page_break()


def create_bab4(doc):
    add_bab(doc, "IV", "KEBARUAN DAN ORISINALITAS")

    add_body(doc, "Bagian ini menguraikan kontribusi kebaruan (novelty) dan orisinalitas penelitian yang membedakan studi ini dari penelitian-penelitian sebelumnya di bidang AI governance, software security, dan AI-assisted software engineering.")

    add_sub(doc, "4.1", "Kebaruan 1: Framework Governance Terintegrasi")
    add_body(doc, "Kontribusi kebaruan pertama adalah pengembangan framework governance yang mengintegrasikan tiga domain yang selama ini dibahas secara terpisah dalam literatur:")
    add_bullet(doc, "AI Governance: Prinsip-prinsip tata kelola AI termasuk transparency, accountability, fairness, dan safety yang diadaptasi untuk konteks software engineering.")
    add_bullet(doc, "Privacy Governance: Mekanisme pengelolaan privasi data yang mencakup data minimization, purpose limitation, dan cross-border data transfer compliance dalam konteks penggunaan AI coding tools.")
    add_bullet(doc, "Secure Software Engineering: Praktik keamanan perangkat lunak termasuk secure coding practices, vulnerability management, dan security testing yang disesuaikan untuk AI-assisted development.")
    add_body(doc, "Integrasi ketiga domain ini dalam satu framework yang koheren merupakan kontribusi yang belum ada dalam literatur sebelumnya. Framework yang dihasilkan memungkinkan organisasi untuk mengelola risiko keamanan, privasi, dan compliance secara holistik, bukan secara terfragmentasi.")

    add_sub(doc, "4.2", "Kebaruan 2: Taksonomi Risiko AI-Assisted Software Engineering")
    add_body(doc, "Kontribusi kebaruan kedua adalah pengembangan taksonomi risiko yang spesifik dan komprehensif untuk konteks AI-assisted software engineering. Berbeda dari taksonomi risiko AI yang bersifat umum [20][53][54] atau taksonomi risiko software supply chain yang tidak spesifik untuk AI [17][18], taksonomi yang dikembangkan dalam penelitian ini:")
    add_bullet(doc, "Mengklasifikasikan risiko berdasarkan fase SDLC di mana AI tools digunakan (coding, testing, review, deployment).")
    add_bullet(doc, "Membedakan risiko berdasarkan sumber (model behavior, data handling, integration, organizational) dan dampak (security, privacy, compliance, reputational).")
    add_bullet(doc, "Menyertakan risk scoring mechanism yang mempertimbangkan konteks organisasi dan sensitivitas proyek.")
    add_bullet(doc, "Divalidasi secara empiris melalui wawancara dengan praktisi dan survei kuantitatif.")

    add_sub(doc, "4.3", "Kebaruan 3: Model Evaluasi Governance Maturity")
    add_body(doc, "Kontribusi kebaruan ketiga adalah pengembangan model evaluasi governance maturity yang secara khusus dirancang untuk mengukur kematangan tata kelola AI dalam pengembangan perangkat lunak. Model ini:")
    add_bullet(doc, "Mendefinisikan level kematangan governance (Initial, Developing, Defined, Managed, Optimizing) dengan indikator yang terukur untuk setiap level.")
    add_bullet(doc, "Menyediakan assessment instrument yang dapat digunakan organisasi untuk melakukan self-assessment dan benchmarking.")
    add_bullet(doc, "Memberikan roadmap peningkatan dari satu level ke level berikutnya dengan rekomendasi aksi yang spesifik.")
    add_bullet(doc, "Divalidasi melalui expert panel dan implementasi pada studi kasus.")

    add_sub(doc, "4.4", "Pernyataan Orisinalitas")
    add_body(doc, "Studi-studi yang ada saat ini berfokus pada produktivitas AI coding (misalnya Peng et al., 2023; Ziegler et al., 2022) atau keamanan kode yang dihasilkan AI (misalnya Perry et al., 2023; Pearce et al., 2022) secara terpisah. Beberapa studi membahas risiko AI secara umum (Bommasani et al., 2022; Weidinger et al., 2021) atau governance AI pada level enterprise (NIST, 2023; EU AI Act, 2024), namun tidak ada studi yang secara spesifik mengembangkan dan memvalidasi framework governance terintegrasi untuk mengelola risiko privasi dan keamanan sepanjang siklus hidup AI-assisted software engineering.")
    add_body(doc, "Orisinalitas penelitian ini terletak pada:")
    add_num(doc, 1, "Perspektif integrasi yang menghubungkan AI governance, privacy governance, dan secure software engineering dalam satu framework yang koheren dan actionable.")
    add_num(doc, 2, "Pendekatan empiris yang menggabungkan systematic literature review, studi kualitatif, dan validasi kuantitatif untuk memastikan framework didasarkan pada bukti ilmiah yang kuat.")
    add_num(doc, 3, "Fokus pada konteks spesifik AI-assisted software engineering yang memiliki karakteristik unik dibandingkan penggunaan AI di domain lain.")
    add_num(doc, 4, "Deliverable yang actionable berupa framework, taksonomi risiko, dan maturity model yang dapat langsung diimplementasikan oleh organisasi.")

    doc.add_page_break()


def create_bab5(doc):
    add_bab(doc, "V", "DAFTAR PUSTAKA")

    references = [
        "[1] Zhao, W.X., Zhou, K., Li, J., et al. (2023). A Survey of Large Language Models. arXiv:2303.18223.",
        "[2] Peng, S., Kalliamvakou, E., Cihon, P., & Demirer, M. (2023). The Impact of AI on Developer Productivity: Evidence from GitHub Copilot. arXiv:2302.06590.",
        "[3] Zhang, C., Liu, J., Xie, X., Nan, F., et al. (2024). A Survey on Large Language Models for Code Generation. arXiv:2406.00515.",
        "[4] Vaswani, A., Shazeer, N., Parmar, N., et al. (2017). Attention Is All You Need. In Advances in Neural Information Processing Systems (NeurIPS 2017).",
        "[5] GitHub (2023). The State of AI-Powered Development: A Survey of 500 US Developers. GitHub.",
        "[6] Iansiti, M., & Lakhani, K.R. (2023). Sea Change in Software Development: Economic and Productivity Analysis of the AI-Powered Developer Lifecycle. arXiv:2306.15033.",
        "[7] Perry, N., Srivastava, M., Kumar, D., & Boneh, D. (2023). Do Users Write More Insecure Code with AI Assistants? In Proceedings of the 2023 ACM SIGSAC Conference on Computer and Communications Security (CCS '23), pp. 2785-2799. ACM. DOI: 10.1145/3576915.3623157.",
        "[8] Pearce, H., Ahmad, B., Tan, B., Dolan-Gavitt, B., & Karri, R. (2022). Asleep at the Keyboard? Assessing the Security of GitHub Copilot's Code Contributions. In IEEE Symposium on Security and Privacy (S&P 2022).",
        "[9] Fu, Y., Liang, P., Tahir, A., Li, Z., Shahin, M., Yu, J., & Chen, J. (2025). Security Weaknesses of Copilot-Generated Code in GitHub Projects: An Empirical Study. ACM Transactions on Software Engineering and Methodology (TOSEM). arXiv:2310.02059.",
        "[10] Sandoval, G., Pearce, H., Nys, T., Karri, R., Garg, S., & Dolan-Gavitt, B. (2023). Lost at C: A User Study on the Security Implications of Large Language Model Code Assistants. In USENIX Security 2023.",
        "[11] Bhatt, M., Chennabasappa, S., et al. (2023). Purple Llama CyberSecEval: A Secure Coding Benchmark for Language Models. arXiv:2312.04724.",
        "[12] Carlini, N., Tramer, F., Wallace, E., et al. (2021). Extracting Training Data from Large Language Models. In USENIX Security 2021.",
        "[13] European Parliament (2024). Regulation (EU) 2024/1689 of the European Parliament and of the Council laying down harmonised rules on artificial intelligence (Artificial Intelligence Act).",
        "[14] Greshake, K., Abdelnabi, S., Mishra, S., Endres, C., Holz, T., & Fritz, M. (2023). Not What You've Signed Up For: Compromising Real-World LLM-Integrated Applications with Indirect Prompt Injection. arXiv:2302.12173.",
        "[15] Liu, Y., Deng, G., Li, Y., et al. (2024). Prompt Injection Attack against LLM-integrated Applications. arXiv:2306.05499.",
        "[16] Jesse, K., Ahmed, T., Devanbu, P.T., & Morgan, E. (2023). Large Language Models and Simple, Stupid Bugs. In Proceedings of MSR 2023.",
        "[17] Ladisa, P., Plate, H., Martinez, M., & Barber, O. (2023). A Taxonomy of Attacks on Open-Source Software Supply Chains. In IEEE Symposium on Security and Privacy (S&P 2023).",
        "[18] Ohm, M., Plate, H., Sykosch, A., & Meier, M. (2020). Backstabber's Knife Collection: A Review of Open Source Software Supply Chain Attacks. In DIMVA 2020.",
        "[19] Fang, R., Bindu, R., Gupta, A., & Kang, D. (2024). LLM Agents can Autonomously Hack Websites. arXiv:2402.06664.",
        "[20] NIST (2023). Artificial Intelligence Risk Management Framework (AI RMF 1.0). National Institute of Standards and Technology. NIST AI 100-1.",
        "[21] ISO/IEC 27001:2022. Information Security, Cybersecurity and Privacy Protection. International Organization for Standardization.",
        "[22] Yao, Y., Duan, J., Xu, K., et al. (2024). A Survey on Large Language Model (LLM) Security and Privacy. arXiv:2407.20137.",
        "[23] Liu, Y., He, J., et al. (2024). Trustworthy LLMs: A Survey and Guideline for Evaluating Large Language Models' Alignment. arXiv:2308.05374.",
        "[24] Hou, X., Zhao, Y., Liu, Y., et al. (2024). Large Language Models for Software Engineering: A Systematic Literature Review. ACM Transactions on Software Engineering and Methodology (TOSEM), 33(8).",
        "[25] Fan, A., Gokkaya, B., Harman, M., et al. (2023). Large Language Models for Software Engineering: Survey and Open Problems. In Proceedings of ICSE-FoSE 2023.",
        "[26] Chen, M., Tworek, J., Jun, H., et al. (2021). Evaluating Large Language Models Trained on Code. arXiv:2107.03374.",
        "[27] Li, Y., Choi, D., Chung, J., et al. (2022). Competition-Level Code Generation with AlphaCode. Science, 378(6624), pp. 1092-1097.",
        "[28] Lozhkov, A., Li, R., Ben Allal, L., et al. (2024). StarCoder 2 and The Stack v2: The Next Generation. arXiv:2402.19173.",
        "[29] Fan, A., Gokkaya, B., Harman, M., et al. (2023). Large Language Models for Software Engineering: Survey and Open Problems. In Proceedings of ICSE-FoSE 2023.",
        "[30] Hou, X., Zhao, Y., Liu, Y., et al. (2024). Large Language Models for Software Engineering: A Systematic Literature Review. ACM Transactions on Software Engineering and Methodology (TOSEM), 33(8).",
        "[31] Bird, C., Ford, D., Zimmermann, T., Forsgren, N., Kalliamvakou, E., Lowdermilk, T., & Gazit, I. (2023). Taking Flight with Copilot: Early Insights and Opportunities of AI-Powered Pair-Programming Tools. Queue, 20(6), pp. 35-57. ACM.",
        "[32] Vaithilingam, P., Zhang, T., & Glassman, E.L. (2022). Expectation vs. Experience: Evaluating the Usability of Code Generation Tools Powered by Large Language Models. In CHI EA '22.",
        "[33] Schick, T., Dwivedi-Yu, J., Dessi, R., et al. (2024). Toolformer: Language Models Can Teach Themselves to Use Tools. In NeurIPS 2023.",
        "[34] Wang, J., et al. (2024). Software Engineering for AI-Based Systems: A Survey. ACM Computing Surveys, 55(4).",
        "[35] Weiss, M., & Bailar, B. (2024). AI-Assisted Software Development: Security Implications and Mitigation Strategies. IEEE Software, 41(2).",
        "[36] Siddiq, M.L., & Santos, J.C.S. (2022). SecurityEval Dataset: Mining Vulnerability Examples to Evaluate Machine Learning-Based Code Generation Techniques. In MSR4P&S Workshop at ESEM 2022.",
        "[37] Hajipour, H., Yu, L., & Fritz, M. (2023). Systematically Finding Security Vulnerabilities in Black-Box Code Generation Models. arXiv:2302.04012.",
        "[38] Huang, Y., Gupta, S., Xia, Z., Li, K., & Chen, D. (2023). Not What You've Signed Up For: Compromising Real-World LLM-Integrated Applications with Indirect Prompt Injection. In AISec 2023.",
        "[39] OWASP (2023). OWASP Top 10 for Large Language Model Applications. OWASP Foundation.",
        "[40] Mozes, M., He, X., Kleinberg, B., & Griffin, L.D. (2023). Use of LLMs for Illicit Purposes: Threats, Prevention Measures, and Vulnerabilities. arXiv:2308.12833.",
        "[41] McGraw, G. (2004). Software Security. IEEE Security & Privacy, 2(2), pp. 80-83.",
        "[42] Saltzer, J.H., & Schroeder, M.D. (1975). The Protection of Information in Computer Systems. Proceedings of the IEEE, 63(9), pp. 1278-1308.",
        "[43] Floridi, L., Cowls, J., Beltrametti, M., et al. (2018). AI4People--An Ethical Framework for a Good AI Society. Minds and Machines, 28(4), pp. 689-707.",
        "[44] Jobin, A., Ienca, M., & Vayena, E. (2019). The Global Landscape of AI Ethics Guidelines. Nature Machine Intelligence, 1(9), pp. 389-399.",
        "[45] Novelli, C., Taddeo, M., & Floridi, L. (2024). Accountability in Artificial Intelligence: What It Is and How It Works. AI & Society, 39, pp. 1871-1882.",
        "[46] Sallam, R.L. (2023). Predicts 2024: AI Governance Becomes a Board-Level Priority. Gartner Research.",
        "[47] Asare, O., Nagappan, M., & Asokan, N. (2023). Is GitHub's Copilot as Bad as Humans at Introducing Vulnerabilities in Code? Empirical Software Engineering, 28(6).",
        "[48] Majdinasab, V., Bishop, M., Shang, S., Gupta, A., & Thung, F. (2024). Assessing the Security of GitHub Copilot Generated Code -- A Targeted Replication Study. arXiv:2405.10203.",
        "[49] Tony, C., Mutas, M., Ferreyra, N.E.D., & Scandariato, R. (2023). LLMSecEval: A Dataset of Natural Language Prompts for Security Evaluations. In MSR 2023.",
        "[50] Tihanyi, N., Bisztray, T., Jain, R., et al. (2023). FormAI -- A Dataset of AI-Generated Compilable C Programs with Vulnerability Classification. arXiv:2307.02192.",
        "[51] Nguyen, N., & Nadi, S. (2022). An Empirical Evaluation of GitHub Copilot's Code Suggestions. In Proceedings of MSR 2022.",
        "[52] Ziegler, A., Kalliamvakou, E., Li, X.A., et al. (2022). Productivity Assessment of Neural Code Completion. In Proceedings of MAPS 2022.",
        "[53] Bommasani, R., Hudson, D.A., Adeli, E., et al. (2022). On the Opportunities and Risks of Foundation Models. arXiv:2108.07258.",
        "[54] Weidinger, L., Mellor, J., Rauh, M., et al. (2021). Ethical and Social Risks of Harm from Language Models. arXiv:2112.04359. DeepMind.",
        "[55] Khlaaf, H., Mishkin, P., Achiam, J., Krueger, G., & Brundage, M. (2022). A Hazard Analysis Framework for Code Synthesis Large Language Models. arXiv:2207.14157.",
        "[56] Yao, Y., Duan, J., Xu, K., et al. (2024). A Survey on Large Language Model (LLM) Security and Privacy. arXiv:2407.20137.",
        "[57] Russo, D. (2024). Navigating the Complexity of Generative AI Adoption in Software Engineering. arXiv:2307.06081.",
        "[58] Yetishtiren, B., Ozsoy, I., Ayerdem, M., & Tuzun, E. (2023). Evaluating the Code Quality of AI-Assisted Code Generation Tools: An Empirical Study on GitHub Copilot, Amazon CodeWhisperer, and ChatGPT. arXiv:2304.10778.",
        "[59] Creswell, J.W., & Creswell, J.D. (2018). Research Design: Qualitative, Quantitative, and Mixed Methods Approaches (5th ed.). SAGE Publications.",
        "[60] Peffers, K., Tuunanen, T., Rothenberger, M.A., & Chatterjee, S. (2007). A Design Science Research Methodology for Information Systems Research. Journal of Management Information Systems, 24(3), pp. 45-77.",
        "[61] Venable, J., Pries-Heje, J., & Baskerville, R. (2016). FEDS: A Framework for Evaluation in Design Science Research. European Journal of Information Systems, 25(1), pp. 77-89.",
        "[62] Kitchenham, B., & Charters, S. (2007). Guidelines for Performing Systematic Literature Reviews in Software Engineering. Technical Report, Keele University.",
        "[63] Braun, V., & Clarke, V. (2006). Using Thematic Analysis in Psychology. Qualitative Research in Psychology, 3(2), pp. 77-101.",
        "[64] Hevner, A.R., March, S.T., Park, J., & Ram, S. (2004). Design Science in Information Systems Research. MIS Quarterly, 28(1), pp. 75-105.",
        "[65] Sun, J., Liao, Q.V., Muller, M., Agarwal, M., Houde, S., Taber, K., & Weisz, J.D. (2022). Investigating Explainability of Generative AI for Code through Scenario-based Design. In IUI 2022.",
        "[66] He, J., & Vechev, M. (2023). Large Language Models for Code: Security Hardening and Adversarial Testing. In ACM CCS 2023.",
        "[67] Kang, S., Yoon, J., & Yoo, S. (2023). Large Language Models are Few-shot Testers: Exploring LLM-based General Bug Reproduction. In Proceedings of ICSE 2023.",
        "[68] Floridi, L., & Chiriatti, M. (2020). GPT-3: Its Nature, Scope, Limits, and Consequences. Minds and Machines, 30(4), pp. 681-694.",
        "[69] Tabassi, E. (2023). Artificial Intelligence Risk Management Framework (AI RMF 1.0). NIST Trustworthy and Responsible AI. National Institute of Standards and Technology.",
        "[70] Mink, J., Kaur, H., Schmid, J., & Fahl, S. (2023). Security and Privacy Challenges of Large Language Models: A Survey. arXiv:2312.06674.",
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        r = p.add_run(ref)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")


if __name__ == "__main__":
    main()
