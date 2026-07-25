#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build: Modul Pelatihan AI untuk BKPSDM Kab. Demak"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK_BLUE  = RGBColor(0x1A, 0x37, 0x6C)
MID_BLUE   = RGBColor(0x27, 0x6F, 0xBF)
LIGHT_BLUE = RGBColor(0xD6, 0xE8, 0xF7)
GOLD       = RGBColor(0xC9, 0xA2, 0x27)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY       = RGBColor(0xF2, 0xF4, 0xF8)

OUT = "/projects/sandbox/untuk-bersama/modul pengabdian"
DOCX_PATH = os.path.join(OUT, "modul_pelatihan_v1.docx")

def shd(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(s)

def ct(cell, text, bold=False, color=None, sz=10, al=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = al
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(sz)
    if color:
        r.font.color.rgb = color

def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    run = p.add_run(text)
    if level == 1:
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(14)
    elif level == 2:
        run.font.color.rgb = MID_BLUE
        run.font.size = Pt(12)
    else:
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(11)
    return p


def body(doc, text, sz=11, space_before=0, space_after=6, bold=False, italic=False, al=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = al
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(sz)
    r.bold = bold
    r.italic = italic
    return p

def bullet(doc, text, level=0, sz=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(sz)
    return p

def add_box(doc, title, content, box_type="INFO"):
    colours = {
        "INFO":    (LIGHT_BLUE, DARK_BLUE),
        "WARNING": (RGBColor(0xFF,0xF3,0xCD), RGBColor(0x85,0x60,0x04)),
        "EXAMPLE": (RGBColor(0xE8,0xF5,0xE9), RGBColor(0x1B,0x5E,0x20)),
        "STEP":    (RGBColor(0xEF,0xEF,0xFF), DARK_BLUE),
    }
    bg, fg = colours.get(box_type, (LIGHT_BLUE, DARK_BLUE))
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cell = tbl.cell(0, 0)
    shd(hdr_cell, fg)
    ct(hdr_cell, f"  {title}", bold=True, color=WHITE, sz=10)
    body_cell = tbl.cell(1, 0)
    shd(body_cell, bg)
    body_cell.text = ''
    p = body_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(content)
    r.font.size = Pt(10)
    doc.add_paragraph()

def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width  = Cm(21.0)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.0)
    for style_name in ['Normal']:
        style = doc.styles[style_name]
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
    return doc


def add_cover(doc):
    """Halaman 1 – Sampul"""
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    # Top colour bar (simulated via table)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = tbl.cell(0, 0)
    shd(c, DARK_BLUE)
    c.width = Cm(16)
    ct(c, "  BKPSDM KABUPATEN DEMAK", bold=True, color=GOLD, sz=11,
       al=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()

    # Logo placeholder
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_logo.add_run("[GAMBAR: Logo Kabupaten Demak + Logo BKPSDM]")
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = MID_BLUE
    doc.add_paragraph()

    # Tag line
    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = p_tag.add_run("MODUL PELATIHAN")
    rt.bold = True
    rt.font.size = Pt(18)
    rt.font.color.rgb = DARK_BLUE

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = p_sub.add_run("Pemanfaatan Kecerdasan Buatan sebagai Asisten Tematik")
    rs.bold = True
    rs.font.size = Pt(14)
    rs.font.color.rgb = MID_BLUE
    doc.add_paragraph()

    p_sub2 = doc.add_paragraph()
    p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs2 = p_sub2.add_run(
        "untuk Meningkatkan Kualitas Konsultasi dan Interpretasi\n"
        "Regulasi Kepegawaian pada BKPSDM Kabupaten Demak"
    )
    rs2.bold = True
    rs2.font.size = Pt(13)
    rs2.font.color.rgb = DARK_BLUE
    doc.add_paragraph()

    # Gold divider
    pd = doc.add_paragraph()
    pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = pd.add_run("━" * 40)
    rd.font.color.rgb = GOLD
    rd.font.size = Pt(12)
    doc.add_paragraph()


    # Metadata block (cover)
    meta = [
        ("Sasaran Peserta",   "Pegawai BKPSDM Kabupaten Demak (Pemula–Menengah)"),
        ("Durasi",            "2 Jam  |  08.00 – 10.00 WIB"),
        ("Versi",             "v1.0 – Juli 2026"),
        ("Penyelenggara",     "Tim Pengabdian Masyarakat  |  [PERLU DATA: nama institusi]"),
        ("Narasumber",        "[PERLU DATA: nama narasumber dan instansi]"),
    ]
    tbl2 = doc.add_table(rows=len(meta), cols=2)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        shd(tbl2.cell(i, 0), DARK_BLUE)
        ct(tbl2.cell(i, 0), k, bold=True, color=WHITE, sz=10)
        ct(tbl2.cell(i, 1), v, sz=10)
    doc.add_paragraph()

    # Disclaimer strip
    tbl3 = doc.add_table(rows=1, cols=1)
    tbl3.style = 'Table Grid'
    c3 = tbl3.cell(0, 0)
    shd(c3, RGBColor(0xFF, 0xF3, 0xCD))
    c3.text = ''
    p3 = c3.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(
        "⚠  PENAFIAN: Seluruh output AI bersifat bantuan, BUKAN keputusan hukum. "
        "Setiap jawaban WAJIB diverifikasi oleh petugas berwenang sebelum digunakan."
    )
    r3.font.size = Pt(9)
    r3.font.italic = True
    r3.font.color.rgb = RGBColor(0x85, 0x60, 0x04)

    add_page_break(doc)

def add_page_break(doc):
    doc.add_page_break()

def add_frontmatter(doc):
    """Halaman 2–4: Identitas, Penafian, Kata Pengantar, Petunjuk, Kebutuhan Alat"""

    # --- Halaman 2: Identitas & Penafian ---
    heading(doc, "Identitas Modul & Penafian", level=1)
    tbl = doc.add_table(rows=8, cols=2)
    tbl.style = 'Table Grid'
    rows_data = [
        ("Judul",           "Pemanfaatan Kecerdasan Buatan sebagai Asisten Tematik untuk "
                            "Meningkatkan Kualitas Konsultasi dan Interpretasi Regulasi "
                            "Kepegawaian pada BKPSDM Kabupaten Demak"),
        ("Versi",           "v1.0 – Juli 2026"),
        ("Penulis",         "[PERLU DATA: nama penyusun modul]"),
        ("Reviewer",        "[PERLU DATA: nama reviewer]"),
        ("Penyelenggara",   "[PERLU DATA: nama institusi penyelenggara]"),
        ("Tanggal Terbit",  "[PERLU DATA: tanggal]"),
        ("Hak Cipta",       "© 2026 [PERLU DATA: nama institusi]. Dilarang menggandakan "
                            "tanpa izin tertulis."),
        ("Kontak",          "[PERLU DATA: email/telepon panitia]"),
    ]
    for i, (k, v) in enumerate(rows_data):
        shd(tbl.cell(i, 0), LIGHT_BLUE)
        ct(tbl.cell(i, 0), k, bold=True, sz=10)
        ct(tbl.cell(i, 1), v, sz=10)
    doc.add_paragraph()


    add_box(doc,
        "⚠ PENAFIAN PENTING",
        "1. Seluruh output yang dihasilkan AI dalam pelatihan ini bersifat DRAF dan BANTUAN saja.\n"
        "2. Tidak ada output AI yang boleh langsung dijadikan keputusan kepegawaian, surat resmi, "
        "atau dasar hukum tanpa verifikasi petugas berwenang.\n"
        "3. Regulasi kepegawaian berubah sewaktu-waktu. Selalu cek sumber resmi: JDIH BKN, "
        "JDIH Kemenpan-RB, JDIH Pemerintah Kabupaten Demak.\n"
        "4. Peserta dan fasilitator bertanggung jawab penuh atas penggunaan materi ini di luar "
        "sesi pelatihan.",
        box_type="WARNING")

    add_page_break(doc)

    # --- Halaman 3: Kata Pengantar ---
    heading(doc, "Kata Pengantar", level=1)
    body(doc,
        "Perkembangan kecerdasan buatan (AI) membuka peluang baru bagi aparatur sipil negara "
        "(ASN) untuk meningkatkan efisiensi dan kualitas pelayanan kepegawaian. Namun, "
        "pemanfaatan AI dalam konteks regulasi pemerintahan memerlukan kehati-hatian, "
        "pemahaman batas kemampuan teknologi, dan komitmen terhadap akuntabilitas manusia.")
    body(doc,
        "Modul ini dirancang khusus untuk pegawai BKPSDM Kabupaten Demak yang ingin memulai "
        "perjalanan menggunakan AI sebagai asisten kerja—bukan pengganti keahlian dan "
        "penilaian profesional Anda. Selama dua jam pelatihan, peserta akan mempelajari cara "
        "memilih tools yang tepat, menyusun prompt yang efektif, mengintegrasikan AI dengan "
        "ekosistem Google Workspace, dan yang terpenting: memverifikasi setiap output AI "
        "sebelum digunakan.")
    body(doc,
        "Semoga modul ini menjadi langkah awal yang bermanfaat dalam transformasi digital "
        "pelayanan kepegawaian di Kabupaten Demak.")
    doc.add_paragraph()
    body(doc, "[PERLU DATA: kota], [PERLU DATA: tanggal]", al=WD_ALIGN_PARAGRAPH.RIGHT)
    body(doc, "[PERLU DATA: nama dan jabatan penyelenggara]", al=WD_ALIGN_PARAGRAPH.RIGHT)
    add_page_break(doc)

    # --- Halaman 4: Petunjuk + Alat + Daftar Isi ---
    heading(doc, "Petunjuk Penggunaan Modul", level=1)

    heading(doc, "Bagi Peserta", level=2)
    items = [
        "Baca seluruh bagian sebelum sesi dimulai untuk orientasi topik.",
        "Siapkan akun Google (Gmail) aktif dan akses internet stabil.",
        "Gunakan data FIKTIF yang tersedia—JANGAN memasukkan data ASN atau dokumen resmi asli.",
        "Catat pertanyaan di kolom refleksi yang tersedia.",
        "Praktikkan langkah-langkah secara berurutan dan verifikasi setiap output AI.",
    ]
    for it in items:
        bullet(doc, it)
    doc.add_paragraph()

    heading(doc, "Bagi Fasilitator", level=2)
    f_items = [
        "Siapkan lingkungan demo (folder DEMO_BKPSDM_AI) minimal 30 menit sebelum sesi.",
        "Pastikan izin OAuth dan App yang diperlukan sudah aktif pada akun demo.",
        "Tampilkan layar melalui proyektor; pastikan peserta dapat melihat setiap langkah.",
        "Hentikan dan ulangi demonstrasi jika ada peserta yang tertinggal.",
        "Tandai output AI dengan label DRAF pada setiap demonstrasi.",
    ]
    for it in f_items:
        bullet(doc, it)


    doc.add_paragraph()
    heading(doc, "Kebutuhan Alat & Akun Demo", level=2)
    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Alat/Akun", "Spesifikasi", "Wajib/Opsional", "Catatan"]
    for i, h in enumerate(headers):
        shd(tbl2.cell(0, i), DARK_BLUE)
        ct(tbl2.cell(0, i), h, bold=True, color=WHITE, sz=10, al=WD_ALIGN_PARAGRAPH.CENTER)
    alat_rows = [
        ("Laptop/PC", "RAM ≥4 GB, layar ≥13\"", "Wajib", "Browser Chrome/Edge terbaru"),
        ("Internet", "≥5 Mbps stabil", "Wajib", "Hotspot cadangan disarankan"),
        ("Akun Google", "Gmail aktif (bukan akun sekolah terbatas)", "Wajib", "Gunakan akun demo panitia jika tidak punya"),
        ("ChatGPT", "Akun gratis atau Pro", "Wajib", "Plus/Pro untuk fitur Projects & Apps [FITUR TERGANTUNG AKUN/ADMIN]"),
        ("Gemini Workspace", "Google Workspace dengan Gemini diaktifkan", "Opsional", "[FITUR TERGANTUNG AKUN/ADMIN]"),
        ("Google Apps Script", "Akses via script.google.com", "Opsional", "Perlu izin OAuth [FITUR TERGANTUNG AKUN/ADMIN]"),
        ("Data Fiktif", "File tersedia di folder DEMO_BKPSDM_AI", "Wajib", "Disediakan panitia; JANGAN gunakan data asli"),
    ]
    for row_data in alat_rows:
        row = tbl2.add_row()
        for j, val in enumerate(row_data):
            ct(row.cells[j], val, sz=10)
    doc.add_paragraph()
    add_page_break(doc)


def add_section_pembelajaran(doc):
    """Halaman 5–6: Kompetensi & Jadwal"""
    heading(doc, "BAGIAN PEMBELAJARAN: Kompetensi & Jadwal", level=1)

    heading(doc, "Capaian Kompetensi", level=2)
    body(doc, "Setelah mengikuti pelatihan ini, peserta mampu:")
    kompetensi = [
        "K1 – Memilih tools AI yang sesuai kebutuhan kerja BKPSDM.",
        "K2 – Menghubungkan aplikasi AI ke ekosistem Google Workspace secara aman.",
        "K3 – Menyusun prompt terstruktur (Peran–Tujuan–Objek–Konteks–Sumber–Batasan–Format–Konfirmasi).",
        "K4 – Menghasilkan draf jawaban berbasis sumber regulasi yang dapat ditelusuri.",
        "K5 – Memverifikasi output AI menggunakan sumber resmi sebelum digunakan.",
    ]
    for k in kompetensi:
        bullet(doc, k)

    heading(doc, "Luaran Peserta", level=2)
    luaran = [
        "✔  Satu prompt terstruktur siap pakai untuk konsultasi kepegawaian.",
        "✔  Satu draf jawaban terverifikasi dengan dasar hukum yang dicantumkan.",
        "✔  Satu rencana pilot 30 hari untuk implementasi di unit kerja.",
    ]
    for l in luaran:
        bullet(doc, l)
    doc.add_paragraph()

    heading(doc, "Jadwal Pelatihan", level=2)
    body(doc, "[TABEL] Jadwal Pelatihan — 08.00–10.00 WIB", bold=True, sz=10)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(["Waktu", "Kegiatan", "Metode", "Durasi"]):
        shd(tbl.cell(0, i), DARK_BLUE)
        ct(tbl.cell(0, i), h, bold=True, color=WHITE, sz=10, al=WD_ALIGN_PARAGRAPH.CENTER)
    jadwal = [
        ("08.00–08.10", "Pembukaan & Pretest (5 soal)", "Tanya jawab lisan / form", "10 mnt"),
        ("08.10–08.25", "Peta Tools AI: Memilih alat yang tepat", "Ceramah + tabel komparasi", "15 mnt"),
        ("08.25–08.40", "Prompt Engineering & Keamanan Data", "Ceramah + kotak contoh", "15 mnt"),
        ("08.40–08.55", "Alur Regulasi: Konsultasi → Verifikasi", "Ceramah + diagram", "15 mnt"),
        ("08.55–09.05", "Penyiapan Alat & Folder Demo", "Panduan langkah", "10 mnt"),
        ("09.05–09.25", "Demo Gmail & Google Drive dengan AI", "Demonstrasi fasilitator", "20 mnt"),
        ("09.25–09.40", "Demo Apps Script + Gemini API", "Demonstrasi + template", "15 mnt"),
        ("09.40–09.50", "Praktik Mandiri & Verifikasi", "Latihan terpandu", "10 mnt"),
        ("09.50–10.00", "Post-test, Refleksi & Penutupan", "Kuesioner + diskusi", "10 mnt"),
    ]
    for row_data in jadwal:
        row = tbl.add_row()
        for j, val in enumerate(row_data):
            ct(row.cells[j], val, sz=10)
    doc.add_paragraph()
    add_page_break(doc)


def add_section_A(doc):
    """Halaman 7–9: Bagian A – AI untuk BKPSDM"""
    heading(doc, "BAGIAN A: Kecerdasan Buatan untuk BKPSDM", level=1)

    heading(doc, "A.1 Masalah Kerja BKPSDM & Peluang AI", level=2)
    body(doc,
        "Pegawai BKPSDM Kabupaten Demak sehari-hari menghadapi volume pertanyaan kepegawaian "
        "yang tinggi: cuti, kenaikan pangkat, disiplin, mutasi, dan penilaian kinerja. "
        "Referensi regulasi tersebar di berbagai peraturan yang kerap diperbarui, "
        "sehingga membutuhkan waktu penelusuran yang signifikan.")
    masalah = [
        "Pertanyaan berulang (cuti tahunan, cuti besar, cuti sakit) menyita waktu layanan.",
        "Regulasi kepegawaian kerap direvisi; risiko mengacu pada pasal yang sudah tidak berlaku.",
        "Pencarian manual di arsip fisik maupun digital memakan waktu 15–30 menit per kasus.",
        "Tidak ada sistem log yang standar untuk pertanyaan-jawaban yang diberikan.",
    ]
    for m in masalah:
        bullet(doc, m)

    heading(doc, "A.2 Manfaat & Batas AI", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    for i, h in enumerate(["✅ Yang Dapat Dilakukan AI", "❌ Yang TIDAK Dapat Dilakukan AI"]):
        shd(tbl.cell(0, i), DARK_BLUE)
        ct(tbl.cell(0, i), h, bold=True, color=WHITE, sz=10)
    manfaat = [
        ("Merangkum pertanyaan dan konteks", "Membuat keputusan kepegawaian resmi"),
        ("Menyusun draf jawaban berdasarkan sumber", "Menjamin kebenaran regulasi terkini"),
        ("Mengelompokkan email berdasarkan urgensi", "Menandatangani atau mengesahkan dokumen"),
        ("Mencari file di Google Drive", "Mengakses sistem SIMPEG/BKN secara langsung"),
        ("Membuat template dan laporan log", "Mewakili keputusan hukum pejabat berwenang"),
    ]
    for b, t in manfaat:
        row = tbl.add_row()
        ct(row.cells[0], b, sz=10)
        ct(row.cells[1], t, sz=10)
    doc.add_paragraph()

    heading(doc, "A.3 Akuntabilitas Manusia & Etika", level=2)
    body(doc,
        "Penggunaan AI dalam layanan kepegawaian tidak mengurangi tanggung jawab petugas. "
        "Setiap output AI adalah DRAF yang harus diverifikasi. Prinsip-prinsip etika berikut "
        "wajib diterapkan dalam setiap penggunaan AI di lingkungan BKPSDM:")
    etika = [
        "Transparansi: Setiap draf AI harus diberi label 'DRAF—WAJIB DIVERIFIKASI'.",
        "Akuntabilitas: Petugas penanda tangan tetap bertanggung jawab penuh.",
        "Kerahasiaan: Data ASN tidak boleh diunggah atau dimasukkan ke prompt AI.",
        "Non-diskriminasi: AI tidak digunakan untuk keputusan yang merugikan individu tanpa bukti.",
        "Verifikasi: Setiap pasal hukum yang dikutip AI wajib dicek ke sumber resmi.",
    ]
    for e in etika:
        bullet(doc, e)
    doc.add_paragraph()

    add_box(doc,
        "⚠ PERINGATAN: Data yang DILARANG Dimasukkan ke AI",
        "DILARANG memasukkan ke prompt AI:\n"
        "• NIP, nama lengkap, alamat, nomor telepon, atau kontak pegawai asli\n"
        "• Data kesehatan, kondisi medis, atau catatan cuti sakit individual\n"
        "• Kasus disiplin aktif dengan identitas nyata\n"
        "• Dokumen rahasia negara atau dokumen berklasifikasi\n"
        "• Kata sandi, OTP, token autentikasi, atau API key\n"
        "• Screenshot sistem SIMPEG atau data BKN yang belum dipublikasikan\n\n"
        "GUNAKAN SELALU: data fiktif atau data yang sudah dipublikasikan secara resmi.",
        box_type="WARNING")
    add_page_break(doc)


def add_section_B(doc):
    """Halaman 10–12: Bagian B – Peta Tools & Prompt Engineering"""
    heading(doc, "BAGIAN B: Peta Tools AI & Rekayasa Prompt", level=1)

    heading(doc, "B.1 Tabel Komparasi Tools AI", level=2)
    body(doc, "[TABEL] Komparasi Tools AI untuk BKPSDM", bold=True, sz=10)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(["Tools", "Fungsi Utama", "Akses", "Risiko", "Contoh Penggunaan"]):
        shd(tbl.cell(0, i), DARK_BLUE)
        ct(tbl.cell(0, i), h, bold=True, color=WHITE, sz=9, al=WD_ALIGN_PARAGRAPH.CENTER)
    tools_data = [
        ("ChatGPT (Free/Plus/Pro)",
         "Chat, upload file, riset web, Projects, Apps, Tasks",
         "Akun openai.com; Plus/Pro untuk Apps [FITUR TERGANTUNG AKUN/ADMIN]",
         "Halusinasi regulasi; data prompt tersimpan di OpenAI",
         "Tanya: 'Apa syarat cuti besar PNS?' + unggah PDF PP"),
        ("Gmail (Google Workspace)",
         "Kelola email, ringkasan, draf balasan",
         "Akun Google; Gemini harus diaktifkan admin [FITUR TERGANTUNG AKUN/ADMIN]",
         "Risiko kirim ke penerima salah; tindakan tidak disengaja",
         "Rangkum 10 email kepegawaian berurgensi tinggi"),
        ("Google Drive/Docs/Sheets",
         "Repositori file, register kasus, template log",
         "Akun Google; izin berbagi sesuai kebutuhan",
         "File sensitif terekspos jika link dibagikan sembarangan",
         "Simpan regulasi aktif; log konsultasi di Sheets"),
        ("Gemini Workspace",
         "Asisten dalam Gmail, Docs, Sheets, Meet",
         "Google Workspace + paket Gemini [FITUR TERGANTUNG AKUN/ADMIN]",
         "TIDAK otomatis memberi akses Gemini API",
         "Ringkas dokumen Docs; sarankan poin agenda rapat"),
        ("Apps Script + Gemini API",
         "Otomasi berulang: log, draf, cek regulasi",
         "script.google.com + API key Gemini / Vertex AI [FITUR TERGANTUNG AKUN/ADMIN]",
         "API key bocor jika disimpan di kode; quota habis",
         "Fungsi tanyaAsisten() di Sheets → Gemini API"),
        ("GitHub (privat)",
         "Versi kode, prompt, SOP, template",
         "Akun GitHub; repo privat wajib",
         "Kode bocor jika repo tidak privat",
         "Simpan versi Apps Script dan template prompt"),
    ]
    for row_data in tools_data:
        row = tbl.add_row()
        for j, val in enumerate(row_data):
            ct(row.cells[j], val, sz=9)
    doc.add_paragraph()

    heading(doc, "B.2 Rekayasa Prompt (Prompt Engineering)", level=2)
    body(doc,
        "Prompt yang baik adalah kunci kualitas output AI. Gunakan kerangka "
        "Peran–Tujuan–Objek–Konteks–Sumber–Batasan–Format–Konfirmasi (P-T-O-K-S-B-F-K):")
    kerangka = [
        ("Peran",      "Tentukan identitas AI: 'Kamu adalah asisten kepegawaian...'"),
        ("Tujuan",     "Apa yang ingin dicapai: 'Tolong jelaskan...'"),
        ("Objek",      "Subjek utama: 'tentang cuti besar PNS golongan III'"),
        ("Konteks",    "Latar belakang: 'untuk pegawai dengan masa kerja 6 tahun'"),
        ("Sumber",     "Regulasi acuan: 'berdasarkan PP 11 Tahun 2017 pasal ...'"),
        ("Batasan",    "Hal yang tidak boleh dilakukan: 'jangan mengarang pasal'"),
        ("Format",     "Bentuk output: 'dalam bentuk poin-poin ringkas'"),
        ("Konfirmasi", "Minta klarifikasi: 'Jika ada yang ambigu, tanyakan dulu.'"),
    ]
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = 'Table Grid'
    for i, h in enumerate(["Elemen", "Penjelasan", "Contoh"]):
        shd(tbl2.cell(0, i), MID_BLUE)
        ct(tbl2.cell(0, i), h, bold=True, color=WHITE, sz=10)
    for k, v in kerangka:
        row = tbl2.add_row()
        ct(row.cells[0], k, bold=True, sz=10)
        ct(row.cells[1], v, sz=10)
        ct(row.cells[2], "—", sz=10)
    doc.add_paragraph()

    add_box(doc,
        "📌 KOTAK CONTOH: Prompt Buruk vs Prompt Baik",
        "❌ PROMPT BURUK:\n"
        "'Apa itu cuti?'\n"
        "→ RISIKO: Jawaban terlalu umum, tidak ada sumber, bisa salah konteks PNS.\n\n"
        "✅ PROMPT BAIK:\n"
        "'Kamu adalah asisten kepegawaian untuk PNS Indonesia. "
        "Tolong jelaskan syarat dan prosedur pengajuan cuti besar bagi PNS golongan III "
        "dengan masa kerja 6 tahun berturut-turut, berdasarkan PP Nomor 11 Tahun 2017 "
        "dan PP Nomor 17 Tahun 2020. "
        "Jangan mengarang pasal. Jika ada hal yang ambigu, tanyakan terlebih dahulu. "
        "Tampilkan dalam format poin-poin ringkas disertai nomor pasal.'",
        box_type="EXAMPLE")

    heading(doc, "B.3 Tabel Risiko & Pengendalian", level=2)
    body(doc, "[TABEL] Risiko Penggunaan AI & Strategi Pengendalian", bold=True, sz=10)
    tbl3 = doc.add_table(rows=1, cols=3)
    tbl3.style = 'Table Grid'
    for i, h in enumerate(["Risiko", "Deskripsi", "Pengendalian"]):
        shd(tbl3.cell(0, i), DARK_BLUE)
        ct(tbl3.cell(0, i), h, bold=True, color=WHITE, sz=10)
    risiko_data = [
        ("Halusinasi", "AI mengarang fakta/pasal yang tidak ada", "Verifikasi setiap pasal ke JDIH; label DRAF wajib"),
        ("Salah pasal", "AI mengutip nomor pasal yang salah", "Cek manual ke PDF regulasi resmi"),
        ("Regulasi kedaluwarsa", "AI mengacu PP yang sudah diubah/dicabut", "Cek status aktif di JDIH sebelum digunakan"),
        ("Bias", "Output tidak proporsional untuk kasus tertentu", "Review manusia wajib; jangan gunakan sebagai satu-satunya acuan"),
        ("Kebocoran data", "Data ASN masuk ke server AI eksternal", "DILARANG memasukkan data pribadi ASN ke prompt"),
        ("Prompt injection (email)", "Email berisi instruksi jahat yang dieksekusi AI", "Tinjau selalu sebelum konfirmasi tindakan"),
        ("Salah penerima email", "AI mengirim ke alamat yang keliru", "Periksa penerima, subjek, isi SEBELUM konfirmasi kirim"),
        ("Salah folder Drive", "File dipindah/dihapus ke lokasi salah", "Operasi hanya dalam folder DEMO_BKPSDM_AI"),
        ("Tindakan tidak disengaja", "AI melakukan tindakan yang tidak diinstruksikan", "Selalu pratinjau dan konfirmasi sebelum eksekusi"),
    ]
    for row_data in risiko_data:
        row = tbl3.add_row()
        for j, val in enumerate(row_data):
            ct(row.cells[j], val, sz=10)
    doc.add_paragraph()
    add_page_break(doc)



def add_section_C(doc):
    """Halaman 13–16: Bagian C – Alur Regulasi & Konsultasi"""
    heading(doc, "BAGIAN C: Alur Konsultasi & Interpretasi Regulasi", level=1)

    heading(doc, "C.1 Kanal & Repositori Kerja", level=2)
    body(doc,
        "Integrasi AI ke dalam alur kerja BKPSDM menggunakan ekosistem Google Workspace "
        "yang sudah ada, diperkuat dengan AI sebagai asisten:")
    kanal = [
        "Email (Gmail): Kanal masuk pertanyaan kepegawaian dari pegawai/unit kerja.",
        "Google Drive: Repositori regulasi aktif, template, dan hasil konsultasi.",
        "Google Sheets: Register kasus konsultasi dan log audit.",
        "ChatGPT: Asisten ad hoc untuk pertanyaan tidak berulang.",
        "Apps Script: Proses berulang otomatis (log, notifikasi, draf).",
        "GitHub (privat): Versioning kode, prompt, SOP, dan template.",
    ]
    for k in kanal:
        bullet(doc, k)
    doc.add_paragraph()

    heading(doc, "C.2 Diagram Alur Konsultasi AI", level=2)
    body(doc, "[DIAGRAM] Alur Konsultasi Kepegawaian dengan Bantuan AI", bold=True, sz=10)
    add_box(doc,
        "📊 DIAGRAM: Alur Konsultasi",
        "Pertanyaan Masuk (Email/Tatap Muka)\n"
        "       ↓\n"
        "Anonimisasi: Hapus NIP, nama, kontak dari pertanyaan\n"
        "       ↓\n"
        "Pilih Regulasi Acuan (dari folder DEMO_BKPSDM_AI/regulasi)\n"
        "       ↓\n"
        "Susun Prompt Terstruktur → Input ke AI\n"
        "       ↓\n"
        "Terima Draf AI (label: DRAF—WAJIB DIVERIFIKASI)\n"
        "       ↓\n"
        "Cek Pasal: Buka PDF regulasi resmi, cocokkan nomor pasal\n"
        "       ↓\n"
        "Validasi oleh Petugas Berwenang\n"
        "       ↓\n"
        "Jawaban Final Dikirim + Log dicatat di Sheets",
        box_type="INFO")

    heading(doc, "C.3 Pengelolaan Regulasi", level=2)
    body(doc,
        "Regulasi kepegawaian dikelola dengan konvensi penamaan standar agar mudah "
        "ditemukan dan statusnya jelas:")
    add_box(doc,
        "📁 Konvensi Nama File Regulasi",
        "Format: JENIS_NOMOR_TAHUN_TOPIK_STATUS.pdf\n\n"
        "Contoh:\n"
        "  PP_011_2017_ManajemenPNS_AKTIF.pdf\n"
        "  PP_017_2020_ManajemenPNS_DIUBAH.pdf\n"
        "  PP_053_2010_DisiplinPNS_DICABUT.pdf\n"
        "  PerBKN_006_2022_PenilaianKinerja_PERLU_VERIFIKASI.pdf\n\n"
        "Status: AKTIF | DIUBAH | DICABUT | PERLU_VERIFIKASI\n"
        "Sumber: JDIH BKN (jdih.bkn.go.id), JDIH Kemenpan-RB, "
        "JDIH Pemerintah Kabupaten Demak [PERLU VERIFIKASI SUMBER RESMI]",
        box_type="INFO")
    add_page_break(doc)



    heading(doc, "C.4 Kasus Fiktif & Format Draf AI", level=2)
    body(doc,
        "Gunakan kasus fiktif berikut untuk latihan. "
        "JANGAN gunakan kasus dengan identitas ASN asli.")
    kasus = [
        ("K1 – Cuti (35%)",    "Pegawai Fiktif A (Gol. III/b, masa kerja 7 thn) mengajukan cuti besar.",
                                "PP 11/2017, PP 17/2020 [PERLU VERIFIKASI SUMBER RESMI]"),
        ("K2 – Disiplin (25%)", "Pegawai Fiktif B tidak masuk 5 hari berturut tanpa keterangan.",
                                "PP 53/2010 [PERLU VERIFIKASI SUMBER RESMI]"),
        ("K3 – Kenaikan Pangkat (20%)", "Pegawai Fiktif C akan naik pangkat dari III/b ke III/c.",
                                "PP 11/2017 Pasal 68–79 [PERLU VERIFIKASI SUMBER RESMI]"),
        ("K4 – Mutasi (10%)",  "Pegawai Fiktif D mengajukan pindah unit kerja.",
                                "PP 11/2017 Pasal 93–104 [PERLU VERIFIKASI SUMBER RESMI]"),
        ("K5 – Kinerja (5%)",  "Unit kerja ingin memahami SKP terbaru.",
                                "PerMenpan-RB 6/2022 [PERLU VERIFIKASI SUMBER RESMI]"),
        ("K6 – Penolakan (5%)", "AI menolak menjawab tanpa menyebut sumber regulasi.",
                                "Uji batas kemampuan AI"),
    ]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    for i, h in enumerate(["Kasus", "Skenario Fiktif", "Regulasi Acuan"]):
        shd(tbl.cell(0, i), DARK_BLUE)
        ct(tbl.cell(0, i), h, bold=True, color=WHITE, sz=10)
    for k, s, r in kasus:
        row = tbl.add_row()
        ct(row.cells[0], k, bold=True, sz=10)
        ct(row.cells[1], s, sz=10)
        ct(row.cells[2], r, sz=10)
    doc.add_paragraph()

    heading(doc, "C.5 Format Draf Output AI", level=2)
    add_box(doc,
        "📋 FORMAT STANDAR DRAF OUTPUT AI",
        "═══════════════════════════════════════════════════════\n"
        "DRAF — WAJIB DIVERIFIKASI PETUGAS BERWENANG\n"
        "═══════════════════════════════════════════════════════\n\n"
        "RINGKASAN:\n"
        "[Jawaban singkat 2-3 kalimat]\n\n"
        "DASAR HUKUM:\n"
        "• [Nama peraturan] Nomor [X] Tahun [Y] Pasal [Z]\n"
        "  Status: [AKTIF/DIUBAH] [PERLU VERIFIKASI SUMBER RESMI]\n\n"
        "PROSEDUR:\n"
        "1. [Langkah pertama]\n"
        "2. [Langkah berikutnya]\n\n"
        "KLARIFIKASI DIPERLUKAN:\n"
        "• [Pertanyaan klarifikasi jika ada]\n\n"
        "VERIFIKATOR: ____________________  Tgl: ____________\n"
        "═══════════════════════════════════════════════════════",
        box_type="STEP")
    add_page_break(doc)


def add_section_D(doc):
    """Halaman 17–23: Bagian D – Demonstrasi Praktis"""
    heading(doc, "BAGIAN D: Demonstrasi Praktis", level=1)

    heading(doc, "D.1 Persiapan Lingkungan Demo", level=2)
    add_box(doc,
        "✅ LANGKAH PRAKTIK: Checklist Sebelum Demo",
        "□ 1. Akun Google demo aktif (bukan akun pribadi/instansi berisi data asli)\n"
        "□ 2. Folder DEMO_BKPSDM_AI sudah dibuat di Google Drive\n"
        "□ 3. Sub-folder: /regulasi_aktif, /regulasi_diubah_dicabut, /kasus_fiktif,\n"
        "     /hasil_verifikasi, /template_kode_prompt_sop\n"
        "□ 4. File kasus fiktif K1–K6 sudah diunggah\n"
        "□ 5. Izin OAuth untuk Apps Script sudah diberikan [FITUR TERGANTUNG AKUN/ADMIN]\n"
        "□ 6. Apps/ChatGPT yang akan digunakan sudah terhubung [FITUR TERGANTUNG AKUN/ADMIN]\n"
        "□ 7. API key Gemini tersimpan di Script Properties (BUKAN di kode)\n"
        "□ 8. Proyektor/layar berbagi siap; peserta dapat melihat dengan jelas",
        box_type="STEP")
    doc.add_paragraph()

    heading(doc, "D.2 Demo Gmail dengan AI", level=2)
    body(doc, "[LANGKAH PRAKTIK] Demo Gmail — Merangkum & Merespons Email Kepegawaian", bold=True, sz=10)
    body(doc,
        "ChatGPT dapat terhubung ke Gmail jika app Gmail sudah diaktifkan dan terhubung "
        "pada akun ChatGPT pengguna. Kemampuan ini bergantung pada akun, paket, wilayah, "
        "perangkat, dan kebijakan administrator. [FITUR TERGANTUNG AKUN/ADMIN]")
    langkah_gmail = [
        "Buka ChatGPT → pastikan app Gmail sudah terhubung (ikon app terlihat di sidebar).",
        "Ketik: 'Cari 5 email terbaru dengan kata kunci cuti atau izin dalam 7 hari terakhir. "
        "Tampilkan pengirim, subjek, dan ringkasan singkat.'",
        "Tinjau hasil: pastikan hanya email fiktif/demo yang ditampilkan.",
        "Kelompokkan berdasarkan urgensi: 'Kelompokkan email tersebut: Segera/Normal/Informasi.'",
        "Ekstrak tindak lanjut: 'Buat daftar tindak lanjut yang diperlukan dari email tersebut.'",
        "Draf balasan (HANYA ke akun uji): prompt → pratinjau → periksa penerima/subjek/isi → konfirmasi.",
    ]
    for i, l in enumerate(langkah_gmail, 1):
        bullet(doc, f"Langkah {i}: {l}")
    doc.add_paragraph()

    add_box(doc,
        "⚠ PERINGATAN: Sebelum Mengirim Email via AI",
        "Sebelum mengkonfirmasi pengiriman email, WAJIB periksa:\n"
        "□ Penerima (To/CC/BCC): apakah benar akun uji?\n"
        "□ Subjek: apakah sesuai dengan tujuan?\n"
        "□ Lampiran: apakah file yang dilampirkan benar?\n"
        "□ Isi: apakah tidak mengandung data pribadi/rahasia?\n"
        "□ Konfirmasi: ketik 'Ya, kirim' HANYA setelah semua di atas diperiksa.",
        box_type="WARNING")
    add_page_break(doc)


    heading(doc, "D.3 Demo Google Drive dengan AI", level=2)
    body(doc, "[LANGKAH PRAKTIK] Demo Drive — Cari, Ringkas, dan Kelola File", bold=True, sz=10)
    langkah_drive = [
        "Pastikan app Google Drive sudah terhubung ke ChatGPT [FITUR TERGANTUNG AKUN/ADMIN].",
        "Cari file: 'Cari file dengan kata kunci cuti di folder DEMO_BKPSDM_AI. Tampilkan nama, tanggal, dan ringkasan.'",
        "Ringkas: 'Ringkas isi file PP_011_2017_ManajemenPNS_AKTIF.pdf dalam 5 poin utama.'",
        "Bandingkan: 'Bandingkan perbedaan utama antara PP_011_2017 dan PP_017_2020.'",
        "Buat file HANYA dalam DEMO_BKPSDM_AI: 'Buat Docs baru berisi template log konsultasi di folder DEMO_BKPSDM_AI.'",
        "Pratinjau sebelum simpan: tinjau isi, nama file, dan lokasi folder.",
        "Operasi file (pindah/salin): HANYA dalam DEMO_BKPSDM_AI; konfirmasi sebelum eksekusi.",
    ]
    for i, l in enumerate(langkah_drive, 1):
        bullet(doc, f"Langkah {i}: {l}")
    doc.add_paragraph()

    heading(doc, "D.4 Kemampuan ChatGPT & Batasannya", level=2)
    body(doc,
        "ChatGPT dapat menjalankan tindakan pada Gmail, Drive, Docs, Sheets, dan Slides "
        "termasuk membuat folder, spreadsheet, dan presentasi—HANYA jika app terkait tersedia, "
        "terhubung, dan sebatas izin pengguna, fitur app, paket, wilayah, perangkat, "
        "dan kebijakan administrator. [FITUR TERGANTUNG AKUN/ADMIN]")
    body(doc, "[TABEL] Perbedaan Mode Akses File di ChatGPT", bold=True, sz=10)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    for i, h in enumerate(["Mode", "Cara", "Baca/Cari", "Tulis/Tindakan"]):
        shd(tbl.cell(0, i), DARK_BLUE)
        ct(tbl.cell(0, i), h, bold=True, color=WHITE, sz=10)
    mode_data = [
        ("Unggah Manual", "Drag & drop file ke chat", "✅ Ya", "❌ Tidak"),
        ("App Baca/Cari", "App Drive/Gmail (baca saja)", "✅ Ya", "❌ Terbatas"),
        ("App Tindakan Tulis", "App Drive/Gmail (aksi lengkap)", "✅ Ya", "✅ Jika tersedia [FITUR TERGANTUNG AKUN/ADMIN]"),
        ("Apps Script/API", "Kode otomasi + API key", "✅ Ya", "✅ Dengan konfirmasi manusia"),
    ]
    for row_data in mode_data:
        row = tbl.add_row()
        for j, val in enumerate(row_data):
            ct(row.cells[j], val, sz=10)
    doc.add_paragraph()

    add_box(doc,
        "📌 PENTING: Perbedaan Langganan vs API",
        "• Langganan ChatGPT Plus/Pro BUKAN kredit API OpenAI.\n"
        "• OAuth (izin akun) BERBEDA dari API key.\n"
        "• Gemini Workspace TIDAK otomatis memberi akses Gemini API.\n"
        "• Untuk Apps Script + Gemini API, diperlukan API key terpisah dari Google AI Studio "
        "atau Vertex AI. [FITUR TERGANTUNG AKUN/ADMIN]",
        box_type="WARNING")
    add_page_break(doc)


    heading(doc, "D.5 Demo Apps Script + Gemini API", level=2)
    body(doc, "[LANGKAH PRAKTIK] Template Apps Script Siap Pakai", bold=True, sz=10)
    body(doc,
        "Template berikut tersedia di folder DEMO_BKPSDM_AI/template_kode_prompt_sop. "
        "Peserta tidak perlu mengetik dari nol—cukup menyalin template dan mengisi "
        "Script Properties dengan API key yang valid.")
    add_box(doc,
        "⚙ Struktur Sheet yang Diperlukan",
        "Sheet KONSULTASI: kolom Timestamp | Pertanyaan (Anonim) | Jawaban AI | Status Verifikasi\n"
        "Sheet REGULASI: kolom Nama File | Judul | Nomor | Tahun | Status | URL | Tgl Akses\n"
        "Sheet LOG: kolom Waktu | Fungsi | Input | Output | HTTP Status | Error\n"
        "Sheet UJI_MODEL: kolom Provider | Model | Pertanyaan Uji | Respons | Waktu (ms)",
        box_type="INFO")
    body(doc, "[KOTAK CONTOH] Template Kode Apps Script (Tersedia di Folder Demo)", bold=True, sz=10)
    add_box(doc,
        "💻 KODE: Template Apps Script — tanyaAsisten()",
        "// ─── KONFIGURASI (simpan di Script Properties, BUKAN di sini) ───\n"
        "// AI_PROVIDER  : 'gemini' atau 'openai'\n"
        "// AI_MODEL     : mis. 'gemini-1.5-flash' atau 'gpt-4o-mini'\n"
        "// AI_API_URL   : endpoint API\n"
        "// AI_API_KEY   : API key (JANGAN simpan di kode/GitHub)\n\n"
        "function onOpen() {\n"
        "  SpreadsheetApp.getUi().createMenu('Asisten AI')\n"
        "    .addItem('Tanya Asisten', 'tanyaAsisten')\n"
        "    .addItem('Uji Koneksi API', 'ujiKoneksiAPI')\n"
        "    .addToUi();\n"
        "}\n\n"
        "function tanyaAsisten() {\n"
        "  const sheet = SpreadsheetApp.getActiveSheet();\n"
        "  const lastRow = sheet.getLastRow();\n"
        "  const pertanyaan = sheet.getRange(lastRow, 2).getValue();\n"
        "  if (!pertanyaan) { Browser.msgBox('Kolom Pertanyaan kosong.'); return; }\n"
        "  const jawaban = callAI(pertanyaan);\n"
        "  sheet.getRange(lastRow, 3).setValue('DRAF—WAJIB DIVERIFIKASI:\\n' + jawaban);\n"
        "  catatLog('tanyaAsisten', pertanyaan, jawaban, 200, '');\n"
        "}\n\n"
        "function callAI(pertanyaan) {\n"
        "  const props = PropertiesService.getScriptProperties();\n"
        "  const provider = props.getProperty('AI_PROVIDER') || 'gemini';\n"
        "  if (provider === 'gemini') return callGemini(pertanyaan);\n"
        "  // Tambahkan adapter lain di sini untuk multi-provider\n"
        "  return '[ERROR: Provider tidak dikenal]';\n"
        "}\n\n"
        "function callGemini(pertanyaan) {\n"
        "  const props = PropertiesService.getScriptProperties();\n"
        "  const key = props.getProperty('AI_API_KEY');\n"
        "  const model = props.getProperty('AI_MODEL') || 'gemini-1.5-flash';\n"
        "  const url = props.getProperty('AI_API_URL') +\n"
        "    '/models/' + model + ':generateContent?key=' + key;\n"
        "  if (!key) return '[ERROR: API key belum diset di Script Properties]';\n"
        "  const systemPrompt = 'Kamu adalah asisten kepegawaian Indonesia. '\n"
        "    + 'JANGAN mengarang regulasi. '\n"
        "    + 'Jika tidak ada sumber, nyatakan tidak tahu. '\n"
        "    + 'DILARANG membuat keputusan kepegawaian. '\n"
        "    + 'Minta klarifikasi jika pertanyaan ambigu.';\n"
        "  const payload = JSON.stringify({\n"
        "    contents: [{role:'user', parts:[{text: systemPrompt + '\\n\\n' + pertanyaan}]}]\n"
        "  });\n"
        "  const options = {method:'post', contentType:'application/json', payload};\n"
        "  try {\n"
        "    const resp = UrlFetchApp.fetch(url, options);\n"
        "    if (resp.getResponseCode() !== 200) {\n"
        "      catatLog('callGemini', pertanyaan, '', resp.getResponseCode(), resp.getContentText());\n"
        "      return '[ERROR HTTP ' + resp.getResponseCode() + '] Cek LOG sheet.';\n"
        "    }\n"
        "    const json = JSON.parse(resp.getContentText());\n"
        "    return json?.candidates?.[0]?.content?.parts?.[0]?.text || '[Jawaban kosong]';\n"
        "  } catch(e) { return '[ERROR: ' + e.message + ']'; }\n"
        "}\n\n"
        "function ujiKoneksiAPI() {\n"
        "  const hasil = callAI('Apa itu cuti tahunan PNS? Jawab dalam satu kalimat.');\n"
        "  Browser.msgBox('Hasil Uji: ' + hasil);\n"
        "}\n\n"
        "function catatLog(fungsi, input, output, httpStatus, error) {\n"
        "  const log = SpreadsheetApp.getActiveSpreadsheet().getSheetByName('LOG');\n"
        "  if (!log) return;\n"
        "  log.appendRow([new Date(), fungsi, input.substring(0,200),\n"
        "    output.substring(0,200), httpStatus, error]);\n"
        "}",
        box_type="STEP")
    add_page_break(doc)


def add_section_assessment(doc):
    """Pretest/Post-test, Refleksi, Glosarium"""
    heading(doc, "PENILAIAN & REFLEKSI", level=1)

    heading(doc, "Pretest & Post-test (masing-masing 5 soal)", level=2)
    soal = [
        ("1", "Apa yang TIDAK boleh dilakukan AI dalam konteks kepegawaian?",
              "A) Menyusun draf jawaban\nB) Membuat keputusan kenaikan pangkat\nC) Merangkum email\nD) Mencari file regulasi",
              "B", "AI tidak memiliki kewenangan membuat keputusan kepegawaian resmi."),
        ("2", "Elemen 'Sumber' dalam kerangka prompt berarti:",
              "A) URL website\nB) Nama peraturan/regulasi yang dijadikan acuan\nC) Nama pengirim email\nD) Nama model AI",
              "B", "Sumber dalam prompt = regulasi acuan seperti PP 11/2017."),
        ("3", "Data manakah yang BOLEH dimasukkan ke prompt AI?",
              "A) NIP pegawai asli\nB) Kasus disiplin dengan nama asli\nC) Teks pertanyaan yang sudah dianonimkan\nD) Kata sandi sistem",
              "C", "Hanya data yang sudah dianonimkan yang boleh digunakan."),
        ("4", "Apa yang harus dilakukan pertama sebelum mengirim email via AI?",
              "A) Langsung klik kirim\nB) Memeriksa penerima, subjek, isi, lampiran\nC) Menutup ChatGPT\nD) Menghapus draft",
              "B", "Verifikasi penerima, subjek, isi, dan lampiran wajib dilakukan."),
        ("5", "Apa bedanya Gemini Workspace dan Gemini API?",
              "A) Sama saja\nB) Gemini Workspace otomatis memberi akses API\n"
              "C) Gemini Workspace adalah asisten dalam Google Workspace; API memerlukan key terpisah\n"
              "D) Gemini API gratis, Workspace berbayar",
              "C", "Gemini Workspace TIDAK otomatis memberi akses Gemini API."),
    ]
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    for i, h in enumerate(["No.", "Soal", "Pilihan", "Kunci", "Penjelasan"]):
        shd(tbl.cell(0, i), DARK_BLUE)
        ct(tbl.cell(0, i), h, bold=True, color=WHITE, sz=9)
    for no, soal_txt, pilihan, kunci, penjelasan in soal:
        row = tbl.add_row()
        ct(row.cells[0], no, sz=9, al=WD_ALIGN_PARAGRAPH.CENTER)
        ct(row.cells[1], soal_txt, sz=9)
        ct(row.cells[2], pilihan, sz=9)
        ct(row.cells[3], kunci, bold=True, sz=9, al=WD_ALIGN_PARAGRAPH.CENTER)
        ct(row.cells[4], penjelasan, sz=9)
    doc.add_paragraph()

    heading(doc, "Rubrik Penilaian", level=2)
    add_box(doc,
        "📊 Rubrik — Total 100 Poin, Batas Lulus 75",
        "• Pretest (5 soal × 5 poin)  = 25 poin\n"
        "• Post-test (5 soal × 5 poin) = 25 poin\n"
        "• Praktik prompt terstruktur  = 20 poin\n"
        "• Draf + verifikasi            = 20 poin\n"
        "• Rencana pilot 30 hari        = 10 poin\n"
        "─────────────────────────────────────────\n"
        "TOTAL                          = 100 poin\n"
        "Batas Lulus: ≥75 poin",
        box_type="INFO")
    doc.add_paragraph()

    heading(doc, "Glosarium", level=2)
    glosarium = [
        ("AI / Kecerdasan Buatan", "Teknologi komputer yang mampu melakukan tugas kognitif seperti memahami bahasa, merangkum teks, dan membuat draf."),
        ("Halusinasi AI", "Kondisi di mana AI menghasilkan informasi yang tampak meyakinkan tetapi tidak akurat atau tidak memiliki dasar."),
        ("Prompt", "Instruksi atau pertanyaan yang diberikan kepada AI untuk menghasilkan respons."),
        ("Apps Script", "Bahasa pemrograman berbasis JavaScript dari Google untuk mengotomasi tugas di Google Workspace."),
        ("API key", "Kode rahasia untuk mengakses layanan AI via program; JANGAN dibagikan atau disimpan di kode."),
        ("OAuth", "Protokol otorisasi yang memungkinkan aplikasi mengakses akun pengguna dengan izin terbatas."),
        ("JDIH", "Jaringan Dokumentasi dan Informasi Hukum; sumber resmi regulasi pemerintah Indonesia."),
        ("Anonimisasi", "Proses menghapus atau mengganti identitas pribadi dari data sebelum diproses lebih lanjut."),
        ("SKP", "Sasaran Kinerja Pegawai; dokumen penilaian kinerja PNS sesuai regulasi terbaru."),
        ("Draf AI", "Output awal yang dihasilkan AI; WAJIB diverifikasi sebelum digunakan sebagai dokumen resmi."),
    ]
    for term, defn in glosarium:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{term}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(defn)
        r2.font.size = Pt(10)
    add_page_break(doc)


def add_section_closing(doc):
    """Halaman 24: Penutup + Rencana Pilot 30 Hari"""
    heading(doc, "PENUTUP & RENCANA PILOT 30 HARI", level=1)

    heading(doc, "Rencana Pilot 30 Hari", level=2)
    body(doc,
        "Gunakan tabel berikut sebagai panduan implementasi AI di unit kerja Anda "
        "setelah pelatihan. Sesuaikan dengan kondisi dan kebutuhan aktual BKPSDM.")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    for i, h in enumerate(["Minggu", "Kegiatan", "Target"]):
        shd(tbl.cell(0, i), DARK_BLUE)
        ct(tbl.cell(0, i), h, bold=True, color=WHITE, sz=10)
    pilot = [
        ("Minggu 1–2\n(Hari 1–14)", "Siapkan folder DEMO_BKPSDM_AI; "
          "unggah 3–5 regulasi aktif; coba prompt terstruktur untuk 1 kasus cuti.",
          "1 prompt siap pakai; 1 draf terverifikasi tersimpan"),
        ("Minggu 3\n(Hari 15–21)", "Coba integrasi ChatGPT + Gmail (rangkum 5 email); "
          "buat log manual di Sheets.",
          "Log 5 konsultasi terdokumentasi"),
        ("Minggu 4\n(Hari 22–30)", "Uji template Apps Script dengan data fiktif; "
          "presentasi hasil pilot ke atasan; susun SOP sederhana.",
          "SOP draft 1 halaman; rencana tindak lanjut disepakati"),
    ]
    for w, k, t in pilot:
        row = tbl.add_row()
        ct(row.cells[0], w, bold=True, sz=10)
        ct(row.cells[1], k, sz=10)
        ct(row.cells[2], t, sz=10)
    doc.add_paragraph()

    heading(doc, "SOP Verifikasi Output AI (Ringkas)", level=2)
    sop = [
        "Terima draf AI dengan label 'DRAF—WAJIB DIVERIFIKASI'.",
        "Identifikasi setiap pasal/regulasi yang dikutip AI.",
        "Buka PDF regulasi dari JDIH (jdih.bkn.go.id atau jdih.demakkab.go.id).",
        "Cocokkan nomor pasal dan isi dengan kutipan AI.",
        "Jika ada ketidaksesuaian: HENTIKAN, perbaiki secara manual.",
        "Jika sesuai: tandai 'TERVERIFIKASI oleh [nama] pada [tanggal]'.",
        "Simpan hasil verifikasi di folder hasil_verifikasi (BUKAN di kode/GitHub).",
        "Catat di log Sheets: pertanyaan, jawaban AI, status verifikasi, verifikator.",
    ]
    for i, s in enumerate(sop, 1):
        bullet(doc, f"{i}. {s}")
    doc.add_paragraph()

    heading(doc, "Refleksi Pelatihan", level=2)
    refleksi = [
        "Hal paling berguna yang saya pelajari hari ini: ___________________________",
        "Satu hal yang masih ingin saya pelajari lebih lanjut: ______________________",
        "Langkah konkret pertama yang akan saya lakukan minggu depan: ______________",
        "Kekhawatiran terbesar saya tentang penggunaan AI di kantor: ________________",
        "Dukungan yang saya butuhkan dari atasan/instansi: __________________________",
    ]
    for r in refleksi:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(r)
        run.font.size = Pt(11)
    doc.add_paragraph()

    add_box(doc,
        "🙏 Terima Kasih & Selamat Berinovasi!",
        "Pelatihan ini adalah awal dari perjalanan transformasi digital BKPSDM Kabupaten Demak.\n"
        "Ingat selalu:\n"
        "  ✅ AI adalah ASISTEN, bukan pengganti penilaian profesional Anda.\n"
        "  ✅ Selalu VERIFIKASI setiap output AI sebelum digunakan.\n"
        "  ✅ JAGA kerahasiaan data ASN; jangan pernah memasukkan data asli ke AI.\n"
        "  ✅ Dokumentasikan, evaluasi, dan terus tingkatkan proses Anda.\n\n"
        "Pertanyaan & dukungan: [PERLU DATA: email/kontak panitia]\n"
        "Repositori kode & template: GitHub privat [PERLU DATA: URL repo privat]",
        box_type="INFO")


def add_footer(doc):
    """Add consistent footer to all sections (except cover)."""
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(
        "Modul Pelatihan AI – BKPSDM Kab. Demak  |  v1.0 – Juli 2026  |  "
        "DRAF—WAJIB DIVERIFIKASI  |  Hal. "
    )
    run.font.size = Pt(8)
    run.font.color.rgb = DARK_BLUE
    # Add auto page number
    fldChar = OxmlElement('w:fldChar')
    fldChar.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'end')
    rPr = footer_para.add_run()
    rPr._r.append(fldChar)
    rPr._r.append(instrText)
    rPr._r.append(fldChar2)
    rPr.font.size = Pt(8)
    rPr.font.color.rgb = DARK_BLUE


def add_toc(doc):
    """Insert a simple manual TOC heading (Word will auto-update on open)."""
    heading(doc, "DAFTAR ISI", level=1)
    body(doc,
        "Daftar Isi akan terisi otomatis ketika dokumen dibuka di Microsoft Word. "
        "Tekan Ctrl+A lalu F9, atau klik kanan → Update Field untuk memperbarui.", italic=True, sz=10)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    docPartObj = OxmlElement('w:docPartObj')
    docPartGallery = OxmlElement('w:docPartGallery')
    docPartGallery.set(qn('w:val'), 'Table of Contents')
    docPartObj.append(docPartGallery)
    sdtPr.append(docPartObj)
    sdt.append(sdtPr)
    sdtContent = OxmlElement('w:sdtContent')
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    p.append(r)
    r2 = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2.append(instrText)
    p.append(r2)
    r3 = OxmlElement('w:r')
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r3.append(fldChar2)
    p.append(r3)
    sdtContent.append(p)
    sdt.append(sdtContent)
    doc.element.body.append(sdt)
    add_page_break(doc)


def main():
    import os
    os.makedirs(OUT, exist_ok=True)
    doc = setup_document()
    add_cover(doc)
    add_toc(doc)
    add_frontmatter(doc)
    add_section_pembelajaran(doc)
    add_section_A(doc)
    add_section_B(doc)
    add_section_C(doc)
    add_section_D(doc)
    add_section_assessment(doc)
    add_section_closing(doc)
    add_footer(doc)
    doc.save(DOCX_PATH)
    print(f"✅ DOCX saved: {DOCX_PATH}")
    size_kb = os.path.getsize(DOCX_PATH) // 1024
    print(f"   File size: {size_kb} KB")


if __name__ == '__main__':
    main()
