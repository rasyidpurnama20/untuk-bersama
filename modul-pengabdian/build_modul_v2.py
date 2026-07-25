#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Build Script v2 – Modul Pelatihan AI untuk BKPSDM Kab. Demak
Target: ~70 halaman, padat contoh dan penjelasan
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Palette
DARK_BLUE  = (0x1A, 0x37, 0x6C)
MID_BLUE   = (0x27, 0x6F, 0xBF)
LIGHT_BLUE = (0xD6, 0xE8, 0xF7)
GOLD       = (0xC9, 0xA2, 0x27)
WHITE      = (0xFF, 0xFF, 0xFF)
WARN_BG    = (0xFF, 0xF3, 0xCD)
WARN_FG    = (0x85, 0x60, 0x04)
GREEN_BG   = (0xE8, 0xF5, 0xE9)
GREEN_FG   = (0x1B, 0x5E, 0x20)
STEP_BG    = (0xEF, 0xEF, 0xFF)

OUT = "/projects/sandbox/untuk-bersama/modul-pengabdian"
DOCX_PATH = os.path.join(OUT, "modul_pelatihan_v2.docx")

def rgb(c):
    return RGBColor(c[0], c[1], c[2])

def shd(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    tcPr.append(s)

def ct(cell, text, bold=False, color=None, sz=10, al=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = al
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(sz)
    r.font.name = 'Calibri'
    if color:
        r.font.color.rgb = rgb(color)



def h1(doc, text):
    p = doc.add_paragraph(style='Heading 1')
    r = p.add_run(text)
    r.font.color.rgb = rgb(DARK_BLUE)
    r.font.size = Pt(16)
    r.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    return p

def h2(doc, text):
    p = doc.add_paragraph(style='Heading 2')
    r = p.add_run(text)
    r.font.color.rgb = rgb(MID_BLUE)
    r.font.size = Pt(13)
    r.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    return p

def h3(doc, text):
    p = doc.add_paragraph(style='Heading 3')
    r = p.add_run(text)
    r.font.color.rgb = rgb(DARK_BLUE)
    r.font.size = Pt(11)
    r.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    return p

def para(doc, text, sz=11, bold=False, italic=False, al=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=6):
    p = doc.add_paragraph()
    p.alignment = al
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text)
    r.font.size = Pt(sz)
    r.font.name = 'Calibri'
    r.bold = bold
    r.italic = italic
    return p



def bul(doc, text, level=0, sz=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(sz)
    r.font.name = 'Calibri'
    return p

def num(doc, text, sz=11):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(sz)
    r.font.name = 'Calibri'
    return p

def box(doc, title, content, btype="INFO"):
    colors = {
        "INFO": (LIGHT_BLUE, DARK_BLUE),
        "WARNING": (WARN_BG, WARN_FG),
        "EXAMPLE": (GREEN_BG, GREEN_FG),
        "STEP": (STEP_BG, DARK_BLUE),
        "CODE": ((0xF5,0xF5,0xF5), (0x33,0x33,0x33)),
    }
    bg, fg = colors.get(btype, (LIGHT_BLUE, DARK_BLUE))
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.cell(0, 0)
    shd(hdr, fg)
    ct(hdr, f"  {title}", bold=True, color=WHITE, sz=10)
    bdy = tbl.cell(1, 0)
    shd(bdy, bg)
    bdy.text = ''
    p = bdy.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(content)
    r.font.size = Pt(10)
    r.font.name = 'Consolas' if btype == "CODE" else 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def pb(doc):
    doc.add_page_break()

def tbl_header(doc, headers, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        shd(tbl.cell(0, i), DARK_BLUE)
        ct(tbl.cell(0, i), h, bold=True, color=WHITE, sz=10, al=WD_ALIGN_PARAGRAPH.CENTER)
    return tbl

def tbl_row(tbl, values, sz=10):
    row = tbl.add_row()
    for j, val in enumerate(values):
        ct(row.cells[j], val, sz=sz)
    return row



def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.0)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    sec.different_first_page_header_footer = True
    return doc

def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("Modul Pelatihan AI – BKPSDM Kab. Demak | v2.0 – Juli 2026 | Hal. ")
        r.font.size = Pt(8)
        r.font.color.rgb = rgb(DARK_BLUE)
        fld1 = OxmlElement('w:fldChar')
        fld1.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.text = ' PAGE '
        fld2 = OxmlElement('w:fldChar')
        fld2.set(qn('w:fldCharType'), 'end')
        rr = fp.add_run()
        rr._r.append(fld1)
        rr._r.append(instr)
        rr._r.append(fld2)
        rr.font.size = Pt(8)

def cover(doc):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = tbl.cell(0, 0)
    shd(c, DARK_BLUE)
    ct(c, "  BKPSDM KABUPATEN DEMAK", bold=True, color=GOLD, sz=12)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[GAMBAR: Logo Kabupaten Demak + Logo BKPSDM]")
    r.font.size = Pt(10); r.italic = True; r.font.color.rgb = rgb(MID_BLUE)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MODUL PELATIHAN"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = rgb(DARK_BLUE)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Pemanfaatan Kecerdasan Buatan\nsebagai Asisten Tematik")
    r.bold = True; r.font.size = Pt(15); r.font.color.rgb = rgb(MID_BLUE)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("untuk Meningkatkan Kualitas Konsultasi dan Interpretasi\nRegulasi Kepegawaian pada BKPSDM Kabupaten Demak")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = rgb(DARK_BLUE)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("━" * 50); r.font.color.rgb = rgb(GOLD); r.font.size = Pt(12)
    doc.add_paragraph()
    meta = [
        ("Sasaran", "Pegawai BKPSDM Kabupaten Demak (Pemula–Menengah)"),
        ("Durasi", "2 Jam | 08.00 – 10.00 WIB"),
        ("Versi", "v2.0 – Juli 2026"),
        ("Penyelenggara", "Tim Pengabdian Masyarakat"),
    ]
    t = doc.add_table(rows=len(meta), cols=2)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        shd(t.cell(i, 0), DARK_BLUE)
        ct(t.cell(i, 0), k, bold=True, color=WHITE, sz=10)
        ct(t.cell(i, 1), v, sz=10)
    doc.add_paragraph()
    box(doc, "⚠ PENAFIAN",
        "Seluruh output AI bersifat DRAF dan BANTUAN. BUKAN keputusan hukum.\n"
        "WAJIB diverifikasi petugas berwenang sebelum digunakan.", "WARNING")
    pb(doc)

def toc(doc):
    h1(doc, "DAFTAR ISI")
    para(doc, "Daftar isi otomatis — tekan Ctrl+A lalu F9 di Microsoft Word untuk memperbarui.", italic=True, sz=10)
    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    dpObj = OxmlElement('w:docPartObj')
    dpGal = OxmlElement('w:docPartGallery')
    dpGal.set(qn('w:val'), 'Table of Contents')
    dpObj.append(dpGal); sdtPr.append(dpObj); sdt.append(sdtPr)
    sdtC = OxmlElement('w:sdtContent')
    p = OxmlElement('w:p')
    r1 = OxmlElement('w:r'); fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin'); r1.append(fc1); p.append(r1)
    r2 = OxmlElement('w:r'); it = OxmlElement('w:instrText'); it.text = ' TOC \\o "1-3" \\h \\z \\u '; r2.append(it); p.append(r2)
    r3 = OxmlElement('w:r'); fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end'); r3.append(fc2); p.append(r3)
    sdtC.append(p); sdt.append(sdtC)
    doc.element.body.append(sdt)
    pb(doc)



# ═══════════════════════════════════════════════════════════════════════════════
# BAGIAN A: DASAR PENGETAHUAN AI (target ~20 halaman = hal 3-22)
# ═══════════════════════════════════════════════════════════════════════════════

def section_A(doc):
    h1(doc, "BAGIAN A: Dasar Pengetahuan Kecerdasan Buatan")
    para(doc,
        "Bagian ini membahas fondasi kecerdasan buatan (AI) secara mendalam namun "
        "praktis. Anda akan memahami bagaimana AI berpikir, berkembang dari waktu ke "
        "waktu, dan bagaimana memilih serta menggunakan tools AI secara aman dan "
        "efektif di lingkungan kerja kepegawaian.")

    # ─── A.1 BAGAIMANA AI BERPIKIR ──────────────────────────────────────────
    h2(doc, "A.1 Bagaimana AI Berpikir dan Berkembang")

    h3(doc, "A.1.1 Apa Itu Kecerdasan Buatan?")
    para(doc,
        "Kecerdasan buatan (Artificial Intelligence/AI) adalah cabang ilmu komputer "
        "yang bertujuan menciptakan sistem mampu melakukan tugas yang biasanya "
        "memerlukan kecerdasan manusia: memahami bahasa, mengenali pola, membuat "
        "keputusan, dan belajar dari pengalaman.")
    para(doc,
        "Dalam konteks pekerjaan sehari-hari BKPSDM, AI yang kita gunakan termasuk "
        "kategori AI Generatif (Generative AI) — sistem yang mampu menghasilkan "
        "teks, gambar, kode, dan konten baru berdasarkan pola yang dipelajari dari "
        "data pelatihan yang sangat besar.")

    box(doc, "📌 Analogi Sederhana",
        "Bayangkan AI seperti pegawai magang yang:\n"
        "• Sudah membaca SELURUH perpustakaan hukum Indonesia (dan dunia)\n"
        "• Sangat cepat mencari dan merangkum informasi\n"
        "• TETAPI: tidak memahami konteks lokal BKPSDM Demak\n"
        "• TETAPI: bisa salah mengingat (halusinasi)\n"
        "• TETAPI: tidak punya kewenangan membuat keputusan\n\n"
        "Kesimpulan: AI adalah ASISTEN, bukan ATASAN dan bukan PENGGANTI.", "INFO")

    h3(doc, "A.1.2 Cara Kerja AI Generatif (LLM)")
    para(doc,
        "Large Language Model (LLM) seperti ChatGPT dan Gemini bekerja dengan prinsip "
        "prediksi kata berikutnya. Model dilatih dengan triliunan kata dari internet, "
        "buku, dan dokumen, kemudian belajar pola statistik bahasa.")
    para(doc, "Proses kerja AI generatif secara sederhana:", bold=True)
    num(doc, "Input: Pengguna mengetik pertanyaan (prompt) ke dalam sistem.")
    num(doc, "Tokenisasi: Teks dipecah menjadi token (potongan kata/sub-kata).")
    num(doc, "Pemrosesan: Model menghitung probabilitas token berikutnya berdasarkan konteks.")
    num(doc, "Generasi: Token-token disusun menjadi kalimat respons.")
    num(doc, "Output: Jawaban ditampilkan kepada pengguna.")

    box(doc, "⚠ PENTING: Mengapa AI Bisa 'Halusinasi'",
        "Karena AI bekerja berdasarkan PROBABILITAS, bukan PEMAHAMAN:\n\n"
        "• AI tidak 'tahu' apakah PP 11/2017 benar ada — ia memprediksi bahwa pola\n"
        "  teks 'PP 11/2017' kemungkinan besar muncul dalam konteks kepegawaian\n"
        "• AI bisa menghasilkan nomor pasal yang TIDAK ADA karena polanya 'masuk akal'\n"
        "• AI tidak memiliki akses ke database hukum real-time (kecuali fitur web search)\n"
        "• AI tidak dapat memverifikasi sendiri apakah jawabannya benar\n\n"
        "IMPLIKASI: Setiap output AI WAJIB diverifikasi oleh manusia menggunakan\n"
        "sumber resmi (JDIH BKN, JDIH Kemenpan-RB, dll).", "WARNING")

    h3(doc, "A.1.3 Evolusi AI: Dari Rule-Based ke Generative")
    tbl = tbl_header(doc, ["Era", "Teknologi", "Kemampuan", "Contoh"])
    data = [
        ("1950–1990", "Rule-Based / Expert System", "Mengikuti aturan IF-THEN yang diprogram manusia", "MYCIN (diagnosis medis)"),
        ("1990–2010", "Machine Learning", "Belajar pola dari data terstruktur", "Spam filter, rekomendasi Netflix"),
        ("2010–2020", "Deep Learning", "Memproses data tidak terstruktur (gambar, suara)", "Google Translate, Siri"),
        ("2020–kini", "Generative AI / LLM", "Menghasilkan teks, kode, gambar baru", "ChatGPT, Gemini, Claude, Copilot"),
        ("2024–kini", "AI Agents & Tool Use", "Mengambil tindakan: baca email, kelola file, jalankan kode", "ChatGPT Pro + Apps, Google Gemini"),
    ]
    for d in data:
        tbl_row(tbl, d)
    doc.add_paragraph()

    h3(doc, "A.1.4 Kemampuan vs Keterbatasan AI untuk Pekerjaan ASN")
    para(doc, "Memahami batas kemampuan AI sama pentingnya dengan memahami manfaatnya:")
    tbl = tbl_header(doc, ["Aspek", "✅ Kemampuan AI", "❌ Keterbatasan AI"])
    rows = [
        ("Bahasa", "Memahami dan menghasilkan teks Indonesia/Inggris", "Tidak memahami dialek lokal atau konteks budaya spesifik"),
        ("Regulasi", "Merangkum dan menjelaskan isi peraturan yang diunggah", "Tidak memiliki akses real-time ke JDIH; bisa salah pasal"),
        ("Dokumen", "Membuat draf, template, ringkasan", "Tidak dapat menandatangani atau mengesahkan dokumen resmi"),
        ("Data", "Menganalisis data yang diberikan", "Tidak boleh diberi data pribadi ASN (NIP, nama asli, dll)"),
        ("Keputusan", "Menyusun opsi dan pertimbangan", "DILARANG membuat keputusan kepegawaian"),
        ("Waktu", "Memproses dalam hitungan detik", "Data pelatihan memiliki cut-off date; regulasi baru mungkin belum tercakup"),
        ("Konsistensi", "Dapat mengikuti format yang diminta", "Jawaban bisa berbeda untuk pertanyaan yang sama (non-deterministik)"),
        ("Konteks lokal", "Dapat diarahkan dengan konteks di prompt", "Tidak mengetahui kondisi spesifik BKPSDM Demak secara bawaan"),
    ]
    for r in rows:
        tbl_row(tbl, r)
    doc.add_paragraph()
    pb(doc)

    # ─── A.2 MEMILIH TOOLS AI ────────────────────────────────────────────────
    h2(doc, "A.2 Memilih Tools AI yang Tepat")
    para(doc,
        "Saat ini tersedia puluhan tools AI dengan berbagai kemampuan. Pemilihan tools "
        "yang tepat bergantung pada kebutuhan kerja, anggaran, dan kebijakan instansi. "
        "Berikut panduan komprehensif untuk memilih tools AI di lingkungan BKPSDM.")

    h3(doc, "A.2.1 Kriteria Pemilihan Tools AI untuk ASN")
    para(doc, "Pertimbangkan faktor-faktor berikut sebelum memilih tools AI:")
    criteria = [
        "Keamanan Data: Apakah data diproses di server luar negeri? Apakah ada kebijakan retensi data?",
        "Ketersediaan: Apakah dapat diakses dari jaringan kantor? Apakah memerlukan VPN?",
        "Biaya: Apakah gratis atau berbayar? Apakah ada batas penggunaan (quota)?",
        "Kemudahan: Apakah mudah digunakan tanpa pelatihan teknis mendalam?",
        "Integrasi: Apakah terhubung dengan Google Workspace atau tools kerja lainnya?",
        "Kebijakan Instansi: Apakah sudah disetujui oleh admin IT instansi?",
        "Bahasa: Apakah mendukung Bahasa Indonesia dengan baik?",
        "Fitur: Apakah mendukung upload file, pencarian web, atau tindakan otomatis?",
    ]
    for c in criteria:
        bul(doc, c)


    h3(doc, "A.2.2 Komparasi Tools AI Populer")
    tbl = tbl_header(doc, ["Tools", "Tipe", "Harga", "Kelebihan", "Kekurangan", "Cocok Untuk"])
    tools = [
        ("ChatGPT Free", "Chat AI", "Gratis", "Mudah digunakan; mendukung file upload", "Batas penggunaan; tidak ada Apps", "Pertanyaan ad hoc sederhana"),
        ("ChatGPT Plus", "Chat AI + Apps", "$20/bln", "Projects, Apps (Gmail/Drive), web search", "Berbayar; fitur tergantung wilayah", "Power user; integrasi workspace"),
        ("ChatGPT Pro", "Chat AI + Full", "$200/bln", "Unlimited; o1-pro; semua fitur", "Mahal; belum tentu perlu untuk ASN", "Riset mendalam; coding intensif"),
        ("Google Gemini", "Chat AI", "Gratis/Workspace", "Terintegrasi Gmail/Drive/Docs", "Perlu admin aktifkan; fitur terbatas", "Pengguna Google Workspace"),
        ("Gemini API", "API", "Pay-per-use", "Otomasi via Apps Script; fleksibel", "Perlu API key; teknis", "Otomasi proses berulang"),
        ("Claude", "Chat AI", "Free/$20/bln", "Konteks panjang; analisis dokumen besar", "Tidak ada integrasi Google", "Analisis dokumen hukum panjang"),
        ("Microsoft Copilot", "Chat + Office", "Termasuk M365", "Terintegrasi Word/Excel/Outlook", "Perlu lisensi M365; privasi Microsoft", "Instansi yang pakai Microsoft 365"),
        ("Perplexity", "Pencarian AI", "Free/$20/bln", "Pencarian web + sitasi sumber", "Tidak untuk tindakan/otomasi", "Riset regulasi dengan sumber URL"),
    ]
    for t in tools:
        tbl_row(tbl, t, sz=9)
    doc.add_paragraph()

    box(doc, "💡 Rekomendasi untuk BKPSDM Kabupaten Demak",
        "Untuk pelatihan ini, kami menggunakan:\n"
        "• ChatGPT (Free/Plus/Pro) — sebagai AI utama untuk demonstrasi\n"
        "• Google Workspace (Gmail, Drive, Sheets, Docs) — ekosistem kerja\n"
        "• Apps Script + Gemini API — untuk otomasi (Bagian C)\n"
        "• GitHub (privat) — untuk version control kode dan prompt\n\n"
        "Pilihan ini berdasarkan:\n"
        "✅ Kemudahan akses (browser-based)\n"
        "✅ Integrasi langsung dengan Google Workspace\n"
        "✅ Tersedia versi gratis untuk uji coba\n"
        "✅ Dokumentasi berbahasa Indonesia tersedia", "INFO")
    pb(doc)

    # ─── A.3 MENGHUBUNGKAN APP SECARA AMAN ────────────────────────────────────
    h2(doc, "A.3 Menghubungkan Aplikasi AI Secara Aman")
    para(doc,
        "Salah satu keunggulan AI modern adalah kemampuannya terhubung dengan "
        "aplikasi lain (Gmail, Drive, Sheets, dll). Namun koneksi ini membawa "
        "risiko keamanan yang harus dipahami dan dikelola dengan baik.")

    h3(doc, "A.3.1 Memahami OAuth (Open Authorization)")
    para(doc,
        "OAuth adalah protokol otorisasi standar yang memungkinkan aplikasi pihak "
        "ketiga (seperti ChatGPT) mengakses akun Google Anda TANPA mengetahui "
        "password Anda. Prosesnya:")
    num(doc, "Anda memilih 'Hubungkan Gmail' di ChatGPT.")
    num(doc, "ChatGPT mengarahkan Anda ke halaman login Google.")
    num(doc, "Google menampilkan daftar izin yang diminta (baca email, kelola file, dll).")
    num(doc, "Anda menyetujui atau menolak izin tersebut.")
    num(doc, "Google memberikan 'token' ke ChatGPT (bukan password Anda).")
    num(doc, "ChatGPT menggunakan token untuk mengakses sesuai izin yang diberikan.")
    num(doc, "Anda dapat mencabut izin kapan saja di myaccount.google.com > Keamanan > Akses pihak ketiga.")

    box(doc, "⚠ PERINGATAN: Perbedaan OAuth vs API Key",
        "OAuth (izin akun):\n"
        "• Mengizinkan aplikasi bertindak ATAS NAMA Anda\n"
        "• Anda login dan menyetujui izin\n"
        "• Bisa dicabut kapan saja\n"
        "• Digunakan: ChatGPT Apps, Gemini Workspace\n\n"
        "API Key (kunci program):\n"
        "• Kode rahasia untuk akses PROGRAMATIK ke layanan AI\n"
        "• Tidak ada login pengguna; langsung dari kode\n"
        "• JANGAN simpan di kode/GitHub/shared drive\n"
        "• Simpan di Script Properties atau Secret Manager\n"
        "• Digunakan: Apps Script → Gemini API\n\n"
        "PENTING: Langganan ChatGPT Plus/Pro BUKAN kredit API.\n"
        "Gemini Workspace BUKAN sama dengan akses Gemini API.", "WARNING")

    h3(doc, "A.3.2 Prinsip Keamanan Koneksi AI")
    para(doc, "Ikuti prinsip-prinsip berikut saat menghubungkan AI ke aplikasi kerja:")
    tbl = tbl_header(doc, ["Prinsip", "Penjelasan", "Contoh Penerapan"])
    prinsip = [
        ("Least Privilege", "Berikan izin MINIMUM yang diperlukan", "Jika hanya perlu baca email, jangan beri izin kirim"),
        ("Akun Demo Terpisah", "Gunakan akun khusus untuk uji coba", "akun.demo.bkpsdm@gmail.com — bukan akun pribadi/dinas"),
        ("Review Berkala", "Periksa izin yang sudah diberikan", "Cek myaccount.google.com/permissions setiap bulan"),
        ("Data Fiktif", "Gunakan data fiktif untuk latihan", "Nama: Pegawai Fiktif A; NIP: 000000000000"),
        ("Jangan Simpan Rahasia", "API key, password tidak boleh di kode", "Gunakan Script Properties di Apps Script"),
        ("Konfirmasi Sebelum Tindakan", "Selalu preview sebelum AI mengirim/membuat", "Baca draf email sebelum klik 'Kirim'"),
        ("Audit Trail", "Catat setiap tindakan AI di log", "Sheet LOG: waktu, fungsi, input, output"),
        ("Cabut Jika Tidak Dipakai", "Hapus izin app yang sudah tidak digunakan", "Cabut akses ChatGPT jika tidak lagi dipakai"),
    ]
    for p in prinsip:
        tbl_row(tbl, p)
    doc.add_paragraph()

    h3(doc, "A.3.3 Data yang DILARANG Diberikan ke AI")
    box(doc, "🚫 LARANGAN ABSOLUT — Data yang TIDAK BOLEH Masuk ke Prompt AI",
        "Kategori IDENTITAS PRIBADI:\n"
        "• NIP (Nomor Induk Pegawai) asli\n"
        "• Nama lengkap pegawai asli\n"
        "• Alamat rumah, nomor telepon, email pribadi\n"
        "• Foto KTP, KK, atau dokumen identitas\n\n"
        "Kategori DATA SENSITIF:\n"
        "• Data kesehatan / catatan medis individual\n"
        "• Kasus disiplin aktif dengan identitas nyata\n"
        "• Hasil penilaian kinerja individual\n"
        "• Data gaji, tunjangan, atau keuangan pegawai\n\n"
        "Kategori KEAMANAN SISTEM:\n"
        "• Kata sandi, OTP, PIN\n"
        "• Token autentikasi, API key\n"
        "• Screenshot sistem SIMPEG/BKN internal\n"
        "• Dokumen berklasifikasi Rahasia/Sangat Rahasia\n\n"
        "GUNAKAN SELALU: Data fiktif, data publik, atau data yang sudah dianonimkan.", "WARNING")
    pb(doc)


    # ─── A.4 MENYUSUN PROMPT ─────────────────────────────────────────────────
    h2(doc, "A.4 Menyusun Prompt yang Efektif (Prompt Engineering)")
    para(doc,
        "Prompt adalah instruksi yang Anda berikan ke AI. Kualitas output AI sangat "
        "bergantung pada kualitas prompt. Bagian ini mengajarkan teknik menyusun "
        "prompt yang menghasilkan jawaban akurat, terstruktur, dan dapat ditelusuri "
        "sumbernya.")

    h3(doc, "A.4.1 Kerangka Prompt P-T-O-K-S-B-F-K")
    para(doc,
        "Gunakan kerangka 8 elemen berikut untuk menyusun prompt terstruktur. "
        "Tidak semua elemen wajib ada di setiap prompt, tetapi semakin lengkap "
        "elemen yang Anda berikan, semakin baik hasilnya.")
    tbl = tbl_header(doc, ["Elemen", "Singkatan", "Pertanyaan Panduan", "Contoh"])
    elemen = [
        ("Peran", "P", "Siapa AI dalam konteks ini?", "Kamu adalah asisten kepegawaian PNS Indonesia..."),
        ("Tujuan", "T", "Apa yang ingin dicapai?", "Tolong jelaskan syarat dan prosedur..."),
        ("Objek", "O", "Apa subjek utamanya?", "...cuti besar bagi PNS golongan III"),
        ("Konteks", "K", "Apa latar belakangnya?", "...dengan masa kerja 6 tahun berturut-turut"),
        ("Sumber", "S", "Regulasi apa yang jadi acuan?", "...berdasarkan PP 11/2017 dan PP 17/2020"),
        ("Batasan", "B", "Apa yang TIDAK boleh dilakukan?", "Jangan mengarang pasal yang tidak ada"),
        ("Format", "F", "Bagaimana bentuk output?", "Tampilkan dalam poin-poin + nomor pasal"),
        ("Konfirmasi", "K", "Kapan harus bertanya?", "Jika ada yang ambigu, tanyakan dulu"),
    ]
    for e in elemen:
        tbl_row(tbl, e)
    doc.add_paragraph()

    h3(doc, "A.4.2 Contoh Prompt Lengkap untuk Konsultasi Kepegawaian")
    box(doc, "✅ CONTOH PROMPT LENGKAP (Kasus Cuti Besar)",
        "[PERAN] Kamu adalah asisten kepegawaian untuk PNS Indonesia yang bekerja\n"
        "di BKPSDM. Kamu membantu menyusun draf jawaban konsultasi.\n\n"
        "[TUJUAN] Tolong jelaskan syarat, prosedur, dan dokumen yang diperlukan\n"
        "untuk mengajukan cuti besar.\n\n"
        "[OBJEK] Cuti besar bagi PNS golongan III/b.\n\n"
        "[KONTEKS] Pegawai memiliki masa kerja 7 tahun berturut-turut dan belum\n"
        "pernah mengambil cuti besar sebelumnya.\n\n"
        "[SUMBER] Gunakan PP Nomor 11 Tahun 2017 tentang Manajemen PNS (Pasal\n"
        "tentang cuti) dan PP Nomor 17 Tahun 2020 sebagai perubahan.\n"
        "Jika kamu tidak yakin pasal mana, nyatakan 'perlu verifikasi'.\n\n"
        "[BATASAN]\n"
        "• Jangan mengarang nomor pasal yang tidak kamu yakin ada\n"
        "• Jangan memberikan keputusan — hanya draf jawaban\n"
        "• Jangan menyebut nama/NIP pegawai (gunakan 'Pegawai Fiktif A')\n\n"
        "[FORMAT]\n"
        "• Mulai dengan: DRAF—WAJIB DIVERIFIKASI PETUGAS BERWENANG\n"
        "• Ringkasan (2-3 kalimat)\n"
        "• Dasar hukum (peraturan + pasal + status)\n"
        "• Syarat (poin-poin)\n"
        "• Prosedur (langkah berurutan)\n"
        "• Dokumen yang diperlukan\n\n"
        "[KONFIRMASI] Jika ada informasi yang kurang jelas atau kamu butuh\n"
        "klarifikasi, tanyakan sebelum menjawab.", "EXAMPLE")


    h3(doc, "A.4.3 Contoh Prompt BURUK vs Prompt BAIK")
    tbl = tbl_header(doc, ["Aspek", "❌ Prompt Buruk", "Risiko", "✅ Prompt Baik"])
    prompt_compare = [
        ("Terlalu pendek", "'Apa itu cuti?'", "Jawaban umum Wikipedia, bukan konteks PNS", "'Jelaskan jenis-jenis cuti PNS berdasarkan PP 11/2017'"),
        ("Tanpa sumber", "'Bagaimana prosedur kenaikan pangkat?'", "AI bisa mengarang prosedur yang tidak sesuai regulasi", "'Jelaskan prosedur kenaikan pangkat reguler PNS berdasarkan PP 11/2017 Pasal 68-79'"),
        ("Tanpa batasan", "'Buat surat keputusan cuti'", "AI membuat SK yang bisa disalahgunakan", "'Buat DRAF template surat pengajuan cuti (bukan SK). Tandai semua field yang perlu diisi manual.'"),
        ("Data asli", "'NIP 123456 mau cuti, gimana?'", "Data ASN bocor ke server AI", "'Pegawai Fiktif A (Gol III/b, 7 thn) ingin mengajukan cuti besar. Jelaskan syaratnya.'"),
        ("Minta keputusan", "'Setujui atau tolak pengajuan ini'", "AI tidak berwenang membuat keputusan", "'Susun pertimbangan pro dan kontra, lalu serahkan keputusan ke atasan berwenang.'"),
    ]
    for p in prompt_compare:
        tbl_row(tbl, p, sz=9)
    doc.add_paragraph()

    h3(doc, "A.4.4 Teknik Prompt Lanjutan")
    para(doc, "Beberapa teknik lanjutan untuk meningkatkan kualitas output AI:")

    para(doc, "1. Chain of Thought (Berpikir Bertahap)", bold=True)
    box(doc, "📌 Contoh: Chain of Thought",
        "Prompt: 'Jelaskan langkah demi langkah bagaimana menentukan apakah\n"
        "seorang PNS berhak atas cuti besar. Pikirkan secara bertahap:\n"
        "1. Periksa masa kerja\n"
        "2. Periksa apakah sudah pernah ambil cuti besar\n"
        "3. Periksa kondisi kepegawaian aktif\n"
        "4. Periksa regulasi yang berlaku\n"
        "Untuk setiap langkah, sebutkan pasal regulasi yang relevan.'", "EXAMPLE")

    para(doc, "2. Few-Shot Learning (Beri Contoh)", bold=True)
    box(doc, "📌 Contoh: Few-Shot Learning",
        "Prompt: 'Buat ringkasan regulasi dengan format berikut:\n\n"
        "CONTOH:\n"
        "Regulasi: PP 53/2010 tentang Disiplin PNS\n"
        "Status: DICABUT (diganti PP 94/2021)\n"
        "Poin utama: Kewajiban dan larangan PNS, jenis hukuman disiplin\n"
        "Catatan: Gunakan PP 94/2021 sebagai acuan terkini\n\n"
        "Sekarang buatkan ringkasan serupa untuk: PP 11/2017 tentang Manajemen PNS'", "EXAMPLE")

    para(doc, "3. Role-Playing dengan Constraints", bold=True)
    box(doc, "📌 Contoh: Role-Playing",
        "Prompt: 'Kamu adalah petugas BKPSDM yang sedang menjawab pertanyaan\n"
        "pegawai via email. Aturan kamu:\n"
        "- Selalu awali dengan salam formal\n"
        "- Cantumkan minimal satu regulasi sebagai dasar\n"
        "- Jika tidak yakin, katakan \"perlu kami konfirmasi lebih lanjut\"\n"
        "- Akhiri dengan \"Demikian informasi awal, untuk kepastian mohon\n"
        "  konsultasi langsung ke bidang terkait.\"\n\n"
        "Pertanyaan pegawai: Apakah saya bisa mengambil cuti menikah?'", "EXAMPLE")
    pb(doc)


    # ─── A.5 MENGHASILKAN DRAF BERBASIS SUMBER ─────────────────────────────────
    h2(doc, "A.5 Menghasilkan Draf Berbasis Sumber Regulasi")
    para(doc,
        "Salah satu penggunaan AI paling berharga untuk BKPSDM adalah menghasilkan "
        "draf jawaban konsultasi yang mengacu pada sumber regulasi yang jelas. "
        "Bagian ini mengajarkan alur lengkap dari pertanyaan hingga draf terverifikasi.")

    h3(doc, "A.5.1 Alur Kerja: Pertanyaan → Draf AI → Verifikasi")
    box(doc, "📊 DIAGRAM: Alur Pembuatan Draf dengan AI",
        "┌─────────────────────────────────────────────────────┐\n"
        "│ 1. PERTANYAAN MASUK (email/tatap muka/telepon)      │\n"
        "└───────────────────────┬─────────────────────────────┘\n"
        "                        ▼\n"
        "┌─────────────────────────────────────────────────────┐\n"
        "│ 2. ANONIMISASI                                      │\n"
        "│    Hapus: NIP, nama, kontak → ganti data fiktif     │\n"
        "└───────────────────────┬─────────────────────────────┘\n"
        "                        ▼\n"
        "┌─────────────────────────────────────────────────────┐\n"
        "│ 3. PILIH REGULASI ACUAN                             │\n"
        "│    Buka folder regulasi → pilih PP/PerMen terkait   │\n"
        "└───────────────────────┬─────────────────────────────┘\n"
        "                        ▼\n"
        "┌─────────────────────────────────────────────────────┐\n"
        "│ 4. SUSUN PROMPT (P-T-O-K-S-B-F-K)                  │\n"
        "│    Lengkapi 8 elemen → kirim ke AI                  │\n"
        "└───────────────────────┬─────────────────────────────┘\n"
        "                        ▼\n"
        "┌─────────────────────────────────────────────────────┐\n"
        "│ 5. TERIMA DRAF AI                                   │\n"
        "│    Label: DRAF—WAJIB DIVERIFIKASI PETUGAS BERWENANG │\n"
        "└───────────────────────┬─────────────────────────────┘\n"
        "                        ▼\n"
        "┌─────────────────────────────────────────────────────┐\n"
        "│ 6. VERIFIKASI MANUAL                                │\n"
        "│    Buka JDIH → cocokkan pasal → cek status regulasi │\n"
        "└───────────────────────┬─────────────────────────────┘\n"
        "                        ▼\n"
        "┌─────────────────────────────────────────────────────┐\n"
        "│ 7. JAWABAN FINAL + LOG                              │\n"
        "│    Kirim jawaban terverifikasi → catat di Sheets    │\n"
        "└─────────────────────────────────────────────────────┘", "INFO")


    h3(doc, "A.5.2 Format Standar Draf Output AI")
    box(doc, "📋 TEMPLATE: Format Draf Output AI",
        "═══════════════════════════════════════════════════════════════\n"
        "DRAF — WAJIB DIVERIFIKASI PETUGAS BERWENANG\n"
        "═══════════════════════════════════════════════════════════════\n\n"
        "RINGKASAN:\n"
        "[Jawaban singkat 2-3 kalimat yang menjawab inti pertanyaan]\n\n"
        "DASAR HUKUM:\n"
        "1. [Nama Peraturan] Nomor [X] Tahun [Y]\n"
        "   Pasal [Z] ayat [(n)]\n"
        "   Status: AKTIF / DIUBAH / DICABUT\n"
        "   Sumber: [URL JDIH jika tersedia]\n"
        "   Tanggal akses: [dd/mm/yyyy]\n\n"
        "PROSEDUR:\n"
        "1. [Langkah pertama]\n"
        "2. [Langkah kedua]\n"
        "3. ...\n\n"
        "DOKUMEN YANG DIPERLUKAN:\n"
        "• [Dokumen 1]\n"
        "• [Dokumen 2]\n\n"
        "KLARIFIKASI DIPERLUKAN:\n"
        "• [Hal yang perlu ditanyakan kembali jika ada]\n\n"
        "CATATAN:\n"
        "• Jawaban ini bersifat DRAF dan PERLU DIVERIFIKASI\n"
        "• [Hal penting lainnya]\n\n"
        "─────────────────────────────────────────────────────────────\n"
        "VERIFIKASI:\n"
        "Verifikator: __________________ Jabatan: __________________\n"
        "Tanggal: __________________     Status: □ VALID  □ REVISI\n"
        "Catatan: __________________________________________________\n"
        "═══════════════════════════════════════════════════════════════", "STEP")

    h3(doc, "A.5.3 Pengelolaan Regulasi Acuan")
    para(doc,
        "Untuk memastikan konsistensi dan kemudahan pencarian, gunakan konvensi "
        "penamaan standar untuk file regulasi:")
    box(doc, "📁 Konvensi Penamaan File Regulasi",
        "Format: JENIS_NOMOR_TAHUN_TOPIK_STATUS.pdf\n\n"
        "Contoh:\n"
        "  PP_011_2017_ManajemenPNS_AKTIF.pdf\n"
        "  PP_017_2020_PerubahanManajemenPNS_AKTIF.pdf\n"
        "  PP_053_2010_DisiplinPNS_DICABUT.pdf\n"
        "  PP_094_2021_DisiplinPNS_AKTIF.pdf\n"
        "  PerBKN_006_2022_PenilaianKinerja_AKTIF.pdf\n"
        "  PerMenpanRB_001_2023_JabatanFungsional_PERLU_VERIFIKASI.pdf\n\n"
        "Status yang digunakan:\n"
        "  AKTIF            = Berlaku penuh\n"
        "  DIUBAH           = Masih berlaku tapi ada pasal yang berubah\n"
        "  DICABUT          = Sudah tidak berlaku\n"
        "  PERLU_VERIFIKASI = Status belum dikonfirmasi dari JDIH", "INFO")
    pb(doc)


    # ─── A.6 MEMVERIFIKASI HASIL ──────────────────────────────────────────────
    h2(doc, "A.6 Memverifikasi Hasil AI")
    para(doc,
        "Verifikasi adalah langkah PALING PENTING dalam penggunaan AI. Tanpa "
        "verifikasi, output AI berpotensi menyesatkan dan menimbulkan masalah "
        "hukum. Bagian ini menyajikan metodologi verifikasi yang sistematis.")

    h3(doc, "A.6.1 Mengapa Verifikasi Wajib?")
    tbl = tbl_header(doc, ["Jenis Kesalahan AI", "Contoh", "Dampak Jika Tidak Diverifikasi"])
    verif_data = [
        ("Halusinasi fakta", "AI menyebut 'PP 11/2017 Pasal 200' — padahal PP hanya punya 183 pasal", "Jawaban resmi mengacu pasal yang tidak ada"),
        ("Regulasi kedaluwarsa", "AI mengacu PP 53/2010 (disiplin PNS) yang sudah dicabut", "Pegawai mendapat informasi berdasarkan aturan lama"),
        ("Salah konteks", "AI menjawab berdasarkan regulasi negara lain", "Prosedur yang diberikan tidak sesuai hukum Indonesia"),
        ("Bias interpretasi", "AI menyederhanakan pasal yang sebenarnya memiliki pengecualian", "Pegawai tidak mengetahui hak/kewajiban lengkap"),
        ("Campuran fakta-fiksi", "Jawaban 80% benar, 20% salah — sulit dibedakan", "Kesalahan lolos karena sebagian besar terlihat benar"),
    ]
    for v in verif_data:
        tbl_row(tbl, v)
    doc.add_paragraph()

    h3(doc, "A.6.2 SOP Verifikasi Output AI (7 Langkah)")
    num(doc, "BACA UTUH: Baca seluruh output AI. Tandai setiap klaim faktual dan pasal yang disebut.")
    num(doc, "IDENTIFIKASI SUMBER: Catat semua regulasi/pasal yang dikutip AI.")
    num(doc, "BUKA JDIH: Akses jdih.bkn.go.id atau jdih.demakkab.go.id — cari peraturan yang dikutip.")
    num(doc, "COCOKKAN PASAL: Buka PDF regulasi asli → cocokkan isi pasal dengan kutipan AI.")
    num(doc, "CEK STATUS: Pastikan regulasi masih AKTIF (belum dicabut atau diubah).")
    num(doc, "VALIDASI ATAU TOLAK: Jika cocok → tandai VALID. Jika ada ketidaksesuaian → REVISI manual.")
    num(doc, "LOG HASIL: Catat di Sheet LOG: pertanyaan, jawaban AI, status verifikasi, nama verifikator, tanggal.")

    box(doc, "🔍 Checklist Verifikasi Cepat",
        "□ Apakah nomor peraturan (PP/PerMen/PerBKN) ada di JDIH?\n"
        "□ Apakah nomor pasal yang dikutip AI benar-benar ada?\n"
        "□ Apakah ISI pasal sesuai dengan apa yang dikatakan AI?\n"
        "□ Apakah peraturan tersebut masih AKTIF (belum dicabut)?\n"
        "□ Apakah ada peraturan PERUBAHANNYA yang lebih baru?\n"
        "□ Apakah ada pengecualian atau ketentuan peralihan?\n"
        "□ Apakah jawaban sudah ditandai DRAF sebelum dikirim?\n"
        "□ Apakah sudah tercatat di log audit?", "STEP")
    pb(doc)


    h3(doc, "A.6.3 Sumber Verifikasi Resmi")
    tbl = tbl_header(doc, ["Sumber", "URL", "Isi", "Catatan"])
    sumber = [
        ("JDIH BKN", "jdih.bkn.go.id", "Regulasi kepegawaian nasional (PP, PerBKN)", "Sumber utama untuk aturan ASN"),
        ("JDIH Kemenpan-RB", "jdih.menpan.go.id", "Peraturan MenPAN-RB (jabatan, kinerja)", "SKP, jabatan fungsional"),
        ("JDIH Nasional", "jdih.setneg.go.id", "Semua peraturan perundang-undangan RI", "Database komprehensif"),
        ("JDIH Kab. Demak", "jdih.demakkab.go.id", "Peraturan daerah Kabupaten Demak", "Perda dan Perbup terkait ASN lokal"),
        ("BKN Portal", "bkn.go.id", "Informasi resmi kepegawaian", "Pengumuman, SE, surat edaran"),
    ]
    for s in sumber:
        tbl_row(tbl, s)
    doc.add_paragraph()
    para(doc, "[PERLU VERIFIKASI SUMBER RESMI: Pastikan URL di atas masih aktif pada saat pelatihan dilaksanakan]", italic=True, sz=10)

    h3(doc, "A.6.4 Tabel Risiko & Pengendalian")
    tbl = tbl_header(doc, ["No.", "Risiko", "Probabilitas", "Dampak", "Pengendalian"])
    risiko = [
        ("1", "Halusinasi pasal", "Tinggi", "Tinggi", "Verifikasi SETIAP pasal ke JDIH; label DRAF wajib"),
        ("2", "Regulasi kedaluwarsa", "Sedang", "Tinggi", "Cek status AKTIF/DICABUT di JDIH sebelum digunakan"),
        ("3", "Kebocoran data ASN", "Rendah", "Sangat Tinggi", "DILARANG memasukkan data pribadi ke prompt AI"),
        ("4", "Prompt injection via email", "Rendah", "Tinggi", "Tinjau selalu sebelum konfirmasi tindakan AI"),
        ("5", "Salah kirim email", "Sedang", "Tinggi", "Periksa penerima, subjek, isi SEBELUM konfirmasi"),
        ("6", "File dihapus/dipindah salah", "Rendah", "Sedang", "Operasi hanya dalam folder DEMO; konfirmasi dulu"),
        ("7", "Tindakan AI tak diinginkan", "Rendah", "Sedang", "Pratinjau + konfirmasi eksplisit sebelum eksekusi"),
        ("8", "API key bocor", "Rendah", "Tinggi", "Simpan di Script Properties; jangan di kode/GitHub"),
        ("9", "Bias dalam jawaban", "Sedang", "Sedang", "Review manusia wajib; cross-check dengan sumber kedua"),
        ("10", "Ketergantungan berlebihan", "Sedang", "Sedang", "AI sebagai ASISTEN, keputusan tetap oleh manusia"),
    ]
    for r in risiko:
        tbl_row(tbl, r, sz=9)
    doc.add_paragraph()
    pb(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# BAGIAN B: CHATGPT PRO (target ~15 halaman = hal 23-37)
# ═══════════════════════════════════════════════════════════════════════════════

def section_B(doc):
    h1(doc, "BAGIAN B: ChatGPT Pro — Kemampuan Lengkap")
    para(doc,
        "ChatGPT Pro adalah versi paling canggih dari ChatGPT yang menawarkan "
        "kemampuan berpikir mendalam (reasoning), terhubung dengan berbagai "
        "aplikasi (Apps), dan dapat melakukan tindakan nyata pada email, file, "
        "dan sistem lainnya. Bagian ini menjelaskan secara detail setiap kemampuan "
        "dan cara memanfaatkannya untuk pekerjaan BKPSDM.")

    # ─── B.1 KEMAMPUAN BERPIKIR ──────────────────────────────────────────────
    h2(doc, "B.1 Kemampuan Berpikir AI (Reasoning)")
    para(doc,
        "ChatGPT memiliki beberapa model dengan tingkat kemampuan berpikir "
        "yang berbeda. Memahami perbedaan ini membantu Anda memilih model "
        "yang tepat untuk setiap tugas.")

    h3(doc, "B.1.1 Jenis Model dan Tingkat Reasoning")
    tbl = tbl_header(doc, ["Model", "Kecepatan", "Reasoning", "Cocok Untuk", "Akses"])
    models = [
        ("GPT-4o", "Cepat", "Standar", "Chat sehari-hari, ringkasan, draf sederhana", "Free/Plus/Pro"),
        ("GPT-4o-mini", "Sangat cepat", "Dasar", "Tugas ringan, brainstorming", "Free/Plus/Pro"),
        ("o1", "Lambat", "Mendalam", "Analisis kompleks, perbandingan regulasi", "Plus/Pro"),
        ("o1-pro", "Sangat lambat", "Paling mendalam", "Riset hukum, analisis multi-dokumen", "Pro only"),
        ("o3-mini", "Sedang", "Menengah-tinggi", "Coding, logika, matematika", "Plus/Pro"),
    ]
    for m in models:
        tbl_row(tbl, m)
    doc.add_paragraph()

    box(doc, "💡 Kapan Menggunakan Model Reasoning Tinggi?",
        "Gunakan o1/o1-pro ketika:\n"
        "• Membandingkan dua regulasi yang saling terkait\n"
        "• Menganalisis kasus yang melibatkan banyak pasal\n"
        "• Menyusun argumentasi hukum yang kompleks\n"
        "• Memverifikasi apakah output sebelumnya konsisten\n\n"
        "Gunakan GPT-4o ketika:\n"
        "• Merangkum email atau dokumen\n"
        "• Membuat draf balasan sederhana\n"
        "• Mencari informasi umum\n"
        "• Tugas-tugas rutin dengan format jelas", "INFO")

    h3(doc, "B.1.2 Cara AI 'Berpikir' — Chain of Thought")
    para(doc,
        "Model reasoning (o1, o1-pro) menggunakan teknik 'Chain of Thought' — "
        "mereka menunjukkan proses berpikirnya langkah demi langkah sebelum "
        "memberikan jawaban akhir. Ini sangat berguna untuk analisis regulasi "
        "yang memerlukan penalaran logis.")
    box(doc, "📌 Contoh: AI Berpikir Tentang Kasus Kepegawaian",
        "PERTANYAAN: Apakah PNS yang sedang menjalani hukuman disiplin ringan\n"
        "berhak mengajukan cuti besar?\n\n"
        "PROSES BERPIKIR AI (ditampilkan):\n"
        "1. Cari definisi 'cuti besar' → PP 11/2017\n"
        "2. Cek syarat cuti besar → Pasal tentang masa kerja & status aktif\n"
        "3. Cek definisi 'hukuman disiplin ringan' → PP 94/2021\n"
        "4. Cek apakah hukuman disiplin ringan mempengaruhi status 'aktif'\n"
        "5. Cek apakah ada larangan eksplisit cuti saat menjalani hukuman\n"
        "6. Kesimpulan dengan ketidakpastian yang ditandai\n\n"
        "JAWABAN: [Draf dengan penanda PERLU VERIFIKASI pada poin-poin\n"
        "yang memerlukan konfirmasi lebih lanjut]", "EXAMPLE")

    # ─── B.2 APPS & KONEKSI ──────────────────────────────────────────────────
    h2(doc, "B.2 ChatGPT Apps — Terhubung dengan Dunia Luar")
    para(doc,
        "ChatGPT dapat terhubung dengan berbagai aplikasi melalui fitur 'Apps'. "
        "Ketika sebuah app terhubung, ChatGPT dapat membaca dan/atau mengambil "
        "tindakan pada aplikasi tersebut atas nama Anda.")

    h3(doc, "B.2.1 Daftar Apps yang Relevan untuk BKPSDM")
    tbl = tbl_header(doc, ["App", "Baca", "Tulis/Tindakan", "Contoh Penggunaan di BKPSDM"])
    apps = [
        ("Gmail", "✅ Baca email, cari, filter", "✅ Draf, kirim, reply, forward", "Rangkum email konsultasi; draf balasan"),
        ("Google Drive", "✅ Cari, buka, baca file", "✅ Buat, edit, pindah, salin file", "Cari regulasi; buat template di folder"),
        ("Google Docs", "✅ Baca dokumen", "✅ Buat, edit dokumen", "Buat draf SOP; edit template surat"),
        ("Google Sheets", "✅ Baca data", "✅ Buat, edit, tambah baris", "Buat log konsultasi; analisis data"),
        ("Google Slides", "✅ Baca presentasi", "✅ Buat slide", "Buat materi pelatihan; presentasi laporan"),
        ("Google Calendar", "✅ Lihat jadwal", "✅ Buat, edit event", "Jadwalkan rapat; reminder deadline"),
        ("GitHub", "✅ Baca repo, issues, PR", "✅ Buat file, commit, PR", "Versioning kode Apps Script; SOP"),
        ("Zapier", "✅ Lihat workflows", "✅ Trigger automations", "Notifikasi email masuk; backup otomatis"),
    ]
    for a in apps:
        tbl_row(tbl, a, sz=9)
    doc.add_paragraph()

    box(doc, "⚠ PENTING: Ketersediaan Apps",
        "Ketersediaan apps bergantung pada:\n"
        "• Paket ChatGPT (Free/Plus/Pro)\n"
        "• Wilayah/negara pengguna\n"
        "• Perangkat (desktop/mobile)\n"
        "• Kebijakan administrator workspace\n"
        "• Versi app yang tersedia di ChatGPT\n\n"
        "JANGAN berasumsi semua app tersedia. Selalu cek di Settings > Apps.\n"
        "[FITUR TERGANTUNG AKUN/ADMIN]", "WARNING")

    h3(doc, "B.2.2 Cara Menghubungkan App")
    para(doc, "Langkah menghubungkan app ke ChatGPT:", bold=True)
    num(doc, "Buka ChatGPT di browser (chat.openai.com).")
    num(doc, "Klik ikon profil → Settings → Apps atau Connected Apps.")
    num(doc, "Pilih app yang ingin dihubungkan (misal: Gmail).")
    num(doc, "Klik 'Connect' → Anda akan diarahkan ke halaman login Google.")
    num(doc, "Login dengan akun Google → Review izin yang diminta.")
    num(doc, "Klik 'Allow' / 'Izinkan' → App terhubung.")
    num(doc, "Kembali ke ChatGPT → app sudah siap digunakan.")
    num(doc, "Untuk mencabut: Google Account → Security → Third-party access → Remove.")

    # ─── B.3 MENGELOLA FILE ──────────────────────────────────────────────────
    h2(doc, "B.3 Mengelola File dengan ChatGPT + Google Drive")
    para(doc,
        "Ketika Google Drive terhubung ke ChatGPT, Anda dapat mengelola file "
        "langsung dari percakapan. Ini sangat berguna untuk mengelola repositori "
        "regulasi, template, dan log konsultasi.")

    h3(doc, "B.3.1 Operasi File yang Dapat Dilakukan")
    tbl = tbl_header(doc, ["Operasi", "Perintah Contoh", "Catatan Penting"])
    ops = [
        ("Mencari file", "'Cari file tentang cuti di Drive saya'", "Pencarian berdasarkan nama dan konten"),
        ("Membaca/Ringkas", "'Ringkas isi file PP_011_2017.pdf'", "File harus bisa dibaca AI (bukan scan)"),
        ("Membuat file baru", "'Buat Google Docs baru berisi template log'", "File dibuat di root atau folder tertentu"),
        ("Membuat folder", "'Buat folder DEMO_BKPSDM_AI di Drive'", "Bisa nested: /folder/subfolder"),
        ("Membuat Spreadsheet", "'Buat Sheets dengan kolom Timestamp, Pertanyaan, Jawaban'", "Termasuk format dan formula"),
        ("Menyalin file", "'Salin template ini ke folder Hasil'", "Konfirmasi lokasi tujuan dulu"),
        ("Memindahkan file", "'Pindahkan laporan ke folder Arsip'", "HATI-HATI: file dipindah permanen"),
        ("Membandingkan", "'Bandingkan PP_011_2017 dan PP_017_2020'", "Butuh kedua file bisa dibaca"),
    ]
    for o in ops:
        tbl_row(tbl, o)
    doc.add_paragraph()
    pb(doc)


    h3(doc, "B.3.2 Struktur Folder yang Disarankan")
    box(doc, "📁 Struktur Folder DEMO_BKPSDM_AI",
        "DEMO_BKPSDM_AI/\n"
        "├── regulasi_aktif/\n"
        "│   ├── PP_011_2017_ManajemenPNS_AKTIF.pdf\n"
        "│   ├── PP_017_2020_PerubahanManajemenPNS_AKTIF.pdf\n"
        "│   ├── PP_094_2021_DisiplinPNS_AKTIF.pdf\n"
        "│   └── PerMenpanRB_006_2022_PenilaianKinerja_AKTIF.pdf\n"
        "├── regulasi_dicabut/\n"
        "│   ├── PP_053_2010_DisiplinPNS_DICABUT.pdf\n"
        "│   └── PP_046_2011_PenilaianPrestasi_DICABUT.pdf\n"
        "├── kasus_fiktif/\n"
        "│   ├── K1_Cuti_Besar.txt\n"
        "│   ├── K2_Disiplin.txt\n"
        "│   ├── K3_KenaikanPangkat.txt\n"
        "│   ├── K4_Mutasi.txt\n"
        "│   ├── K5_Kinerja.txt\n"
        "│   └── K6_Penolakan.txt\n"
        "├── hasil_verifikasi/\n"
        "│   └── (output terverifikasi disimpan di sini)\n"
        "├── template/\n"
        "│   ├── Template_Log_Konsultasi.xlsx\n"
        "│   ├── Template_Prompt.md\n"
        "│   └── Template_SOP_Verifikasi.docx\n"
        "└── kode/\n"
        "    ├── apps_script_template.gs\n"
        "    └── README.md", "CODE")

    # ─── B.4 OTOMATISASI ─────────────────────────────────────────────────────
    h2(doc, "B.4 Otomatisasi dengan ChatGPT")
    para(doc,
        "ChatGPT dapat membantu mengotomasi tugas-tugas berulang, terutama "
        "ketika terhubung dengan apps seperti Gmail, Drive, dan GitHub.")

    h3(doc, "B.4.1 Jenis Otomatisasi yang Dapat Dilakukan")
    tbl = tbl_header(doc, ["Tugas", "Tanpa AI (Manual)", "Dengan ChatGPT + Apps", "Penghematan Waktu"])
    auto = [
        ("Rangkum 10 email", "Buka satu-satu, baca, catat poin", "Satu perintah: 'Rangkum 10 email terbaru tentang cuti'", "~30 mnt → 2 mnt"),
        ("Cari regulasi", "Buka JDIH, ketik keyword, buka PDF satu-satu", "'Cari file tentang disiplin di folder regulasi_aktif'", "~15 mnt → 1 mnt"),
        ("Buat draf jawaban", "Tulis manual, cari pasal, format", "Prompt terstruktur → draf + pasal dalam 30 detik", "~45 mnt → 5 mnt"),
        ("Log konsultasi", "Buka Sheets, isi kolom manual", "'Catat di Sheet LOG: [pertanyaan, jawaban, status]'", "~5 mnt → 30 dtk"),
        ("Buat presentasi", "Buka Slides, tulis slide per slide", "'Buat 10 slide tentang prosedur cuti PNS'", "~60 mnt → 10 mnt"),
    ]
    for a in auto:
        tbl_row(tbl, a, sz=9)
    doc.add_paragraph()


    h3(doc, "B.4.2 ChatGPT Tasks (Penjadwalan)")
    para(doc,
        "ChatGPT Tasks memungkinkan Anda menjadwalkan tugas yang dijalankan "
        "otomatis pada waktu tertentu. [FITUR TERGANTUNG AKUN/ADMIN]")
    box(doc, "📌 Contoh Penggunaan Tasks untuk BKPSDM",
        "TASK 1: Rangkuman Email Harian (08.00 WIB)\n"
        "'Setiap hari kerja jam 08.00, rangkum email masuk 24 jam terakhir\n"
        " yang berisi kata kunci: cuti, disiplin, kenaikan pangkat, mutasi.\n"
        " Kirim rangkuman ke email saya.'\n\n"
        "TASK 2: Reminder Deadline (Setiap Senin 07.30)\n"
        "'Setiap Senin, cek Google Calendar saya untuk 7 hari ke depan.\n"
        " Buat ringkasan deadline dan rencana prioritas minggu ini.'\n\n"
        "TASK 3: Backup Mingguan (Jumat 16.00)\n"
        "'Setiap Jumat, buat ringkasan semua konsultasi yang terjadi\n"
        " minggu ini berdasarkan Sheet LOG.'", "EXAMPLE")

    # ─── B.5 PENGEMBANGAN APLIKASI ──────────────────────────────────────────
    h2(doc, "B.5 Mengembangkan Aplikasi dengan ChatGPT + GitHub")
    para(doc,
        "ChatGPT dapat membantu menulis, mengedit, dan mengelola kode "
        "di GitHub. Untuk BKPSDM, ini berguna untuk mengelola kode "
        "Apps Script, template prompt, dan dokumentasi SOP.")

    h3(doc, "B.5.1 GitHub sebagai Version Control")
    para(doc,
        "GitHub adalah platform untuk menyimpan versi kode dan dokumen. "
        "Dengan GitHub PRIVAT, Anda dapat menyimpan:")
    bul(doc, "Kode Apps Script (template otomasi)")
    bul(doc, "Template prompt yang sudah tervalidasi")
    bul(doc, "SOP verifikasi dan panduan kerja")
    bul(doc, "Catatan perubahan (kapan, siapa, apa yang berubah)")

    box(doc, "⚠ PENTING: GitHub Privat SAJA",
        "• Selalu gunakan repositori PRIVAT — JANGAN publik\n"
        "• JANGAN simpan: API key, data ASN, kasus individual, log\n"
        "• BOLEH simpan: kode, template prompt, SOP, dokumentasi\n"
        "• Aktifkan 2FA (Two-Factor Authentication)\n"
        "• Batasi akses hanya untuk tim yang relevan", "WARNING")
    pb(doc)


    h3(doc, "B.5.2 Workflow ChatGPT + GitHub")
    para(doc, "Contoh alur kerja menggunakan ChatGPT dengan GitHub:")
    num(doc, "Minta ChatGPT membuat/merevisi kode Apps Script.")
    num(doc, "Review kode yang dihasilkan — pastikan tidak ada API key atau data sensitif.")
    num(doc, "Minta ChatGPT commit ke repo GitHub privat: 'Commit file ini ke repo bkpsdm-ai dengan pesan: update template prompt cuti'")
    num(doc, "Jika ada perubahan, minta ChatGPT buat branch baru: 'Buat branch fitur-log-otomatis dan commit perubahan'")
    num(doc, "Review perubahan di GitHub → merge jika sudah benar.")
    num(doc, "Jika ada masalah, rollback ke versi sebelumnya di GitHub.")

    h3(doc, "B.5.3 ChatGPT Projects")
    para(doc,
        "ChatGPT Projects memungkinkan Anda menyimpan konteks, file, dan "
        "instruksi dalam satu 'proyek' yang persisten. Ini berguna untuk "
        "menjaga konsistensi saat bekerja pada topik yang sama berulang kali.")
    box(doc, "📌 Contoh Project untuk BKPSDM",
        "PROJECT: 'Asisten Kepegawaian BKPSDM Demak'\n\n"
        "Instruksi Project:\n"
        "'Kamu adalah asisten kepegawaian untuk BKPSDM Kabupaten Demak.\n"
        " Selalu:\n"
        " - Awali draf dengan DRAF—WAJIB DIVERIFIKASI\n"
        " - Cantumkan dasar hukum (PP/PerMen + pasal)\n"
        " - Tandai yang tidak pasti dengan [PERLU VERIFIKASI]\n"
        " - Jangan mengarang regulasi\n"
        " - Gunakan bahasa formal Indonesia\n"
        " - Minta klarifikasi jika ambigu'\n\n"
        "File yang diunggah ke Project:\n"
        " - PP_011_2017_ManajemenPNS.pdf\n"
        " - PP_094_2021_DisiplinPNS.pdf\n"
        " - Template_Format_Draf.md\n"
        " - Daftar_Regulasi_BKPSDM.xlsx\n\n"
        "Manfaat:\n"
        " ✅ Tidak perlu mengulang instruksi setiap sesi\n"
        " ✅ AI sudah memiliki konteks regulasi\n"
        " ✅ Format output konsisten", "EXAMPLE")
    pb(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# BAGIAN C: OTOMATISASI GRATIS (target ~15 halaman = hal 38-52)
# ═══════════════════════════════════════════════════════════════════════════════

def section_C(doc):
    h1(doc, "BAGIAN C: Otomatisasi Gratis — Apps Script & GitHub")
    para(doc,
        "Bagian ini membahas cara memanfaatkan Google Apps Script dan GitHub "
        "untuk mengotomasi proses kerja BKPSDM TANPA biaya tambahan. Apps Script "
        "gratis untuk pengguna Google Workspace, dan GitHub menyediakan repositori "
        "privat gratis.")

    # ─── C.1 PENGENALAN APPS SCRIPT ──────────────────────────────────────────
    h2(doc, "C.1 Pengenalan Google Apps Script")
    para(doc,
        "Google Apps Script adalah platform pengembangan berbasis JavaScript yang "
        "terintegrasi langsung dengan Google Workspace. Dengan Apps Script, Anda "
        "dapat membuat otomasi yang menghubungkan Gmail, Drive, Sheets, Docs, "
        "dan layanan AI seperti Gemini API.")

    h3(doc, "C.1.1 Apa yang Bisa Dilakukan Apps Script?")
    tbl = tbl_header(doc, ["Kemampuan", "Contoh Penggunaan di BKPSDM", "Koneksi"])
    caps = [
        ("Otomasi Sheets", "Auto-generate draf jawaban dari pertanyaan di Sheets", "Sheets → AI → Sheets"),
        ("Trigger waktu", "Kirim rangkuman harian setiap pagi", "Timer → Sheets → Gmail"),
        ("Email otomatis", "Notifikasi ketika ada konsultasi baru masuk", "Sheets → Gmail"),
        ("Parsing email", "Ekstrak pertanyaan dari email masuk", "Gmail → Sheets"),
        ("File management", "Otomatis pindah file ke folder yang tepat", "Drive → Drive"),
        ("OCR sederhana", "Baca teks dari gambar/PDF scan", "Drive → Vision API"),
        ("Web scraping", "Ambil info dari JDIH (jika tersedia API)", "URL → Sheets"),
        ("Integrasi AI", "Kirim prompt ke Gemini API, terima jawaban", "Sheets → Gemini → Sheets"),
        ("Log audit", "Catat setiap tindakan otomatis dengan timestamp", "Semua → Sheet LOG"),
    ]
    for c in caps:
        tbl_row(tbl, c)
    doc.add_paragraph()

    h3(doc, "C.1.2 Cara Mengakses Apps Script")
    num(doc, "Buka Google Sheets/Docs/Slides yang ingin Anda otomasi.")
    num(doc, "Klik menu Extensions → Apps Script.")
    num(doc, "Editor kode akan terbuka di tab baru (script.google.com).")
    num(doc, "Tulis atau tempel kode → klik ▶ Run → izinkan akses jika diminta.")
    num(doc, "Untuk URL langsung: ketik script.google.com di browser → New Project.")
    para(doc, "[FITUR TERGANTUNG AKUN/ADMIN: Admin workspace mungkin membatasi akses Apps Script]", italic=True, sz=10)
    pb(doc)

    # ─── C.2 TEMPLATE APPS SCRIPT ────────────────────────────────────────────
    h2(doc, "C.2 Template Apps Script Siap Pakai")
    para(doc,
        "Template berikut dirancang khusus untuk BKPSDM Kabupaten Demak. "
        "Peserta TIDAK PERLU mengetik dari nol — cukup salin template ke "
        "Apps Script Editor dan isi konfigurasi yang diperlukan.")

    h3(doc, "C.2.1 Struktur Sheet yang Diperlukan")
    box(doc, "⚙ Buat Google Sheets dengan 4 Sheet Berikut",
        "1. Sheet 'KONSULTASI'\n"
        "   Kolom: A=Timestamp | B=Pertanyaan (Anonim) | C=Jawaban AI |\n"
        "          D=Status Verifikasi | E=Verifikator | F=Tanggal Verifikasi\n\n"
        "2. Sheet 'REGULASI'\n"
        "   Kolom: A=Nama File | B=Judul Resmi | C=Nomor | D=Tahun |\n"
        "          E=Status (AKTIF/DIUBAH/DICABUT) | F=URL JDIH | G=Tgl Akses\n\n"
        "3. Sheet 'LOG'\n"
        "   Kolom: A=Waktu | B=Fungsi | C=Input (200 karakter) |\n"
        "          D=Output (200 karakter) | E=HTTP Status | F=Error Message\n\n"
        "4. Sheet 'UJI_MODEL'\n"
        "   Kolom: A=Provider | B=Model | C=Pertanyaan Uji |\n"
        "          D=Respons | E=Waktu Respons (ms) | F=Tanggal Uji", "STEP")

    h3(doc, "C.2.2 Konfigurasi Script Properties")
    para(doc,
        "API key dan konfigurasi WAJIB disimpan di Script Properties, "
        "BUKAN di dalam kode. Cara mengatur:")
    num(doc, "Di Apps Script Editor, klik ⚙ (Settings) di sidebar kiri.")
    num(doc, "Scroll ke bagian 'Script Properties'.")
    num(doc, "Klik 'Add script property' → isi key dan value:")
    box(doc, "🔐 Script Properties yang Diperlukan",
        "AI_PROVIDER  = gemini\n"
        "AI_MODEL     = gemini-1.5-flash\n"
        "AI_API_URL   = https://generativelanguage.googleapis.com/v1beta\n"
        "AI_API_KEY   = [API key Anda dari Google AI Studio]\n\n"
        "Cara mendapatkan API key Gemini:\n"
        "1. Buka aistudio.google.com\n"
        "2. Login dengan akun Google\n"
        "3. Klik 'Get API key' → 'Create API key'\n"
        "4. Salin key → tempel di Script Properties\n\n"
        "⚠ JANGAN simpan API key di:\n"
        "  ❌ Dalam kode (.gs)\n"
        "  ❌ Di GitHub (meski privat)\n"
        "  ❌ Di Google Docs/Sheets biasa\n"
        "  ❌ Di email atau chat\n"
        "  ✅ HANYA di Script Properties", "WARNING")
    pb(doc)

    h3(doc, "C.2.3 Template Kode Lengkap")
    para(doc,
        "Berikut template lengkap yang siap digunakan. Salin seluruh kode "
        "ke Apps Script Editor (ganti isi file Code.gs default).")
    box(doc, "💻 KODE: Template Apps Script — Asisten Kepegawaian BKPSDM",
        "// ═══════════════════════════════════════════════════════════════\n"
        "// TEMPLATE APPS SCRIPT: Asisten Kepegawaian AI\n"
        "// BKPSDM Kabupaten Demak — v2.0 Juli 2026\n"
        "// ═══════════════════════════════════════════════════════════════\n"
        "// KONFIGURASI: Simpan di Script Properties (Settings > Script Properties)\n"
        "//   AI_PROVIDER  : 'gemini' atau 'openai'\n"
        "//   AI_MODEL     : 'gemini-1.5-flash' / 'gemini-1.5-pro' / 'gpt-4o-mini'\n"
        "//   AI_API_URL   : endpoint sesuai provider\n"
        "//   AI_API_KEY   : API key (JANGAN tulis di sini!)\n"
        "// ═══════════════════════════════════════════════════════════════\n\n"
        "function onOpen() {\n"
        "  const ui = SpreadsheetApp.getUi();\n"
        "  ui.createMenu('🤖 Asisten AI')\n"
        "    .addItem('Tanya Asisten', 'tanyaAsisten')\n"
        "    .addItem('Uji Koneksi API', 'ujiKoneksiAPI')\n"
        "    .addSeparator()\n"
        "    .addItem('Lihat Konfigurasi', 'lihatKonfigurasi')\n"
        "    .addToUi();\n"
        "}\n\n"
        "function tanyaAsisten() {\n"
        "  const sheet = SpreadsheetApp.getActiveSpreadsheet()\n"
        "    .getSheetByName('KONSULTASI');\n"
        "  if (!sheet) {\n"
        "    Browser.msgBox('ERROR: Sheet KONSULTASI tidak ditemukan!');\n"
        "    return;\n"
        "  }\n"
        "  const lastRow = sheet.getLastRow();\n"
        "  if (lastRow < 2) {\n"
        "    Browser.msgBox('Belum ada pertanyaan. Isi kolom B terlebih dahulu.');\n"
        "    return;\n"
        "  }\n"
        "  const pertanyaan = sheet.getRange(lastRow, 2).getValue();\n"
        "  if (!pertanyaan || pertanyaan.toString().trim() === '') {\n"
        "    Browser.msgBox('Kolom Pertanyaan (B) pada baris terakhir kosong.');\n"
        "    return;\n"
        "  }\n"
        "  // Set timestamp\n"
        "  sheet.getRange(lastRow, 1).setValue(new Date());\n"
        "  // Call AI\n"
        "  const jawaban = callAI(pertanyaan);\n"
        "  // Set jawaban dengan label DRAF\n"
        "  sheet.getRange(lastRow, 3).setValue(\n"
        "    'DRAF—WAJIB DIVERIFIKASI PETUGAS BERWENANG:\\n\\n' + jawaban\n"
        "  );\n"
        "  // Set status\n"
        "  sheet.getRange(lastRow, 4).setValue('BELUM DIVERIFIKASI');\n"
        "  // Log\n"
        "  catatLog('tanyaAsisten', pertanyaan, jawaban, 200, '');\n"
        "  Browser.msgBox('Draf jawaban berhasil dibuat. WAJIB DIVERIFIKASI!');\n"
        "}", "CODE")
    pb(doc)


    box(doc, "💻 KODE (lanjutan): callAI() dan callGemini()",
        "function callAI(pertanyaan) {\n"
        "  const props = PropertiesService.getScriptProperties();\n"
        "  const provider = (props.getProperty('AI_PROVIDER') || 'gemini').toLowerCase();\n"
        "  try {\n"
        "    switch (provider) {\n"
        "      case 'gemini': return callGemini(pertanyaan);\n"
        "      // Tambahkan provider lain di sini untuk multi-provider\n"
        "      // case 'openai': return callOpenAI(pertanyaan);\n"
        "      default: return '[ERROR: Provider \"' + provider + '\" tidak dikenal. Cek AI_PROVIDER di Script Properties.]';\n"
        "    }\n"
        "  } catch (e) {\n"
        "    catatLog('callAI', pertanyaan, '', 0, e.message);\n"
        "    return '[ERROR: ' + e.message + ']';\n"
        "  }\n"
        "}\n\n"
        "function callGemini(pertanyaan) {\n"
        "  const props = PropertiesService.getScriptProperties();\n"
        "  const key = props.getProperty('AI_API_KEY');\n"
        "  const model = props.getProperty('AI_MODEL') || 'gemini-1.5-flash';\n"
        "  const baseUrl = props.getProperty('AI_API_URL')\n"
        "    || 'https://generativelanguage.googleapis.com/v1beta';\n\n"
        "  // Validasi\n"
        "  if (!key || key === '[API key Anda dari Google AI Studio]') {\n"
        "    return '[ERROR: API key belum diset. Buka Settings > Script Properties > AI_API_KEY]';\n"
        "  }\n\n"
        "  const url = baseUrl + '/models/' + model + ':generateContent?key=' + key;\n\n"
        "  // System prompt: aturan untuk AI\n"
        "  const systemPrompt = [\n"
        "    'Kamu adalah asisten kepegawaian untuk PNS Indonesia di BKPSDM Kabupaten Demak.',\n"
        "    'ATURAN WAJIB:',\n"
        "    '1. JANGAN mengarang regulasi, pasal, atau nomor peraturan.',\n"
        "    '2. Jika tidak yakin atau tidak ada sumber, nyatakan \"saya tidak memiliki informasi yang cukup\".',\n"
        "    '3. DILARANG membuat keputusan kepegawaian (kenaikan pangkat, hukuman, dll).',\n"
        "    '4. Selalu minta klarifikasi jika pertanyaan ambigu.',\n"
        "    '5. Selalu cantumkan sumber regulasi jika menyebut aturan.',\n"
        "    '6. Awali jawaban dengan: DRAF—WAJIB DIVERIFIKASI PETUGAS BERWENANG.',\n"
        "    '7. Gunakan bahasa Indonesia formal dan ringkas.',\n"
        "  ].join('\\n');\n\n"
        "  const payload = JSON.stringify({\n"
        "    contents: [{\n"
        "      role: 'user',\n"
        "      parts: [{ text: systemPrompt + '\\n\\nPERTANYAAN:\\n' + pertanyaan }]\n"
        "    }]\n"
        "  });\n\n"
        "  const options = {\n"
        "    method: 'post',\n"
        "    contentType: 'application/json',\n"
        "    payload: payload,\n"
        "    muteHttpExceptions: true  // jangan throw, tangkap error\n"
        "  };\n\n"
        "  const resp = UrlFetchApp.fetch(url, options);\n"
        "  const code = resp.getResponseCode();\n\n"
        "  if (code !== 200) {\n"
        "    const errMsg = handleHTTPError(code, resp.getContentText());\n"
        "    catatLog('callGemini', pertanyaan, '', code, errMsg);\n"
        "    return errMsg;\n"
        "  }\n\n"
        "  const json = JSON.parse(resp.getContentText());\n"
        "  const text = json?.candidates?.[0]?.content?.parts?.[0]?.text;\n"
        "  if (!text) {\n"
        "    catatLog('callGemini', pertanyaan, '', code, 'Jawaban kosong dari API');\n"
        "    return '[Jawaban kosong — coba ulangi atau ubah pertanyaan]';\n"
        "  }\n"
        "  return text;\n"
        "}", "CODE")


    box(doc, "💻 KODE (lanjutan): Error Handling, Uji Koneksi, & Log",
        "function handleHTTPError(code, responseText) {\n"
        "  const errors = {\n"
        "    400: 'Bad Request — periksa format pertanyaan',\n"
        "    401: 'Unauthorized — API key salah atau expired',\n"
        "    403: 'Forbidden — API key tidak punya akses ke model ini',\n"
        "    429: 'Rate Limit — terlalu banyak request, tunggu beberapa menit',\n"
        "    500: 'Server Error — Google sedang bermasalah, coba lagi nanti',\n"
        "    503: 'Service Unavailable — server overload, coba lagi nanti',\n"
        "  };\n"
        "  const msg = errors[code] || 'Error tidak dikenal (HTTP ' + code + ')';\n"
        "  return '[ERROR HTTP ' + code + '] ' + msg + '\\nDetail: ' +\n"
        "    responseText.substring(0, 200);\n"
        "}\n\n"
        "function ujiKoneksiAPI() {\n"
        "  const pertanyaanUji = 'Apa itu cuti tahunan PNS? Jawab maksimal 2 kalimat.';\n"
        "  const start = new Date().getTime();\n"
        "  const hasil = callAI(pertanyaanUji);\n"
        "  const elapsed = new Date().getTime() - start;\n\n"
        "  // Log ke sheet UJI_MODEL\n"
        "  const uji = SpreadsheetApp.getActiveSpreadsheet().getSheetByName('UJI_MODEL');\n"
        "  if (uji) {\n"
        "    const props = PropertiesService.getScriptProperties();\n"
        "    uji.appendRow([\n"
        "      props.getProperty('AI_PROVIDER'),\n"
        "      props.getProperty('AI_MODEL'),\n"
        "      pertanyaanUji,\n"
        "      hasil.substring(0, 500),\n"
        "      elapsed,\n"
        "      new Date()\n"
        "    ]);\n"
        "  }\n"
        "  Browser.msgBox('Hasil Uji (' + elapsed + 'ms):\\n\\n' + hasil.substring(0, 300));\n"
        "}\n\n"
        "function lihatKonfigurasi() {\n"
        "  const props = PropertiesService.getScriptProperties().getProperties();\n"
        "  let msg = 'KONFIGURASI SAAT INI:\\n\\n';\n"
        "  for (const [k, v] of Object.entries(props)) {\n"
        "    msg += k + ' = ' + (k.includes('KEY') ? '***' + v.slice(-4) : v) + '\\n';\n"
        "  }\n"
        "  Browser.msgBox(msg);\n"
        "}\n\n"
        "function catatLog(fungsi, input, output, httpStatus, error) {\n"
        "  const log = SpreadsheetApp.getActiveSpreadsheet().getSheetByName('LOG');\n"
        "  if (!log) return;\n"
        "  log.appendRow([\n"
        "    new Date(),\n"
        "    fungsi,\n"
        "    (input || '').toString().substring(0, 200),\n"
        "    (output || '').toString().substring(0, 200),\n"
        "    httpStatus,\n"
        "    error || ''\n"
        "  ]);\n"
        "}", "CODE")
    pb(doc)


    h3(doc, "C.2.4 System Prompt yang Diterapkan")
    para(doc,
        "System prompt berikut ditanamkan dalam kode untuk memastikan AI "
        "berperilaku sesuai aturan yang telah ditetapkan:")
    box(doc, "🛡️ System Prompt — Aturan untuk AI",
        "Kamu adalah asisten kepegawaian untuk PNS Indonesia\n"
        "di BKPSDM Kabupaten Demak.\n\n"
        "ATURAN WAJIB:\n"
        "1. JANGAN mengarang regulasi, pasal, atau nomor peraturan\n"
        "   yang tidak kamu yakini kebenarannya.\n\n"
        "2. Jika tidak ada sumber yang cukup untuk menjawab,\n"
        "   nyatakan: 'Saya tidak memiliki informasi yang cukup\n"
        "   untuk menjawab pertanyaan ini dengan akurat.'\n\n"
        "3. DILARANG membuat keputusan kepegawaian dalam bentuk\n"
        "   apapun — kamu hanya menyusun DRAF untuk diverifikasi.\n\n"
        "4. Selalu MINTA KLARIFIKASI jika pertanyaan ambigu atau\n"
        "   informasi yang diberikan tidak lengkap.\n\n"
        "5. Selalu CANTUMKAN sumber regulasi (nama, nomor, tahun,\n"
        "   pasal) jika menyebut aturan tertentu.\n\n"
        "6. Awali SETIAP jawaban dengan:\n"
        "   'DRAF—WAJIB DIVERIFIKASI PETUGAS BERWENANG'\n\n"
        "7. Gunakan bahasa Indonesia formal, ringkas, dan mudah\n"
        "   dipahami oleh ASN tingkat pemula-menengah.", "INFO")

    h3(doc, "C.2.5 Adapter Multi-Provider (Pengembangan Lanjutan)")
    para(doc,
        "Template di atas menggunakan Gemini sebagai provider utama. "
        "Untuk pengembangan lanjutan, Anda dapat menambahkan adapter "
        "untuk provider lain seperti OpenAI, Claude, atau Groq:")
    box(doc, "💡 Konsep Multi-Provider (Opsional)",
        "Ide dasar:\n"
        "• callAI() membaca AI_PROVIDER dari Script Properties\n"
        "• Berdasarkan provider, memanggil fungsi spesifik:\n"
        "  - 'gemini'  → callGemini()\n"
        "  - 'openai'  → callOpenAI()  ← tambahkan sendiri\n"
        "  - 'groq'    → callGroq()    ← tambahkan sendiri\n\n"
        "Manfaat:\n"
        "• Bisa ganti model tanpa ubah kode\n"
        "• Bisa bandingkan hasil antar-provider\n"
        "• Fallback jika satu provider sedang down\n\n"
        "Catatan: Untuk pelatihan ini, cukup gunakan Gemini.\n"
        "Multi-provider adalah materi lanjutan.", "INFO")
    pb(doc)


    # ─── C.3 GITHUB PRIVAT ────────────────────────────────────────────────────
    h2(doc, "C.3 GitHub Privat untuk Version Control")
    para(doc,
        "GitHub adalah platform untuk menyimpan dan mengelola versi kode, "
        "dokumen, dan konfigurasi. Dengan menggunakan repositori privat, "
        "tim BKPSDM dapat melacak setiap perubahan pada kode otomasi, "
        "template prompt, dan SOP.")

    h3(doc, "C.3.1 Apa yang Disimpan di GitHub (dan yang TIDAK)")
    tbl = tbl_header(doc, ["✅ BOLEH Disimpan", "❌ DILARANG Disimpan"])
    gh_data = [
        ("Kode Apps Script (.gs)", "API key / secret"),
        ("Template prompt (.md)", "Data ASN individual"),
        ("SOP verifikasi (.md, .docx)", "Kasus konsultasi individual"),
        ("Dokumentasi alur kerja", "Log audit (simpan di Sheets)"),
        ("Konfigurasi (tanpa rahasia)", "Password, token, OTP"),
        ("README dan catatan perubahan", "File regulasi PDF (simpan di Drive)"),
    ]
    for g in gh_data:
        tbl_row(tbl, g)
    doc.add_paragraph()

    h3(doc, "C.3.2 Alur Kerja Git Sederhana")
    box(doc, "📊 DIAGRAM: Alur Git untuk Tim BKPSDM",
        "1. TULIS/EDIT kode di Apps Script Editor atau VS Code\n"
        "       ↓\n"
        "2. REVIEW — pastikan tidak ada API key/data sensitif\n"
        "       ↓\n"
        "3. COMMIT — simpan versi dengan pesan deskriptif\n"
        "   Contoh: 'update: perbaiki format draf output'\n"
        "       ↓\n"
        "4. PUSH ke GitHub privat\n"
        "       ↓\n"
        "5. Jika ada masalah → ROLLBACK ke versi sebelumnya\n\n"
        "Tips Pesan Commit:\n"
        "  'fix: perbaiki error HTTP 429 handling'\n"
        "  'feat: tambahkan validasi email sebelum kirim'\n"
        "  'docs: update SOP verifikasi'\n"
        "  'refactor: pisahkan system prompt ke konstanta'", "STEP")

    h3(doc, "C.3.3 Integrasi GitHub dengan ChatGPT")
    para(doc,
        "Jika Anda menggunakan ChatGPT Pro dengan app GitHub terhubung, "
        "Anda dapat mengelola repositori langsung dari percakapan:")
    bul(doc, "'Commit file apps_script_template.gs ke repo bkpsdm-ai'")
    bul(doc, "'Buat branch baru bernama fitur-notifikasi-email'")
    bul(doc, "'Lihat 5 commit terakhir di repo bkpsdm-ai'")
    bul(doc, "'Buat issue: bug pada handling HTTP 429'")
    bul(doc, "'Review pull request #3 dan beri komentar'")
    para(doc, "[FITUR TERGANTUNG AKUN/ADMIN]", italic=True, sz=10)
    pb(doc)

    # ─── C.4 KONEKSI APPS SCRIPT DENGAN LAYANAN LAIN ──────────────────────────
    h2(doc, "C.4 Menghubungkan Apps Script dengan Layanan Lain")
    para(doc,
        "Apps Script dapat terhubung dengan berbagai layanan Google dan "
        "eksternal. Berikut peta koneksi yang relevan untuk BKPSDM:")

    tbl = tbl_header(doc, ["Layanan", "Cara Koneksi", "Contoh Penggunaan", "Izin Diperlukan"])
    koneksi = [
        ("Google Drive", "DriveApp service", "Baca/tulis file di folder DEMO_BKPSDM_AI", "Drive (bawaan)"),
        ("Gmail", "GmailApp service", "Parsing email masuk, kirim notifikasi", "Gmail (izinkan saat run)"),
        ("Google Sheets", "SpreadsheetApp", "Baca pertanyaan, tulis jawaban, log", "Sheets (bawaan)"),
        ("Google Docs", "DocumentApp", "Generate draf surat, template", "Docs (bawaan)"),
        ("Google Calendar", "CalendarApp", "Jadwalkan reminder deadline", "Calendar (izinkan)"),
        ("Gemini API", "UrlFetchApp + API key", "Kirim prompt, terima jawaban AI", "External URL fetch"),
        ("OCR (Vision API)", "UrlFetchApp + API key", "Baca teks dari gambar/PDF scan", "Cloud Vision API key"),
        ("JDIH (jika ada API)", "UrlFetchApp", "Cek status regulasi terbaru", "External URL fetch"),
    ]
    for k in koneksi:
        tbl_row(tbl, k, sz=9)
    doc.add_paragraph()

    h3(doc, "C.4.1 Contoh: Parsing Email Masuk + Auto-Log")
    box(doc, "💻 KODE: Parsing Email Konsultasi ke Sheets",
        "function parseEmailKonsultasi() {\n"
        "  // Cari email dengan label 'konsultasi-kepegawaian'\n"
        "  const threads = GmailApp.search('label:konsultasi-kepegawaian is:unread', 0, 10);\n"
        "  const sheet = SpreadsheetApp.getActiveSpreadsheet()\n"
        "    .getSheetByName('KONSULTASI');\n\n"
        "  threads.forEach(thread => {\n"
        "    const msgs = thread.getMessages();\n"
        "    const lastMsg = msgs[msgs.length - 1];\n"
        "    const body = lastMsg.getPlainBody().substring(0, 500);\n\n"
        "    // Anonimisasi: hapus pola NIP dan nama\n"
        "    const anonBody = body\n"
        "      .replace(/\\d{18}/g, '[NIP_DIHAPUS]')\n"
        "      .replace(/Yth\\.\\s+\\w+\\s+\\w+/g, '[NAMA_DIHAPUS]');\n\n"
        "    // Tambah ke sheet KONSULTASI\n"
        "    sheet.appendRow([\n"
        "      new Date(),                    // Timestamp\n"
        "      anonBody,                      // Pertanyaan (anonim)\n"
        "      '',                            // Jawaban AI (isi nanti)\n"
        "      'BARU',                        // Status\n"
        "      '',                            // Verifikator\n"
        "      ''                             // Tgl verifikasi\n"
        "    ]);\n"
        "    // Tandai sudah dibaca\n"
        "    thread.markRead();\n"
        "  });\n\n"
        "  catatLog('parseEmailKonsultasi', threads.length + ' email diproses', '', 200, '');\n"
        "}", "CODE")
    pb(doc)



# ═══════════════════════════════════════════════════════════════════════════════
# BAGIAN D: DEMONSTRASI (target ~15 halaman = hal 53-67)
# ═══════════════════════════════════════════════════════════════════════════════

def section_D(doc):
    h1(doc, "BAGIAN D: Demonstrasi Praktis")
    para(doc,
        "Bagian ini berisi panduan langkah demi langkah untuk demonstrasi "
        "penggunaan AI di lingkungan kerja BKPSDM. Setiap demo dirancang "
        "agar peserta dapat mengikuti secara langsung dan mempraktikkan "
        "sendiri setelah pelatihan.")

    # ─── D.1 PERSIAPAN DEMO ──────────────────────────────────────────────────
    h2(doc, "D.1 Persiapan Lingkungan Demo")
    box(doc, "✅ CHECKLIST: Persiapan Sebelum Demo",
        "AKUN & AKSES:\n"
        "□ Akun Google demo aktif (BUKAN akun pribadi/dinas berisi data asli)\n"
        "□ ChatGPT sudah login (Free/Plus/Pro)\n"
        "□ App Gmail terhubung ke ChatGPT [FITUR TERGANTUNG AKUN/ADMIN]\n"
        "□ App Google Drive terhubung ke ChatGPT [FITUR TERGANTUNG AKUN/ADMIN]\n"
        "□ Apps Script dapat diakses (script.google.com)\n"
        "□ API key Gemini tersimpan di Script Properties\n\n"
        "FOLDER & FILE:\n"
        "□ Folder DEMO_BKPSDM_AI sudah dibuat di Google Drive\n"
        "□ Sub-folder: /regulasi_aktif, /regulasi_dicabut, /kasus_fiktif,\n"
        "  /hasil_verifikasi, /template\n"
        "□ File kasus fiktif K1-K6 sudah diunggah\n"
        "□ Minimal 3 file regulasi (PDF) sudah diunggah\n"
        "□ Template Sheets (4 sheet) sudah disiapkan\n\n"
        "PERANGKAT:\n"
        "□ Proyektor/layar berbagi siap\n"
        "□ Internet stabil (≥5 Mbps)\n"
        "□ Browser Chrome/Edge terbaru\n"
        "□ Backup hotspot jika WiFi bermasalah", "STEP")
    pb(doc)

    # ─── D.2 DEMO GMAIL ──────────────────────────────────────────────────────
    h2(doc, "D.2 Demo Gmail — Merangkum & Merespons Email")
    para(doc,
        "Demonstrasi ini menunjukkan bagaimana ChatGPT dapat membantu "
        "mengelola email konsultasi kepegawaian: mencari, merangkum, "
        "mengelompokkan urgensi, mengekstrak tindak lanjut, dan menyusun "
        "draf balasan.")
    para(doc, "Prasyarat: App Gmail sudah terhubung ke ChatGPT. [FITUR TERGANTUNG AKUN/ADMIN]", italic=True, sz=10)

    h3(doc, "D.2.1 Langkah 1: Mencari Email")
    box(doc, "🎯 LANGKAH PRAKTIK: Cari Email Fiktif",
        "PERINTAH ke ChatGPT:\n"
        "'Cari 5 email terbaru di inbox saya yang mengandung kata\n"
        " kunci: cuti, izin, atau mutasi. Untuk setiap email, tampilkan:\n"
        " - Pengirim\n"
        " - Tanggal\n"
        " - Subjek\n"
        " - Ringkasan isi (maks 2 kalimat)'\n\n"
        "YANG DIHARAPKAN:\n"
        "ChatGPT akan membaca inbox → filter email → tampilkan tabel ringkasan.\n\n"
        "VERIFIKASI:\n"
        "□ Apakah hanya email fiktif/demo yang ditampilkan?\n"
        "□ Apakah tidak ada data pribadi ASN asli yang muncul?\n"
        "□ Jika ada data asli → STOP, ganti ke akun demo.", "STEP")

    h3(doc, "D.2.2 Langkah 2: Mengelompokkan Urgensi")
    box(doc, "🎯 LANGKAH PRAKTIK: Kelompokkan Email",
        "PERINTAH ke ChatGPT:\n"
        "'Dari email-email yang kamu temukan tadi, kelompokkan berdasarkan\n"
        " urgensi:\n"
        " 🔴 SEGERA (perlu ditindaklanjuti hari ini)\n"
        " 🟡 NORMAL (perlu ditindaklanjuti minggu ini)\n"
        " 🟢 INFORMASI (tidak perlu tindak lanjut segera)\n\n"
        " Berikan alasan singkat mengapa email tersebut masuk kategori itu.'\n\n"
        "YANG DIHARAPKAN:\n"
        "Tabel dengan kolom: Email | Urgensi | Alasan", "STEP")

    h3(doc, "D.2.3 Langkah 3: Ekstrak Tindak Lanjut")
    box(doc, "🎯 LANGKAH PRAKTIK: Buat Daftar Tindak Lanjut",
        "PERINTAH ke ChatGPT:\n"
        "'Dari email-email kategori SEGERA dan NORMAL, buatkan daftar\n"
        " tindak lanjut yang diperlukan. Format:\n"
        " - Nomor\n"
        " - Tindakan yang harus dilakukan\n"
        " - Deadline yang disarankan\n"
        " - PIC (posisi/jabatan, bukan nama)'\n\n"
        "YANG DIHARAPKAN:\n"
        "Daftar terstruktur yang bisa dijadikan to-do list.", "STEP")

    h3(doc, "D.2.4 Langkah 4: Draf Balasan Email")
    box(doc, "🎯 LANGKAH PRAKTIK: Susun Draf Balasan",
        "PERINTAH ke ChatGPT:\n"
        "'Untuk email dari [pengirim fiktif] tentang pertanyaan cuti besar,\n"
        " susun draf balasan dengan ketentuan:\n"
        " - Awali dengan salam formal\n"
        " - Jawab berdasarkan PP 11/2017 (jika tahu)\n"
        " - Tandai dengan DRAF—WAJIB DIVERIFIKASI\n"
        " - Akhiri dengan catatan: untuk kepastian silakan konsultasi langsung\n"
        " - Format email resmi ASN'\n\n"
        "SEBELUM KIRIM — WAJIB PERIKSA:\n"
        "□ Penerima (To): apakah benar akun uji/fiktif?\n"
        "□ CC/BCC: apakah ada yang tidak seharusnya?\n"
        "□ Subjek: apakah sesuai?\n"
        "□ Lampiran: apakah benar (tidak ada dokumen rahasia)?\n"
        "□ Isi: apakah tidak mengandung data pribadi asli?\n"
        "□ DRAF label: apakah sudah ditandai?\n"
        "□ KONFIRMASI: ketik 'Ya, kirim' HANYA setelah semua diperiksa.", "WARNING")
    pb(doc)

    # ─── D.3 DEMO DRIVE ──────────────────────────────────────────────────────
    h2(doc, "D.3 Demo Google Drive — Cari, Ringkas, dan Kelola File")
    para(doc,
        "Demonstrasi ini menunjukkan bagaimana ChatGPT dapat membantu "
        "mengelola file regulasi dan dokumen kerja di Google Drive.")
    para(doc, "Prasyarat: App Google Drive sudah terhubung ke ChatGPT. [FITUR TERGANTUNG AKUN/ADMIN]", italic=True, sz=10)

    h3(doc, "D.3.1 Mencari dan Meringkas File")
    box(doc, "🎯 LANGKAH PRAKTIK: Cari & Ringkas Regulasi",
        "PERINTAH 1 (Cari):\n"
        "'Cari file di folder DEMO_BKPSDM_AI yang berkaitan dengan\n"
        " cuti PNS. Tampilkan nama file, lokasi folder, dan tanggal modifikasi.'\n\n"
        "PERINTAH 2 (Ringkas):\n"
        "'Ringkas isi file PP_011_2017_ManajemenPNS_AKTIF.pdf dalam 10 poin\n"
        " utama, khususnya bagian tentang cuti.'\n\n"
        "PERINTAH 3 (Bandingkan):\n"
        "'Bandingkan PP_011_2017 dan PP_017_2020 — apa saja yang berubah?\n"
        " Tampilkan dalam tabel: Aspek | PP 11/2017 | PP 17/2020 | Perubahan'", "STEP")

    h3(doc, "D.3.2 Membuat dan Mengelola File")
    box(doc, "🎯 LANGKAH PRAKTIK: Buat File Baru di Drive",
        "PERINTAH 1 (Buat Docs):\n"
        "'Buat Google Docs baru di folder DEMO_BKPSDM_AI/template\n"
        " dengan nama: Template_Jawaban_Konsultasi\n"
        " Isinya: format standar draf output AI yang kita bahas di Bagian A.'\n\n"
        "PERINTAH 2 (Buat Sheets):\n"
        "'Buat Google Sheets baru di folder DEMO_BKPSDM_AI\n"
        " dengan nama: Log_Konsultasi_Demo\n"
        " Buat 4 sheet: KONSULTASI, REGULASI, LOG, UJI_MODEL\n"
        " Isi header kolom sesuai struktur yang kita bahas di Bagian C.'\n\n"
        "PERINTAH 3 (Buat Folder):\n"
        "'Buat folder baru di DEMO_BKPSDM_AI bernama: laporan_bulanan'\n\n"
        "SEBELUM EKSEKUSI — WAJIB:\n"
        "□ Periksa nama file/folder yang akan dibuat\n"
        "□ Periksa LOKASI folder tujuan (harus di DEMO_BKPSDM_AI)\n"
        "□ Pratinjau isi jika membuat dokumen\n"
        "□ Konfirmasi sebelum ChatGPT mengeksekusi", "STEP")

    h3(doc, "D.3.3 Memindahkan dan Menyalin File")
    box(doc, "⚠ PERINGATAN: Operasi File Berbahaya",
        "Operasi PINDAH dan HAPUS bersifat permanen!\n\n"
        "ATURAN:\n"
        "• Operasi file HANYA dalam folder DEMO_BKPSDM_AI\n"
        "• SELALU pratinjau sebelum konfirmasi\n"
        "• JANGAN pindahkan file ke luar folder demo\n"
        "• JANGAN hapus file regulasi asli\n"
        "• Jika ragu → BATALKAN dan tanya fasilitator\n\n"
        "PERINTAH AMAN:\n"
        "'Salin file Template_Jawaban_Konsultasi ke folder hasil_verifikasi'\n"
        "(salin = duplikat, file asli tetap ada)\n\n"
        "PERINTAH BERBAHAYA (hati-hati):\n"
        "'Pindahkan file X ke folder Y'\n"
        "(pindah = file hilang dari lokasi semula)", "WARNING")
    pb(doc)


    # ─── D.4 DEMO CHATGPT PRO ────────────────────────────────────────────────
    h2(doc, "D.4 Demo ChatGPT Pro — Kemampuan Lengkap")
    para(doc,
        "ChatGPT Pro dapat menjalankan tindakan pada email, file Drive, "
        "Docs, Sheets, Slides, dan Calendar — termasuk membuat folder, "
        "spreadsheet, dan presentasi — HANYA jika app terkait tersedia "
        "dan terhubung.")
    box(doc, "⚠ PENTING: Batasan ChatGPT Apps",
        "Kemampuan ChatGPT untuk bertindak TERGANTUNG pada:\n"
        "• App tersedia di paket Anda (Free/Plus/Pro)\n"
        "• App sudah dihubungkan (connected)\n"
        "• Izin yang diberikan (baca saja / baca+tulis)\n"
        "• Kebijakan admin workspace\n"
        "• Wilayah/negara (beberapa app belum tersedia)\n"
        "• Perangkat (desktop vs mobile)\n"
        "• Versi ChatGPT yang digunakan\n\n"
        "JANGAN berasumsi semua fitur tersedia.\n"
        "Selalu cek: Settings > Connected Apps\n"
        "[FITUR TERGANTUNG AKUN/ADMIN]", "WARNING")

    h3(doc, "D.4.1 Contoh Perintah Aksi ChatGPT Pro")
    tbl = tbl_header(doc, ["Kategori", "Perintah Contoh", "App Diperlukan", "Risiko"])
    cmds = [
        ("Email: Cari", "'Cari email dari Budi tentang mutasi bulan lalu'", "Gmail (baca)", "Rendah"),
        ("Email: Draf", "'Buat draf balasan tentang prosedur cuti'", "Gmail (tulis)", "Sedang — cek sebelum kirim"),
        ("Email: Kirim", "'Kirim email ini ke akun-uji@gmail.com'", "Gmail (kirim)", "Tinggi — wajib konfirmasi!"),
        ("Drive: Cari", "'Cari file regulasi tentang disiplin'", "Drive (baca)", "Rendah"),
        ("Drive: Buat", "'Buat spreadsheet log konsultasi baru'", "Drive (tulis)", "Sedang — cek lokasi"),
        ("Drive: Pindah", "'Pindahkan file ke folder arsip'", "Drive (tulis)", "Tinggi — permanen!"),
        ("Docs: Buat", "'Buat dokumen SOP verifikasi'", "Docs (tulis)", "Sedang"),
        ("Sheets: Edit", "'Tambahkan baris baru di sheet LOG'", "Sheets (tulis)", "Rendah"),
        ("Slides: Buat", "'Buat presentasi 10 slide tentang regulasi cuti'", "Slides (tulis)", "Sedang"),
        ("Calendar: Buat", "'Jadwalkan rapat evaluasi pilot AI hari Jumat jam 10'", "Calendar (tulis)", "Sedang"),
    ]
    for c in cmds:
        tbl_row(tbl, c, sz=9)
    doc.add_paragraph()

    h3(doc, "D.4.2 Perbedaan Mode Akses File")
    tbl = tbl_header(doc, ["Mode", "Cara Kerja", "Baca", "Tulis", "Contoh"])
    modes = [
        ("Upload Manual", "Drag & drop file ke chat", "✅", "❌", "Upload PDF regulasi → minta ringkasan"),
        ("App Baca", "Drive app (read-only)", "✅", "❌", "'Cari file tentang cuti di Drive saya'"),
        ("App Tulis", "Drive app (full access)", "✅", "✅", "'Buat folder baru di DEMO_BKPSDM_AI'"),
        ("Apps Script", "Kode otomasi + API", "✅", "✅", "Fungsi tanyaAsisten() berjalan otomatis"),
    ]
    for m in modes:
        tbl_row(tbl, m)
    doc.add_paragraph()

    h3(doc, "D.4.3 Demo Langkah demi Langkah")
    para(doc, "Berikut demonstrasi lengkap kemampuan ChatGPT Pro:", bold=True)
    num(doc, "Buka ChatGPT → pastikan app Gmail dan Drive terhubung (ikon di sidebar).")
    num(doc, "DEMO 1 — Cari email: 'Cari 3 email terakhir tentang cuti di inbox saya. Tampilkan pengirim dan ringkasan.'")
    num(doc, "DEMO 2 — Rangkum: 'Rangkum email dari [pengirim fiktif] dalam 3 poin tindak lanjut.'")
    num(doc, "DEMO 3 — Cari file: 'Cari file PP_011_2017 di folder DEMO_BKPSDM_AI. Ringkas pasal tentang cuti.'")
    num(doc, "DEMO 4 — Buat Sheets: 'Buat Google Sheets baru di DEMO_BKPSDM_AI dengan 4 sheet sesuai template kita.'")
    num(doc, "DEMO 5 — Buat Slides: 'Buat presentasi 5 slide tentang alur konsultasi AI di BKPSDM.'")
    num(doc, "DEMO 6 — Draf email: 'Susun draf balasan untuk pertanyaan cuti besar. Format formal ASN.'")
    num(doc, "DEMO 7 — Verifikasi sebelum kirim: Tinjau penerima, subjek, isi → konfirmasi atau batalkan.")
    num(doc, "DEMO 8 — Log: 'Tambahkan baris baru di Sheet LOG: tanyaAsisten, [pertanyaan], [jawaban], 200, -'")

    box(doc, "🔑 PRINSIP UTAMA DEMO",
        "Setiap tindakan AI harus mengikuti alur:\n\n"
        "  PERINTAH → PRATINJAU → PERIKSA → KONFIRMASI → EKSEKUSI\n\n"
        "JANGAN PERNAH konfirmasi tindakan tanpa membaca pratinjau terlebih dahulu.\n"
        "Jika ragu → BATALKAN → tanya fasilitator → ulangi dengan perintah yang lebih spesifik.", "WARNING")
    pb(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# PENUTUP (target ~3 halaman = hal 68-70)
# ═══════════════════════════════════════════════════════════════════════════════

def section_closing(doc):
    h1(doc, "PENUTUP")

    h2(doc, "Ringkasan Materi")
    tbl = tbl_header(doc, ["Bagian", "Materi Utama", "Luaran"])
    ringkasan = [
        ("A: Dasar AI", "Cara kerja AI, memilih tools, keamanan koneksi, prompt engineering, verifikasi", "Pemahaman fundamental AI untuk ASN"),
        ("B: ChatGPT Pro", "Model reasoning, Apps, file management, otomatisasi, GitHub, Projects", "Kemampuan menggunakan ChatGPT secara produktif"),
        ("C: Otomatisasi", "Apps Script, template kode, system prompt, GitHub, koneksi layanan", "Template otomasi siap pakai"),
        ("D: Demo", "Gmail, Drive, ChatGPT Pro actions, langkah praktik terstruktur", "Pengalaman praktik langsung"),
    ]
    for r in ringkasan:
        tbl_row(tbl, r)
    doc.add_paragraph()

    h2(doc, "Luaran Peserta")
    bul(doc, "✅ Satu prompt terstruktur (P-T-O-K-S-B-F-K) siap pakai untuk konsultasi kepegawaian.")
    bul(doc, "✅ Satu draf jawaban terverifikasi dengan dasar hukum yang dicantumkan.")
    bul(doc, "✅ Satu rencana pilot 30 hari untuk implementasi di unit kerja.")
    doc.add_paragraph()

    h2(doc, "Rencana Pilot 30 Hari")
    tbl = tbl_header(doc, ["Minggu", "Kegiatan", "Target", "Dukungan"])
    pilot = [
        ("1\n(Hari 1-7)", "Setup: Buat folder DEMO, unggah 3 regulasi, coba 1 prompt terstruktur", "1 prompt + 1 draf terverifikasi", "Template dari pelatihan"),
        ("2\n(Hari 8-14)", "Eksplorasi: Coba ChatGPT + Gmail (rangkum 5 email), buat log manual", "Log 5 konsultasi terdokumentasi", "SOP verifikasi"),
        ("3\n(Hari 15-21)", "Otomasi: Setup Apps Script template, uji koneksi API, coba 2 kasus", "Template berjalan, 2 kasus teruji", "Kode template"),
        ("4\n(Hari 22-30)", "Evaluasi: Presentasi ke atasan, susun SOP, rencanakan scale-up", "SOP 1 halaman, rencana tindak lanjut", "Bimbingan tim"),
    ]
    for p in pilot:
        tbl_row(tbl, p, sz=9)
    doc.add_paragraph()


    h2(doc, "Glosarium")
    glos = [
        ("AI (Artificial Intelligence)", "Teknologi komputer yang mampu melakukan tugas kognitif: memahami bahasa, mengenali pola, menghasilkan teks."),
        ("LLM (Large Language Model)", "Model AI skala besar yang dilatih pada triliunan kata untuk memahami dan menghasilkan teks."),
        ("Halusinasi AI", "Output AI yang tampak meyakinkan tetapi tidak akurat atau tidak memiliki dasar faktual."),
        ("Prompt", "Instruksi atau pertanyaan yang diberikan pengguna kepada AI untuk menghasilkan respons."),
        ("Prompt Engineering", "Teknik menyusun prompt yang terstruktur agar output AI lebih akurat dan berguna."),
        ("OAuth", "Protokol otorisasi yang memungkinkan aplikasi mengakses akun pengguna tanpa mengetahui password."),
        ("API Key", "Kode rahasia untuk akses programatik ke layanan AI; JANGAN disimpan di kode."),
        ("Apps Script", "Platform pengembangan berbasis JavaScript dari Google untuk mengotomasi Google Workspace."),
        ("JDIH", "Jaringan Dokumentasi dan Informasi Hukum; sumber resmi regulasi pemerintah Indonesia."),
        ("Anonimisasi", "Proses menghapus identitas pribadi dari data sebelum diproses AI."),
        ("System Prompt", "Instruksi yang ditanamkan dalam kode untuk mengatur perilaku AI secara konsisten."),
        ("Version Control (Git)", "Sistem pelacakan perubahan pada kode/dokumen; memungkinkan rollback ke versi sebelumnya."),
        ("Chain of Thought", "Teknik di mana AI menunjukkan proses berpikirnya langkah demi langkah."),
        ("Token", "Unit terkecil teks yang diproses AI (bisa berupa kata, sub-kata, atau karakter)."),
        ("Generative AI", "AI yang mampu menghasilkan konten baru (teks, gambar, kode) berdasarkan pola yang dipelajari."),
    ]
    for term, defn in glos:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{term}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(defn)
        r2.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    pb(doc)

    h2(doc, "Penutup")
    box(doc, "🙏 Terima Kasih & Selamat Berinovasi!",
        "Pelatihan ini adalah langkah awal transformasi digital BKPSDM Kabupaten Demak.\n\n"
        "INGAT SELALU:\n"
        "✅ AI adalah ASISTEN — bukan pengganti penilaian profesional Anda\n"
        "✅ VERIFIKASI setiap output AI sebelum digunakan\n"
        "✅ JAGA kerahasiaan data ASN — jangan pernah masukkan data asli ke AI\n"
        "✅ DOKUMENTASIKAN setiap penggunaan AI di log\n"
        "✅ EVALUASI dan tingkatkan proses secara berkala\n\n"
        "Pertanyaan & dukungan pasca-pelatihan:\n"
        "[PERLU DATA: email/kontak panitia]\n\n"
        "Repositori kode & template:\n"
        "[PERLU DATA: URL repo GitHub privat tim]\n\n"
        "Jadwal bimbingan lanjutan:\n"
        "[PERLU DATA: jadwal konsultasi pasca-pelatihan]", "INFO")


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    os.makedirs(OUT, exist_ok=True)
    print("🔨 Building Modul Pelatihan AI v2.0...")
    doc = setup_doc()

    # Build all sections
    cover(doc)
    toc(doc)
    section_A(doc)
    section_B(doc)
    section_C(doc)
    section_D(doc)
    section_closing(doc)

    # Add footer
    add_footer(doc)

    # Save
    doc.save(DOCX_PATH)
    size_kb = os.path.getsize(DOCX_PATH) // 1024
    print(f"✅ DOCX saved: {DOCX_PATH}")
    print(f"   File size: {size_kb} KB")
    print(f"   Target: ~70 pages (verify in Word)")
    print(f"\n📌 To get PDF: Open in Word → File → Export → PDF")


if __name__ == '__main__':
    main()
